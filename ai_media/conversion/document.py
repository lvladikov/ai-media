"""
Document conversion module for AI-Media.

Supports: md, html, pdf, docx, rtf, txt, json conversions.
Uses Markdown as intermediate format for hub-and-spoke conversion.
"""

import os
import json as json_module
from pathlib import Path

from ..utils.interaction import check_overwrite

SUPPORTED_FORMATS = ["md", "html", "pdf", "docx", "rtf", "txt", "json", "xhtml"]


def _read_to_markdown(input_path, input_format, ocr_model="florence"):
    """Read input file and convert to Markdown."""
    markdown_content = ""
    
    if input_format == "md":
        with open(input_path, "r", encoding="utf-8") as f:
            markdown_content = f.read()
    
    elif input_format in ["html", "xhtml"]:
        try:
            import html2text
            h = html2text.HTML2Text()
            h.ignore_links = False
            h.ignore_images = False
            with open(input_path, "r", encoding="utf-8") as f:
                markdown_content = h.handle(f.read())
        except ImportError:
            from bs4 import BeautifulSoup
            with open(input_path, "r", encoding="utf-8") as f:
                soup = BeautifulSoup(f.read(), "html.parser")
                markdown_content = soup.get_text()
            print("   ⚠️ html2text not installed, using basic text extraction")
    
    elif input_format == "docx":
        import docx
        doc = docx.Document(input_path)
        lines = []
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading 1'):
                lines.append(f"# {para.text}")
            elif para.style.name.startswith('Heading 2'):
                lines.append(f"## {para.text}")
            elif para.style.name.startswith('Heading 3'):
                lines.append(f"### {para.text}")
            else:
                lines.append(para.text)
        markdown_content = "\n\n".join(lines)
    
    elif input_format == "txt":
        with open(input_path, "r", encoding="utf-8") as f:
            markdown_content = f.read()
    
    elif input_format == "json":
        with open(input_path, "r", encoding="utf-8") as f:
            data = json_module.load(f)
        if isinstance(data, dict):
            markdown_content = data.get("content", "") or data.get("markdown", "") or data.get("text", "") or str(data)
        else:
            markdown_content = str(data)
    
    elif input_format == "pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(input_path)
            text_parts = []
            for page in reader.pages:
                text_parts.append(page.extract_text() or "")
            markdown_content = "\n\n".join(text_parts)
            
            if not markdown_content.strip():
                print(f"   ⚠️ No text layer found. Attempting OCR ({ocr_model})...")
                try:
                    from .ocr import image_to_text
                    from PIL import Image
                    import io
                    import tempfile
                    
                    ocr_texts = []
                    
                    # Extract images from each page
                    for page_num, page in enumerate(reader.pages):
                        page_images = []
                        
                        # Try to extract images from the page
                        if "/XObject" in page["/Resources"]:
                            x_objects = page["/Resources"]["/XObject"].get_object()
                            for obj_name in x_objects:
                                obj = x_objects[obj_name]
                                if obj["/Subtype"] == "/Image":
                                    try:
                                        # Get image data
                                        if "/Filter" in obj:
                                            if obj["/Filter"] == "/DCTDecode":
                                                # JPEG
                                                img_data = obj._data
                                                img = Image.open(io.BytesIO(img_data))
                                            elif obj["/Filter"] == "/FlateDecode":
                                                # Raw image data
                                                width = obj["/Width"]
                                                height = obj["/Height"]
                                                color_space = obj.get("/ColorSpace", "/DeviceRGB")
                                                mode = "RGB" if "RGB" in str(color_space) else "L"
                                                img = Image.frombytes(mode, (width, height), obj._data)
                                            else:
                                                continue
                                            page_images.append(img)
                                    except Exception:
                                        continue
                        
                        if page_images:
                            # OCR each image from the page
                            for img_idx, img in enumerate(page_images):
                                print(f"      Scanning page {page_num+1}, image {img_idx+1}...")
                                
                                # Resize large images to prevent GPU OOM
                                max_dim = 2048
                                if img.width > max_dim or img.height > max_dim:
                                    ratio = min(max_dim / img.width, max_dim / img.height)
                                    new_size = (int(img.width * ratio), int(img.height * ratio))
                                    img = img.resize(new_size, Image.Resampling.LANCZOS)
                                    print(f"         (Resized to {new_size[0]}x{new_size[1]} for OCR)")
                                
                                with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as tmp:
                                    # Convert to RGB if needed (JPEG doesn't support RGBA)
                                    if img.mode in ('RGBA', 'P'):
                                        img = img.convert('RGB')
                                    img.save(tmp.name, "JPEG", quality=90)
                                    text = image_to_text(tmp.name, model_type=ocr_model)
                                    if text:
                                        ocr_texts.append(text)
                                    os.unlink(tmp.name)
                    
                    if ocr_texts:
                        markdown_content = "\n\n".join(ocr_texts)
                    else:
                        raise ValueError("No extractable images found in PDF for OCR")
                        
                except Exception as e:
                    raise ValueError(f"No text found and OCR failed: {e}")
                
            print("   ⚠️ PDF conversion extract text (native or OCR)")
        except ImportError:
            raise ImportError("pypdf required for PDF reading. Install: pip install pypdf")
    
    elif input_format == "rtf":
        try:
            from striprtf.striprtf import rtf_to_text
            with open(input_path, "r", encoding="utf-8") as f:
                markdown_content = rtf_to_text(f.read())
            print("   ⚠️ RTF conversion extracts text only (formatting lost)")
        except ImportError:
            # Fallback for when striprtf is not available
            import re
            def fallback_rtf_to_text(rtf_str):
                # Very basic RTF stripper
                pattern = re.compile(r"\\([a-z]{1,32})(-?\d{1,10})?[ ]?|\\'([0-9a-f]{2})|\\([^a-z])|([{}])|[\r\n]+|(.)", re.I)
                destinations = frozenset((
                    'ftnid','ftnsep','ftnsepc','annotation','atnid','atnref','atntime','atrfend','atrfstart',
                    'author','background','bkmkend','bkmkstart','blipuid','buptim','category','colorschememapping',
                    'colortbl','comment','company','creatim','datafield','datastore','defchp','defpap','do',
                    'doccomm','docvar','dptxbxtext','ebcend','ebcstart','factoidname','falt','fchars','ffdeftext',
                    'ffentrymcr','ffexitmcr','ffformat','ffhelptext','ffl','ffname','ffstattext','field',
                    'filetbl','fldinst','fldrslt','fldtype','fname','fontemb','fontfile','fonttbl','footer',
                    'footerf','footerl','footerr','footnote','formfield','generator','gridtbl','header',
                    'headerf','headerl','headerr','hl','hlfr','hllink','hlsrc','hsv','htmltag','info',
                    'keycode','keywords','latency','lchars','levelnumbers','leveltext','list','listlevel',
                    'listname','listoverride','listoverridetable','listpicture','liststylename','listtable',
                    'listtext','lsdlockedexcept','macc','maccPr','mailmerge','malformed','margins','mati',
                    'matn','mhtmltag','mmath','mmathPr','mnum','mpocket','mtype','mxml','nesttableprops',
                    'nextfile','nonesttables','obj','objalias','objclass','objdata','object','objname',
                    'objsect','objtime','oldcprops','oldpprops','oldsprops','oldtprops','oleclsid','operator',
                    'panose','password','passwordhash','pgp','pgptbl','picprop','pict','pn','pnseclvl',
                    'pntext','pntxta','pntxtb','printim','private','propname','protend','protstart','protusertbl',
                    'pxe','result','revtbl','revtim','rsidtbl','rxe','shp','shpgrp','shpinst','shppict',
                    'shprslt','shptxt','sn','sp','staticval','stylesheet','subject','sv','svb','tc',
                    'template','themedata','title','txe','ud','upr','userprops','wgrffmtfilter','windowcaption',
                    'writereservation','writereservhash','xe','xform','xmlattrname','xmlattrvalue','xmlclose',
                    'xmlname','xmlnstbl','xmlopen',
                ))
                stack = []
                ignorable = False       # Whether this group (and all inside it) are "ignorable".
                ucskip = 1              # Number of ASCII characters to skip after a unicode character.
                curskip = 0             # Number of ASCII characters left to skip
                out = []                # Output buffer.

                for match in pattern.finditer(rtf_str):
                    word,arg,hex,char,brace,tchar = match.groups()
                    if brace:
                        curskip = 0
                        if brace == '{':
                            stack.append((ucskip,ignorable))
                        elif brace == '}':
                            if stack:
                                ucskip,ignorable = stack.pop()
                    elif char: # \x (not a letter)
                        curskip = 0
                        if char == '~':
                            if not ignorable: out.append('\xA0')
                        elif char in '{}\\':
                            if not ignorable: out.append(char)
                        elif char == '*':
                            ignorable = True
                    elif word: # \foo
                        curskip = 0
                        if word in destinations:
                            ignorable = True
                        elif ignorable:
                            pass
                        elif word == 'par' or word == 'line' or word == 'row':
                            out.append('\n')
                        elif word == 'tab':
                            out.append('\t')
                    elif hex: # \'xx
                        if curskip > 0:
                            curskip -= 1
                        elif not ignorable:
                            out.append(chr(int(hex,16)))
                    elif tchar:
                        if curskip > 0:
                            curskip -= 1
                        elif not ignorable:
                            out.append(tchar)
                return "".join(out)
            
            with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
                markdown_content = fallback_rtf_to_text(f.read())
            print("   ⚠️ RTF conversion using fallback regex parser (formatting lost)")

    
    return markdown_content


def _write_from_markdown(markdown_content, output_path, output_format):
    """Write Markdown content to output format."""
    import markdown as md_module
    
    if output_format == "md":
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(markdown_content)
            
    elif output_format in ["html", "xhtml"]:
        html = md_module.markdown(markdown_content, extensions=['extra', 'codehilite'])
        full_html = (
            f"<!DOCTYPE html><html><head><meta charset='utf-8'><title>Document</title>"
            f"<style>body{{font-family:sans-serif;max-width:800px;margin:2em auto;padding:1em;line-height:1.6}}"
            f"pre{{background:#f4f4f4;padding:1em;border-radius:5px}}</style></head>"
            f"<body>{html}</body></html>"
        )
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(full_html)
            
    elif output_format == "json":
        data = {"content": markdown_content, "html": md_module.markdown(markdown_content)}
        with open(output_path, "w", encoding="utf-8") as f:
            json_module.dump(data, f, indent=2)

    elif output_format == "txt":
        from bs4 import BeautifulSoup
        html = md_module.markdown(markdown_content)
        text = BeautifulSoup(html, "html.parser").get_text()
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text)
            
    elif output_format == "docx":
        import docx
        doc = docx.Document()
        for line in markdown_content.split('\n'):
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            else:
                doc.add_paragraph(line)
        doc.save(output_path)
        
    elif output_format == "pdf":
        from xhtml2pdf import pisa
        html_content = md_module.markdown(markdown_content, extensions=['extra', 'fenced_code', 'tables'])
        full_html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>body {{ font-family: Helvetica, sans-serif; font-size: 10pt; }}</style>
</head><body>{html_content}</body></html>"""
        with open(output_path, "wb") as f:
            pisa.CreatePDF(full_html, dest=f)
            
    elif output_format == "rtf":
        # Basic RTF generation
        rtf_lines = [r'{\rtf1\ansi\deff0', r'{\fonttbl{\f0 Helvetica;}}', r'\f0\fs24']
        for line in markdown_content.split('\n'):
            escaped = line.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
            if line.startswith('# '):
                rtf_lines.append(r'\pard\sb400\sa200\b\fs48 ' + escaped[2:] + r'\b0\fs24\par')
            elif line.startswith('## '):
                rtf_lines.append(r'\pard\sb300\sa150\b\fs36 ' + escaped[3:] + r'\b0\fs24\par')
            else:
                rtf_lines.append(r'\pard\sa100 ' + escaped + r'\par')
        rtf_lines.append('}')
        with open(output_path, "w", encoding="utf-8") as f:
            f.write('\n'.join(rtf_lines))


def convert_document(input_path, output_path, target_format=None, ocr_enabled=False, ocr_model="qwen-vl"):
    """
    Convert a document or image to another format.
    
    Args:
        input_path (str): Path to input file
        output_path (str): Path to output file
        target_format (str, optional): Target format extension
        ocr_enabled (bool): Whether to use OCR for images/scanned PDFs
        ocr_model (str): OCR model to use ('florence', 'qwen-vl')
        
    Returns:
        bool: True if successful
    """
    # Determine output format
    if target_format:
        output_format = target_format.strip().lower()
    elif output_path and ('/' not in output_path and '\\' not in output_path and len(output_path) <= 6):
        # output_path is just a format string like "md" or "pdf"
        output_format = output_path.strip().lstrip('.').lower()
    else:
        output_format = Path(output_path).suffix.lstrip('.').lower()
    
    if output_format not in SUPPORTED_FORMATS:
        print(f"❌ Unsupported output format: {output_format}")
        print(f"   Supported: {', '.join(SUPPORTED_FORMATS)}")
        return False
    
    # Determine input format
    input_format = Path(input_path).suffix.lstrip('.').lower()
    
    # Allow image formats if OCR is enabled
    image_exts = ['jpg', 'jpeg', 'png', 'webp', 'bmp', 'tiff', 'gif']
    if input_format in image_exts:
        ocr_enabled = True # Use OCR for images
    
    if input_format in image_exts and ocr_enabled:
        pass # Valid for OCR
    elif input_format not in SUPPORTED_FORMATS:
        print(f"❌ Unsupported input format: {input_format}")
        print(f"   Supported: {', '.join(SUPPORTED_FORMATS)}")
        return False
    
    print(f"📄 Converting Document: {input_path}")
    print(f"   {input_format.upper()} → {output_format.upper()}")
    
    # Generate output path if only format string was provided
    if output_path and ('/' not in output_path and '\\' not in output_path and len(output_path) <= 6):
        output_path = f"{Path(input_path).stem}.{output_format}"
    
    should_write, output_path, _, _ = check_overwrite(output_path, always_overwrite=os.environ.get("AI_MEDIA_FORCE") == "1")
    if not should_write:
        return False
    
    try:
        # Step 1: Read to Markdown
        # Check if input is an image and OCR is enabled
        image_exts = ['jpg', 'jpeg', 'png', 'webp', 'bmp', 'tiff', 'gif']
        if input_format in image_exts and ocr_enabled:
            print(f"   📷 Image input detected. Using OCR ({ocr_model})...")
            from .ocr import image_to_text
            from PIL import Image
            import tempfile
            
            # Temporarily disable PIL's decompression bomb limit for large images
            # (we'll resize them down anyway)
            original_max = Image.MAX_IMAGE_PIXELS
            Image.MAX_IMAGE_PIXELS = None
            
            try:
                # Load and resize if needed to prevent GPU OOM
                img = Image.open(input_path)
                max_dim = 2048
                needs_resize = img.width > max_dim or img.height > max_dim
                
                if needs_resize:
                    ratio = min(max_dim / img.width, max_dim / img.height)
                    new_size = (int(img.width * ratio), int(img.height * ratio))
                    img = img.resize(new_size, Image.Resampling.LANCZOS)
                    print(f"      (Resized to {new_size[0]}x{new_size[1]} for OCR)")
                    
                    # Save resized image to temp file
                    if img.mode in ('RGBA', 'P'):
                        img = img.convert('RGB')
                    with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as tmp:
                        img.save(tmp.name, "JPEG", quality=90)
                        markdown_content = image_to_text(tmp.name, model_type=ocr_model)
                        os.unlink(tmp.name)
                else:
                    markdown_content = image_to_text(input_path, model_type=ocr_model)
            finally:
                Image.MAX_IMAGE_PIXELS = original_max
        else:
            markdown_content = _read_to_markdown(input_path, input_format, ocr_model=ocr_model)
        
        # Determine output format early for writing
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        # Step 2: Write to target format
        _write_from_markdown(markdown_content, output_path, output_format)
        
        print(f"✅ Saved to {output_path}")
        return True
        
    except Exception as e:
        # Re-raise the exception so the task runner gets the specific error message
        # The runner will handle logging and status update
        raise e
