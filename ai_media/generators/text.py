"""
Text/Article/Code generation module for AI-Media.

Supports: Article generation, code generation, chat sessions, and deep research.
Uses LLMs like Llama, Qwen, DeepSeek, and Mistral.
"""

import os
import re
import time
from datetime import datetime

from ..models import TEXT_MODELS, get_model_id
from ..utils.system import get_optimal_device_and_dtype


from ..utils.interaction import check_overwrite, prompt_choice


class ArticleGenerator:
    """Text generation for articles, code, research, and chat using LLMs."""
    
    def __init__(self, model_name="llama-3.1-8b", device=None, args=None):
        """Initialize the article generator.
        
        Args:
            model_name: Model short code or HF ID
            device: Torch device (auto-detected if None)
            args: Optional argparse namespace for flags like --force
        """
        import torch
        self.torch = torch
        
        self.model_name = get_model_id(model_name, TEXT_MODELS)
        self.device = device or get_optimal_device_and_dtype(quiet=True)[0]
        self.pipeline = None
        self.args = args
        
        # Import DDGS for web search
        # Try importing 'ddgs' first (new package name), then 'duckduckgo_search' (legacy)
        try:
            try:
                from ddgs import DDGS
            except ImportError:
                from duckduckgo_search import DDGS
            self.ddgs = DDGS()
        except ImportError:
            self.ddgs = None
            print("⚠️  duckduckgo_search/ddgs not installed. Online search unavailable.")
        
    def _load_model(self):
        """Load the LLM pipeline."""
        if self.pipeline:
            return
        
        from transformers import AutoTokenizer, pipeline
        
        print(f"📚 Loading Text Model: {self.model_name}...")
        try:
            # Use bfloat16 on CUDA if supported (Ampere+), otherwise float16
            from ..utils.system import is_bfloat16_supported
            if self.device.type == "cuda":
                dtype = self.torch.bfloat16 if is_bfloat16_supported() else self.torch.float16
            elif self.device.type == "mps":
                dtype = self.torch.float32  # MPS uses float32 for stability
            else:
                dtype = self.torch.float32
            
            # Workaround: Qwen3/Llama have numerical instability on MPS float16
            if self.device.type == "mps" and any(m in self.model_name.lower() for m in ["qwen3", "llama"]):
                print(f"   ⚠️  {self.model_name} detected on MPS - using fp32 for stability...")
                dtype = self.torch.float32
                os.environ["PYTORCH_ENABLE_MPS_FALLBACK"] = "1"
            
            tokenizer = AutoTokenizer.from_pretrained(self.model_name)
            
            # Memory optimization: 4-bit on CUDA
            quantization_config = None
            if self.device.type == "cuda":
                try:
                    from transformers import BitsAndBytesConfig
                    quantization_config = BitsAndBytesConfig(load_in_4bit=True)
                except ImportError:
                    pass
            
            model_kwargs = {"dtype": dtype}
            if quantization_config:
                model_kwargs["quantization_config"] = quantization_config
                model_kwargs["device_map"] = "auto"
            elif self.device.type == 'cuda':
                # Use 'auto' to enable offloading to CPU/RAM if the model is too large for VRAM
                model_kwargs["device_map"] = "auto"
            else:
                model_kwargs["device_map"] = self.device
                
            self.pipeline = pipeline(
                "text-generation",
                model=self.model_name,
                tokenizer=tokenizer,
                **model_kwargs
            )
            dtype_name = str(dtype).replace("torch.", "")
            print(f"   Platform: {self.device.type.upper()} | Dtype: {dtype_name}")
            print("✅ Model loaded.")
            
        except RuntimeError as e:
            error_msg = str(e)
            if "Invalid buffer size" in error_msg or "out of memory" in error_msg.lower():
                size_match = re.search(r'(\d+\.?\d*)\s*(GiB|GB|MiB|MB)', error_msg)
                
                print(f"\n❌ Model too large for this system.")
                if size_match:
                     print(f"   Allocation failed when trying to reserve {size_match.group(0)}.")
                     
                print(f"   The model '{self.model_name}' cannot fit in available memory.")
                print(f"   💡 Try a smaller model like:")
                print(f"      - deepseek-r1-qwen-7b (~7GB)")
                print(f"      - deepseek-r1-llama-8b (~8GB)")
                print(f"      - llama-3.1-8b (~16GB)")
                return
            else:
                print(f"❌ Failed to load model: {e}")
                raise
                
        except OSError as e:
            error_msg = str(e)
            if "not a valid model identifier" in error_msg or "Repository Not Found" in error_msg:
                print(f"\n❌ Model not found: '{self.model_name}'")
                print(f"   This model doesn't exist on HuggingFace.")
                print(f"   💡 Available models:")
                print(f"      - deepseek-r1-qwen-7b, deepseek-r1-qwen-14b, deepseek-r1-qwen-32b")
                print(f"      - deepseek-r1-llama-8b, deepseek-r1-llama-70b")
                print(f"      - llama-3.1-8b, qwen3-8b, qwen-2.5-14b, mistral-nemo-12b")
                return
            else:
                print(f"❌ Failed to load model: {e}")
                raise
                
        except (ValueError, Exception) as e:
            error_msg = str(e)
            if "Invalid buffer size" in error_msg or "out of memory" in error_msg.lower():
                size_match = re.search(r'(\d+\.?\d*)\s*(GiB|GB|MiB|MB)', error_msg)
                
                print(f"\n❌ Model too large for this system.")
                if size_match:
                     print(f"   Allocation failed when trying to reserve {size_match.group(0)}.")
                     
                print(f"   The model '{self.model_name}' cannot fit in available memory.")
                print(f"   💡 Try a smaller model like:")
                print(f"      - deepseek-r1-qwen-7b (~7GB)")
                print(f"      - deepseek-r1-llama-8b (~8GB)")
                print(f"      - llama-3.1-8b (~16GB)")
                return
            else:
                print(f"❌ Failed to load model: {e}")
                raise

    def deep_research(self, query, iterations=3):
        """Perform recursive web search and summarization."""
        if not self.ddgs:
            print("❌ Online search unavailable (duckduckgo_search not installed)")
            return ""
            
        print(f"\n🔎 Deep Researching: '{query}' ({iterations} iterations)...")
        results = []
        
        # 1. Initial Broad Search
        try:
            search_results = list(self.ddgs.text(query, max_results=iterations))
            pad_width = len(str(iterations))
            for i, res in enumerate(search_results, 1):
                num_str = str(i).zfill(pad_width)
                print(f"   Reading [{num_str}]: {res['title']}...")
                content = res.get('body', '') or res.get('snippet', '')
                
                # Attempt deep scraping for better context
                try:
                    import requests
                    from bs4 import BeautifulSoup
                    
                    page = requests.get(res['href'], timeout=5, headers={"User-Agent": "Mozilla/5.0"})
                    if page.status_code == 200:
                        soup = BeautifulSoup(page.text, 'html.parser')
                        paragraphs = [p.get_text().strip() for p in soup.find_all('p')]
                        full_text = ' '.join(p for p in paragraphs if p)
                        
                        if len(full_text) > 200:
                            content = full_text[:4000] + "..."
                except Exception:
                    pass
                
                results.append(f"Source: {res['title']}\nURL: {res['href']}\nContent: {content}\n")
                time.sleep(1.0)
        except Exception as e:
            print(f"⚠️ Search error: {e}")
        
        # 2. Image Search
        image_results = []
        try:
            image_query = f"{query} photos"
            print(f"   🖼️  Searching for images...")
            image_search = list(self.ddgs.images(image_query, max_results=5))
            
            for img in image_search:
                img_url = img.get('image', '')
                img_title = img.get('title', 'Image')
                if img_url and img_url.startswith('http'):
                    image_results.append(f"![{img_title}]({img_url})")
            
            if image_results:
                print(f"   ✅ Found {len(image_results)} images")
        except Exception as e:
            print(f"   ⚠️  Image search failed: {e}")
            
        research_context = "\n\n".join(results)
        
        if image_results:
            research_context += "\n\n## Available Images (use these exact URLs)\n"
            research_context += "\n".join(image_results)
        
        return research_context

    def generate_article(self, topic, output_file, format="md", online=False, 
                        research_iter=3, length="quick"):
        """Generate full article with optional research.
        
        Args:
            topic: Article topic/title
            output_file: Path to save output
            format: Output format (md, html, pdf, docx, rtf, txt, json)
            online: Enable deep research mode
            research_iter: Number of search iterations
            length: 'quick' (512 tokens), 'standard' (2048), 'detailed' (4096)
        """
        from rich.console import Console
        console = Console()
        
        length_config = {
            "quick": {"tokens": 512, "desc": "concise"},
            "standard": {"tokens": 2048, "desc": "balanced"},
            "detailed": {"tokens": 4096, "desc": "comprehensive"},
        }
        config = length_config.get(length, length_config["detailed"])
        max_tokens = config["tokens"]
        style = config["desc"]
        
        research_data = ""
        if online:
            with console.status(f"[bold green]Thinking... (Deep Research Iterations {research_iter})[/bold green]", spinner="dots"):
                research_data = self.deep_research(topic, iterations=research_iter)
        
        self._load_model()
        if not self.pipeline:
            return
        
        print(f"✍️  Writing {style} article on '{topic}'...")
        
        # Prompt Engineering
        if research_data:
            system_prompt = (
                f"You are an expert investigative journalist. Write a {style}, well-structured "
                "article based on the following research context. Use Markdown formatting. "
                "Cite sources where appropriate."
            )
            user_prompt = f"Topic: {topic}\n\nResearch Context:\n{research_data}\n\nArticle:"
        else:
            system_prompt = (
                f"You are a creative writer and expert knowledge base. Write a {style}, "
                "well-structured article on the following topic. Use Markdown formatting."
            )
            user_prompt = f"Topic: {topic}\n\nArticle:"

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
        
        full_prompt = self.pipeline.tokenizer.apply_chat_template(
            messages, tokenize=False, add_generation_prompt=True
        )
        
        with console.status("[bold green]Thinking... (Writing Article)[/bold green]", spinner="dots"):
            outputs = self.pipeline(
                full_prompt, 
                max_new_tokens=max_tokens, 
                do_sample=True, 
                temperature=0.7
            )
        
        final_md = outputs[0]['generated_text'][len(full_prompt):].strip()
        
        # Extract and save <think> blocks separately
        think_matches = re.findall(r'<think>(.*?)</think>', final_md, re.DOTALL)
        if think_matches:
            final_md = re.sub(r'<think>.*?</think>\s*', '', final_md, flags=re.DOTALL).strip()
            base, ext = os.path.splitext(output_file)
            think_file = f"{base}-think.md"
            try:
                with open(think_file, "w", encoding="utf-8") as f:
                    f.write("# Reasoning Process\n\n")
                    for i, block in enumerate(think_matches, 1):
                        if len(think_matches) > 1:
                            f.write(f"## Block {i}\n\n")
                        f.write(block.strip() + "\n\n")
                print(f"💭 Reasoning saved to: {think_file}")
            except Exception as e:
                print(f"⚠️  Could not save reasoning: {e}")
        
        # Save in requested format
        failed_images = self._save_formatted(final_md, output_file, format, online=online)
        
        # Offer retry if offline and images failed
        if not online and failed_images > 0:
            print(f"\n⚠️  {failed_images} image(s) could not be fetched (hallucinated URLs).")
            print("💡 Tip: Offline models (-ga) cannot provide real image URLs.")
            print("   Options:")
            print("   • Use Deep Research (-gr) to find real images from the web")
            print("   • Remove 'images' from your prompt for text-only articles")
            
            retry = prompt_choice("What would you like to do?", [
                ("Retry with Deep Research (online)", "y"),
                ("Keep current output (no images)", "n")
            ])
            
            if retry == "y":
                print("\n🔄 Retrying with Deep Research...")
                self.generate_article(
                    topic=topic,
                    output_file=output_file,
                    format=format,
                    online=True,
                    research_iter=research_iter,
                    length=length
                )

    def generate_code(self, prompt, output_file=None):
        """Generate Code from Prompt (supports multi-file output)."""
        self._load_model()
        if not self.pipeline:
            return
        
        from rich.console import Console
        console = Console()
        console.print(f"💻 Generating Code for: '{prompt}'...")
        
        system_prompt = (
            "You are an expert coding assistant. Write clean, efficient, and well-commented code "
            "based on the user's request. Return ONLY the code blocks. "
            "IMPORTANT: Before EACH code file, include a comment line with the filename, "
            "e.g., '# filename: my_script.py' or '// filename: src/utils.js'. "
            "You can use folder paths like 'src/module/file.py'. "
            "If multiple files are needed, separate them with filename comments. "
            "Do not include markdown backticks or explanations unless asked. "
            "Make sure the filename extension matches the code language."
        )
        user_prompt = f"Request: {prompt}\n\nCode:"
        
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
        
        full_prompt = self.pipeline.tokenizer.apply_chat_template(
            messages, tokenize=False, add_generation_prompt=True
        )
        
        console.print()
        with console.status("[yellow] Thinking...[/yellow]", spinner="dots"):
            outputs = self.pipeline(
                full_prompt, 
                max_new_tokens=4096, 
                do_sample=True, 
                temperature=0.2,
                top_p=0.9,
            )
        
        generated_text = outputs[0]['generated_text']
        response = generated_text[len(full_prompt):].strip()
        
        # Remove markdown code fences if present
        response = re.sub(r"```\w*\n?", "", response)
        
        # Parse multiple files from response
        file_pattern = re.compile(
            r"^(?:#|//)\s*(?:filename:\s*)?([^\s]+\.(?:py|js|ts|jsx|tsx|html|css|java|cpp|c|h|go|rs|rb|php|sh|sql|json|yaml|yml|md|txt))\s*$",
            re.IGNORECASE | re.MULTILINE
        )
        
        parts = file_pattern.split(response)
        files_to_write = []
        
        if len(parts) > 1:
            for i in range(1, len(parts), 2):
                if i + 1 < len(parts):
                    filename = parts[i].strip()
                    content = parts[i + 1].strip()
                    if filename and content:
                        files_to_write.append((filename, content))
        else:
            content = response.strip()
            if output_file:
                if "." in os.path.basename(output_file):
                    files_to_write.append((output_file, content))
                else:
                    ext = self._infer_extension(content)
                    files_to_write.append((f"{output_file}{ext}", content))
            else:
                ext = self._infer_extension(content)
                files_to_write.append((f"generated_code_{int(time.time())}{ext}", content))
        
        # Write all files
        output_is_dir = output_file and os.path.isdir(output_file)
        always_overwrite = self.args.force if self.args and hasattr(self.args, 'force') else False
        never_overwrite = False
        
        for filepath, content in files_to_write:
            try:
                final_path = filepath
                
                if output_file:
                    if output_is_dir:
                        final_path = os.path.join(output_file, filepath)
                    elif len(files_to_write) == 1:
                        final_path = output_file

                should_write, final_path, always_overwrite, never_overwrite = check_overwrite(
                    final_path, always_overwrite, never_overwrite
                )
                
                if final_path is None:
                    print("🛑 Code generation cancelled.")
                    break
                
                if not should_write:
                    continue

                dir_path = os.path.dirname(final_path)
                if dir_path:
                    os.makedirs(dir_path, exist_ok=True)
                
                with open(final_path, "w", encoding="utf-8") as f:
                    f.write(content)
                print(f"✅ Code saved to: {final_path}")
            except Exception as e:
                print(f"❌ Error saving {filepath}: {e}")
                
    def _infer_extension(self, code_content):
        """Infer file extension from code content."""
        if "import " in code_content or "def " in code_content or "print(" in code_content:
            return ".py"
        elif "function " in code_content or "const " in code_content or "console.log" in code_content:
            return ".js"
        elif "#include" in code_content:
            return ".cpp"
        elif "public class" in code_content:
            return ".java"
        elif "<html" in code_content:
            return ".html"
        elif "package main" in code_content:
            return ".go"
        elif "fn main" in code_content or "use std::" in code_content:
            return ".rs"
        return ".txt"

    def _save_formatted(self, markdown_text, filename, fmt, online=False):
        """Convert and save to specific format.
        
        Returns:
            int: Number of failed image fetches
        """
        import markdown as md_module
        
        failed_image_count = 0
        base, _ = os.path.splitext(filename)
        
        if not filename.lower().endswith(f".{fmt}"):
            filename = f"{base}.{fmt}"
            
        print(f"💾 Saving as {fmt.upper()}...")
        
        if fmt == "md":
            with open(filename, "w", encoding="utf-8") as f:
                f.write(markdown_text)
                
        elif fmt == "html" or fmt == "xhtml":
            html = md_module.markdown(markdown_text, extensions=['extra', 'codehilite'])
            full_html = (
                f"<!DOCTYPE html><html><head><meta charset='utf-8'><title>Article</title>"
                f"<style>body{{font-family:sans-serif;max-width:800px;margin:2em auto;padding:1em;line-height:1.6}}"
                f"pre{{background:#f4f4f4;padding:1em;border-radius:5px}}</style></head>"
                f"<body>{html}</body></html>"
            )
            with open(filename, "w", encoding="utf-8") as f:
                f.write(full_html)
                
        elif fmt == "json":
            import json
            data = {"content": markdown_text, "html": md_module.markdown(markdown_text)}
            with open(filename, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2)

        elif fmt == "txt":
            from bs4 import BeautifulSoup
            html = md_module.markdown(markdown_text)
            text = BeautifulSoup(html, "html.parser").get_text()
            with open(filename, "w", encoding="utf-8") as f:
                f.write(text)
                
        elif fmt == "docx":
            import io
            import urllib.request
            import re as re_module
            import docx
            from docx.shared import Inches
            
            doc = docx.Document()
            MIN_IMAGE_SIZE = 5 * 1024  # 5KB threshold
            
            def fetch_image_for_docx(url):
                """Fetch image and return as BytesIO for docx embedding."""
                nonlocal failed_image_count
                try:
                    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
                    with urllib.request.urlopen(req, timeout=10) as response:
                        image_data = response.read()
                    
                    if len(image_data) < MIN_IMAGE_SIZE:
                        failed_image_count += 1
                        print(f"⚠️  Image too small (likely placeholder): {url[:50]}...")
                        return None
                    
                    return io.BytesIO(image_data)
                except Exception as e:
                    failed_image_count += 1
                    print(f"⚠️  Could not fetch image: {url[:50]}... ({e})")
                    return None
            
            # Process markdown line by line
            for line in markdown_text.split('\n'):
                # Check for markdown image: ![alt](url)
                img_match = re_module.match(r'!\[([^\]]*)\]\((https?://[^\)]+)\)', line)
                if img_match:
                    alt_text = img_match.group(1)
                    img_url = img_match.group(2)
                    img_stream = fetch_image_for_docx(img_url)
                    if img_stream:
                        try:
                            doc.add_picture(img_stream, width=Inches(5))
                            if alt_text:
                                caption = doc.add_paragraph(alt_text)
                                caption.alignment = 1  # Center
                        except Exception as e:
                            print(f"⚠️  Could not embed image: {e}")
                            doc.add_paragraph(f"[Image: {alt_text}]")
                    else:
                        doc.add_paragraph(f"[Image: {alt_text}]")
                elif line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                else:
                    doc.add_paragraph(line)
            doc.save(filename)
        
        elif fmt == "rtf":
            import urllib.request
            import re as re_module
            import binascii
            
            MIN_IMAGE_SIZE = 5 * 1024
            
            def rtf_escape(text):
                return text.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
            
            def fetch_image_for_rtf(url):
                nonlocal failed_image_count
                try:
                    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
                    with urllib.request.urlopen(req, timeout=10) as response:
                        image_data = response.read()
                    
                    if len(image_data) < MIN_IMAGE_SIZE:
                        failed_image_count += 1
                        return None
                    
                    if image_data[:3] == b'\xff\xd8\xff':
                        img_format = 'jpegblip'
                    elif image_data[:8] == b'\x89PNG\r\n\x1a\n':
                        img_format = 'pngblip'
                    else:
                        img_format = 'jpegblip'
                    
                    hex_data = binascii.hexlify(image_data).decode('ascii')
                    return (img_format, hex_data)
                except Exception as e:
                    failed_image_count += 1
                    return None
            
            rtf_lines = []
            rtf_lines.append(r'{\rtf1\ansi\deff0')
            rtf_lines.append(r'{\fonttbl{\f0 Helvetica;}{\f1 Courier;}}')
            rtf_lines.append(r'{\colortbl;\red0\green0\blue0;\red51\green51\blue51;}')
            rtf_lines.append(r'\f0\fs24')
            
            for line in markdown_text.split('\n'):
                img_match = re_module.match(r'!\[([^\]]*)\]\((https?://[^\)]+)\)', line)
                if img_match:
                    alt_text = img_match.group(1)
                    img_url = img_match.group(2)
                    img_data = fetch_image_for_rtf(img_url)
                    if img_data:
                        img_format, hex_data = img_data
                        rtf_lines.append(r'\pard\qc\sb200\sa100')
                        rtf_lines.append(r'{\pict\\' + img_format + r'\picwgoal6000\pichgoal4000')
                        rtf_lines.append(hex_data)
                        rtf_lines.append(r'}')
                        if alt_text:
                            rtf_lines.append(r'\pard\qc\i\fs20 ' + rtf_escape(alt_text) + r'\i0\fs24\par')
                    else:
                        rtf_lines.append(r'\pard\sa100 [Image: ' + rtf_escape(alt_text) + r']\par')
                    continue
                
                line = rtf_escape(line)
                if line.startswith('# '):
                    rtf_lines.append(r'\pard\sb400\sa200\b\fs48 ' + line[2:] + r'\b0\fs24\par')
                elif line.startswith('## '):
                    rtf_lines.append(r'\pard\sb300\sa150\b\fs36 ' + line[3:] + r'\b0\fs24\par')
                elif line.startswith('### '):
                    rtf_lines.append(r'\pard\sb200\sa100\b\fs28 ' + line[4:] + r'\b0\fs24\par')
                elif line.startswith('- ') or line.startswith('* '):
                    rtf_lines.append(r'\pard\li720\fi-360\bullet  ' + line[2:] + r'\par')
                elif line.startswith('```'):
                    continue
                elif line.strip():
                    rtf_lines.append(r'\pard\sa100 ' + line + r'\par')
                else:
                    rtf_lines.append(r'\par')
            
            rtf_lines.append('}')
            
            with open(filename, "w", encoding="utf-8") as f:
                f.write('\n'.join(rtf_lines))
                
        elif fmt == "pdf":
            import base64
            import urllib.request
            import re as re_module
            from xhtml2pdf import pisa
            
            # Pre-process markdown
            processed_md = re_module.sub(
                r'\[([Ii]mage[^\]]*)\]\((https?://[^\)]+\.(jpg|jpeg|png|gif|webp)[^\)]*)\)',
                r'![\1](\2)',
                markdown_text
            )
            
            # Remove bolding from table rows
            md_lines = processed_md.split('\n')
            for idx, line in enumerate(md_lines):
                if line.strip().startswith('|'):
                    md_lines[idx] = line.replace('**', '')
            processed_md = '\n'.join(md_lines)
            
            # Convert MD -> HTML
            html_content = md_module.markdown(processed_md, extensions=['extra', 'fenced_code', 'tables', 'toc'])
            
            # Strip emojis
            def strip_emojis(text):
                result = []
                for char in text:
                    code = ord(char)
                    if (0x1F300 <= code <= 0x1FFFF or 0x2600 <= code <= 0x27BF or
                        0x2300 <= code <= 0x23FF or 0xFE00 <= code <= 0xFE0F):
                        continue
                    result.append(char)
                return ''.join(result)
            
            html_content = strip_emojis(html_content)
            
            # Fetch and embed images
            def fetch_and_encode_image(url):
                nonlocal failed_image_count
                MIN_IMAGE_SIZE = 5 * 1024
                try:
                    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
                    with urllib.request.urlopen(req, timeout=10) as response:
                        image_data = response.read()
                    
                    if len(image_data) < MIN_IMAGE_SIZE:
                        failed_image_count += 1
                        return None
                    
                    content_type = 'image/jpeg'
                    if 'png' in url.lower():
                        content_type = 'image/png'
                    elif 'gif' in url.lower():
                        content_type = 'image/gif'
                    
                    b64_data = base64.b64encode(image_data).decode('utf-8')
                    return f'data:{content_type};base64,{b64_data}'
                except Exception as e:
                    failed_image_count += 1
                    return None
            
            def replace_src(match):
                url = match.group(1)
                data_uri = fetch_and_encode_image(url)
                if data_uri:
                    return f'src="{data_uri}"'
                return match.group(0)
            
            html_content = re_module.sub(r'src="(https?://[^"]+)"', replace_src, html_content)
            
            # Full HTML document for xhtml2pdf
            full_html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
@page {{ size: a4 portrait; margin: 1.5cm; }}
body {{ font-family: Helvetica, sans-serif; font-size: 9pt; line-height: 1.4; }}
h1 {{ font-size: 18pt; color: #333; margin-top: 0.8em; margin-bottom: 0.4em; }}
h2 {{ font-size: 14pt; color: #444; margin-top: 0.6em; margin-bottom: 0.3em; }}
h3 {{ font-size: 12pt; color: #555; margin-top: 0.5em; margin-bottom: 0.2em; }}
pre {{ background: #f4f4f4; padding: 6px; font-family: Courier, monospace; font-size: 6pt; }}
code {{ background: #f0f0f0; padding: 1px 2px; font-family: Courier, monospace; font-size: 6pt; }}
table {{ border-collapse: collapse; width: 100%; font-size: 6pt; margin: 0.4em 0; }}
th, td {{ border: 1px solid #999; padding: 2px 4px; text-align: left; }}
th {{ background: #e8e8e8; font-weight: bold; }}
img {{ max-width: 100%; height: auto; }}
a {{ color: #0066cc; text-decoration: underline; }}
</style>
</head>
<body>
{html_content}
</body>
</html>"""
            
            with open(filename, "wb") as f:
                pisa_status = pisa.CreatePDF(full_html, dest=f)
            
            if pisa_status.err:
                print("❌ PDF conversion failed")
                return failed_image_count
                
        else:
            print(f"⚠️ Unknown format '{fmt}', saving as MD.")
            with open(f"{base}.md", "w", encoding="utf-8") as f:
                f.write(markdown_text)

        print(f"✅ Saved to {filename}")
        return failed_image_count

    def chat_session(self):
        """Interactive Chat Loop."""
        from rich.console import Console
        from rich.markdown import Markdown
        from prompt_toolkit import PromptSession, HTML
        from prompt_toolkit.history import InMemoryHistory
        from prompt_toolkit.lexers import Lexer
        from prompt_toolkit.styles import Style
        from prompt_toolkit.completion import NestedCompleter, PathCompleter, FuzzyCompleter
        
        console = Console()
        
        self._load_model()
        if not self.pipeline:
            return
            
        history = []
        pending_context = ""
        
        # Detect location once at session start
        user_location = "Unknown"
        try:
            import json
            from urllib.request import urlopen
            # Short timeout to avoid hanging
            with urlopen("http://ip-api.com/json/", timeout=1.5) as response:
                ip_data = json.loads(response.read().decode())
                if ip_data.get('status') == 'success':
                    user_location = f"{ip_data.get('city')}, {ip_data.get('regionName')}, {ip_data.get('country')}"
        except:
            pass
        
        class ChatLexer(Lexer):
            def lex_document(self, document):
                def get_line_tokens(line_number):
                    line = document.lines[line_number]
                    for cmd in ['/read', '/save', '/search', '/online-search']:
                        if line.startswith(cmd):
                            base_len = len(cmd)
                            if len(line) > base_len and line[base_len] == '|':
                                end_pos = line.find(' ', base_len)
                                if end_pos == -1:
                                    end_pos = len(line)
                                return [
                                    ('class:command', line[:end_pos]),
                                    ('', line[end_pos:])
                                ]
                            return [
                                ('class:command', cmd),
                                ('', line[base_len:])
                            ]
                    return [('', line)]
                return get_line_tokens

        chat_style = Style.from_dict({
            'command': '#ff00ff bold',
        })

        session = PromptSession(
            history=InMemoryHistory(),
            lexer=ChatLexer(),
            style=chat_style
        )
        
        path_completer = FuzzyCompleter(PathCompleter(expanduser=True))
        completer = NestedCompleter.from_nested_dict({
            '/read': path_completer,
            '/save': path_completer,
            '/search': None,
            '/online-search': None,
            'exit': None,
            'quit': None,
        })
        
        console.print(f"\n💬 [bold]Chat Session Started[/bold] (Model: [bold cyan]{self.model_name}[/bold cyan])")
        console.print("   Type '[bold]exit[/bold]' or '[bold]quit[/bold]' to end.")
        console.print("   Commands: [bold]/read <path>[/bold], [bold]/save[/bold][bold]|all[/bold] [bold]<path>[/bold], [bold]/search[/bold][bold]|N[/bold] [bold]<query>[/bold]")
        console.print("   [dim]💡 Tip: /save saves last code or full response. Use |all for full history.[/dim]")
        console.print("   [dim]💡 Tip: Use /search query or /search|5 query for deeper results.[/dim]\n")
        
        while True:
            try:
                user_input = session.prompt(HTML('<b fg="blue">You:</b> '), completer=completer, complete_while_typing=True)
                if user_input.strip().lower() in ['exit', 'quit']:
                    break
                
                # Handle slash commands
                if user_input.startswith("/read "):
                    file_path = user_input[6:].strip()
                    if os.path.exists(file_path):
                        try:
                            with open(file_path, "r", encoding="utf-8") as f:
                                content = f.read()
                            pending_context += f"\n\n[File Context: {file_path}]\n{content}\n"
                            console.print(f"📄 [bold green]Added file context:[/bold green] {file_path}")
                        except Exception as e:
                            console.print(f"[bold red]❌ Error reading file:[/bold red] {e}")
                    else:
                        console.print(f"[bold red]❌ File not found:[/bold red] {file_path}")
                    continue

                is_search = False
                for s_cmd in ["/search", "/online-search"]:
                    if user_input.startswith(s_cmd + " ") or user_input.startswith(s_cmd + "|"):
                        is_search = True
                        break

                if is_search:
                    parts = user_input.split(' ', 1)
                    cmd_part = parts[0]
                    query = parts[1].strip() if len(parts) > 1 else ""
                    
                    iterations = 3
                    if '|' in cmd_part:
                        try:
                            iter_str = cmd_part.split('|', 1)[1]
                            if iter_str.isdigit():
                                iterations = int(iter_str)
                        except:
                            pass
                    
                    if query:
                        try:
                            search_results = self.deep_research(query, iterations=iterations)
                            if search_results:
                                pending_context += f"\n\n[Online Search Context: '{query}']\n{search_results}\n"
                                console.print(f"🌍 [bold green]Added search results for:[/bold green] '{query}'")
                            else:
                                console.print(f"[bold yellow]⚠️ No results found for:[/bold yellow] '{query}'")
                        except Exception as e:
                            console.print(f"[bold red]❌ Search error:[/bold red] {e}")
                    else:
                        console.print("[bold red]❌ Please provide a search query.[/bold red]")
                    continue

                if user_input.startswith("/save"):
                    # Handle /save|all syntax or regular /save
                    parts = user_input.split(' ', 1)
                    cmd_part = parts[0]
                    file_path = parts[1].strip() if len(parts) > 1 else ""
                    
                    save_all = "|all" in cmd_part.lower()
                    content_to_save = None
                    label = "response"
                    ext_suggestion = ".md"

                    if save_all:
                        # Format full history as markdown
                        history_content = "# Chat Conversation History\n\n"
                        for msg in history:
                            role = "User" if msg["role"] == "user" else "Assistant"
                            history_content += f"## {role}\n{msg['content']}\n\n"
                        content_to_save = history_content
                        label = "full conversation history"
                        ext_suggestion = ".md"
                    else:
                        # 1. Try to find last code block
                        for msg in reversed(history):
                            if msg["role"] == "assistant":
                                content = msg["content"]
                                # Attempt to find code blocks and language
                                matches = re.findall(r"```(.*?)\n(.*?)```", content, re.DOTALL)
                                if matches:
                                    lang, code = matches[-1]
                                    content_to_save = code
                                    label = "code block"
                                    # Suggest extension based on language
                                    lang_map = {"python": ".py", "bash": ".sh", "javascript": ".js", "html": ".html", "css": ".css", "markdown": ".md"}
                                    ext_suggestion = lang_map.get(lang.strip().lower(), ".txt")
                                    
                                    # Heuristic: look for filename in the text before the block
                                    if not file_path:
                                        fn_match = re.search(r"(\w+[\.\w]+)", content[:content.find("```")].split("\n")[-1])
                                        if fn_match and "." in fn_match.group(1):
                                            suggested_fn = fn_match.group(1)
                                            if os.path.splitext(suggested_fn)[1] in lang_map.values():
                                                file_path = suggested_fn
                                    break
                                else:
                                    # 2. Fallback to full last response
                                    content_to_save = content
                                    label = "last response"
                                    ext_suggestion = ".md"
                                    break

                    if not file_path:
                        # Try to get a descriptive name from context
                        context_str = ""
                        if save_all and history:
                            for msg in history:
                                if msg["role"] == "user" and not msg["content"].strip().startswith("/"):
                                    context_str = msg["content"]
                                    break
                        elif history:
                            for msg in reversed(history):
                                if msg["role"] == "user" and not msg["content"].strip().startswith("/"):
                                    context_str = msg["content"]
                                    break
                        
                        # Slugify: lowercase, alphanumeric and underscores only
                        clean_text = re.sub(r'[^a-zA-Z0-9\s]', '', context_str).strip()
                        slug = re.sub(r'\s+', '_', clean_text[:25]).lower()
                        
                        ts = int(time.time())
                        if save_all:
                            prefix = f"chat_{slug}" if slug else "chat_all"
                        else:
                            type_prefix = "code" if label == "code block" else "resp"
                            prefix = f"{type_prefix}_{slug}" if slug else type_prefix
                            
                        file_path = f"{prefix}_{ts}{ext_suggestion}"
                    elif "." not in os.path.basename(file_path):
                        # Add suggested extension if missing
                        file_path += ext_suggestion

                    if content_to_save:
                        # Check overwrite before saving
                        always_overwrite = self.args.force if self.args and hasattr(self.args, 'force') else False
                        should_write, final_path, _, _ = check_overwrite(file_path, always_overwrite=always_overwrite)
                        if should_write:
                            try:
                                with open(final_path, "w", encoding="utf-8") as f:
                                    f.write(content_to_save)
                                console.print(f"💾 [bold green]Exported {label} to:[/bold green] {final_path}")
                            except Exception as e:
                                console.print(f"[bold red]❌ Error saving file:[/bold red] {e}")
                        else:
                            console.print(f"\n[bold yellow]⏭️  Save cancelled (skipped).[/bold yellow]")
                    else:
                         console.print("[bold red]❌ No conversation content found to save.[/bold red]")
                    continue

                # Construct prompt
                final_content = user_input
                if pending_context:
                    final_content = pending_context + "\n" + user_input
                    pending_context = ""
                
                history.append({"role": "user", "content": final_content})
                
                if len(history) > 20:
                    history = history[-10:]
                
                # Build prompt with dynamic system context (Time/Location)
                current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                system_prompt = f"You are a helpful assistant. Current date and time: {current_time}. User location: {user_location}."
                
                prompt_messages = [{"role": "system", "content": system_prompt}] + history
                
                prompt = self.pipeline.tokenizer.apply_chat_template(
                    prompt_messages, tokenize=False, add_generation_prompt=True
                )
                
                with console.status("[yellow]Thinking...[/yellow]", spinner="dots"):
                    outputs = self.pipeline(
                        prompt, 
                        max_new_tokens=512, 
                        do_sample=True, 
                        temperature=0.7,
                        top_p=0.9,
                    )
                
                console.print("[bold green]Bot:[/bold green]")
                
                generated_text = outputs[0]['generated_text']
                response = generated_text[len(prompt):].strip()
                
                # Handle DeepSeek R1 reasoning
                if '</think>' in response:
                    parts = response.split('</think>', 1)
                    reasoning = parts[0].replace('<think>', '').strip()
                    final_answer = parts[1].strip() if len(parts) > 1 else ""
                    
                    console.print("[dim italic]💭 Reasoning:[/dim italic]")
                    console.print(f"[dim italic]{reasoning}[/dim italic]")
                    console.print("")  # Spacer
                    if final_answer:
                        console.print("[bold]Answer:[/bold]")
                        console.print(Markdown(final_answer))
                else:
                    console.print(Markdown(response))
                console.print("")
                
                # Keep original response in history to maintain reasoning context
                history.append({"role": "assistant", "content": response})
                
            except KeyboardInterrupt:
                console.print("\n")
                break
            except Exception as e:
                console.print(f"[bold red]❌ Error:[/bold red] {e}")
