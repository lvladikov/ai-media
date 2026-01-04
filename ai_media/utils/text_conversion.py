"""Utility for converting markdown text to various formats."""

import os
import io
import re
import json
import base64
import urllib.request
import binascii
from bs4 import BeautifulSoup
import markdown as md_module
from xhtml2pdf import pisa

def convert_text(markdown_text: str, fmt: str, filename: str = None) -> bytes:
    """Convert markdown text to the specified format and return as bytes.
    
    Args:
        markdown_text: The source markdown text.
        fmt: Target format (md, txt, html, json, pdf, rtf, docx).
        filename: Optional filename for internal use (e.g. title).
        
    Returns:
        bytes: The converted content.
    """
    fmt = fmt.lower()
    
    if fmt == "md":
        return markdown_text.encode("utf-8")
        
    elif fmt == "html":
        html = md_module.markdown(markdown_text, extensions=['extra', 'codehilite'])
        full_html = (
            f"<!DOCTYPE html><html><head><meta charset='utf-8'><title>Export</title>"
            f"<style>body{{font-family:sans-serif;max-width:800px;margin:2em auto;padding:1em;line-height:1.6}}"
            f"pre{{background:#f4f4f4;padding:1em;border-radius:5px}}</style></head>"
            f"<body>{html}</body></html>"
        )
        return full_html.encode("utf-8")

    elif fmt == "xhtml":
        # Strict XHTML 1.1
        html = md_module.markdown(markdown_text, extensions=['extra', 'codehilite'], output_format='xhtml')
        full_html = (
            f'<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.1//EN" "http://www.w3.org/TR/xhtml11/DTD/xhtml11.dtd">\n'
            f'<html xmlns="http://www.w3.org/1999/xhtml">\n'
            f'<head>\n'
            f'<meta http-equiv="Content-Type" content="text/html; charset=utf-8" />\n'
            f'<title>Export</title>\n'
            f'<style type="text/css">\n'
            f'body{{font-family:sans-serif;max-width:800px;margin:2em auto;padding:1em;line-height:1.6}}\n'
            f'pre{{background:#f4f4f4;padding:1em;border-radius:5px}}\n'
            f'</style>\n'
            f'</head>\n'
            f'<body>\n{html}\n</body>\n</html>'
        )
        return full_html.encode("utf-8")
        
    elif fmt == "json":
        data = {"content": markdown_text, "html": md_module.markdown(markdown_text)}
        return json.dumps(data, indent=2).encode("utf-8")

    elif fmt == "txt":
        html = md_module.markdown(markdown_text)
        text = BeautifulSoup(html, "html.parser").get_text()
        return text.encode("utf-8")
        
    elif fmt == "docx":
        import docx
        from docx.shared import Inches
        
        doc = docx.Document()
        MIN_IMAGE_SIZE = 5 * 1024
        
        def fetch_image(url):
            try:
                req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
                with urllib.request.urlopen(req, timeout=10) as response:
                    image_data = response.read()
                if len(image_data) < MIN_IMAGE_SIZE:
                    return None
                return io.BytesIO(image_data)
            except:
                return None
        
        for line in markdown_text.split('\n'):
            img_match = re.match(r'!\[([^\]]*)\]\((https?://[^\)]+)\)', line)
            if img_match:
                alt_text = img_match.group(1)
                img_url = img_match.group(2)
                img_stream = fetch_image(img_url)
                if img_stream:
                    try:
                        doc.add_picture(img_stream, width=Inches(5))
                        if alt_text:
                            caption = doc.add_paragraph(alt_text)
                            caption.alignment = 1
                    except:
                        doc.add_paragraph(f"[Image: {alt_text}]")
                else:
                    doc.add_paragraph(f"[Image: {alt_text}]")
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            else:
                doc.add_paragraph(line)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    
    elif fmt == "rtf":
        MIN_IMAGE_SIZE = 5 * 1024
        
        def rtf_escape(text):
            return text.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
        
        def fetch_image(url):
            try:
                req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
                with urllib.request.urlopen(req, timeout=10) as response:
                    image_data = response.read()
                if len(image_data) < MIN_IMAGE_SIZE:
                    return None
                
                if image_data[:3] == b'\xff\xd8\xff':
                    img_format = 'jpegblip'
                elif image_data[:8] == b'\x89PNG\r\n\x1a\n':
                    img_format = 'pngblip'
                else:
                    img_format = 'jpegblip'
                
                hex_data = binascii.hexlify(image_data).decode('ascii')
                return (img_format, hex_data)
            except:
                return None
        
        rtf_lines = [
            r'{\rtf1\ansi\deff0',
            r'{\fonttbl{\f0 Helvetica;}{\f1 Courier;}}',
            r'{\colortbl;\red0\green0\blue0;\red51\green51\blue51;}',
            r'\f0\fs24'
        ]
        
        for line in markdown_text.split('\n'):
            img_match = re.match(r'!\[([^\]]*)\]\((https?://[^\)]+)\)', line)
            if img_match:
                alt_text = img_match.group(1)
                img_url = img_match.group(2)
                img_data = fetch_image(img_url)
                if img_data:
                    img_format, hex_data = img_data
                    rtf_lines.append(r'\pard\qc\sb200\sa100')
                    rtf_lines.append(r'{\pict\\' + img_format + r'\picwgoal6000\pichgoal4000')
                    rtf_lines.append(hex_data)
                    rtf_lines.append(r'}')
                    if alt_text:
                        rtf_lines.append(r'\pard\qc\i\fs20 ' + rtf_escape(alt_text) + r'\i0\fs24\par')
                else:
                    rtf_lines.append(r'\pard\sa100 [Image: ' + rtf_escape(alt_text) + r']\par')
                continue
            
            line = rtf_escape(line)
            if line.startswith('# '):
                rtf_lines.append(r'\pard\sb400\sa200\b\fs48 ' + line[2:] + r'\b0\fs24\par')
            elif line.startswith('## '):
                rtf_lines.append(r'\pard\sb300\sa150\b\fs36 ' + line[3:] + r'\b0\fs24\par')
            elif line.startswith('### '):
                rtf_lines.append(r'\pard\sb200\sa100\b\fs28 ' + line[4:] + r'\b0\fs24\par')
            elif line.startswith('- ') or line.startswith('* '):
                rtf_lines.append(r'\pard\li720\fi-360\bullet  ' + line[2:] + r'\par')
            elif line.startswith('```'):
                continue
            elif line.strip():
                rtf_lines.append(r'\pard\sa100 ' + line + r'\par')
            else:
                rtf_lines.append(r'\par')
        
        rtf_lines.append('}')
        return '\n'.join(rtf_lines).encode("utf-8")
            
    elif fmt == "pdf":
        # Pre-process markdown
        processed_md = re.sub(
            r'\[([Ii]mage[^\]]*)\]\((https?://[^\)]+\.(jpg|jpeg|png|gif|webp)[^\)]*)\)',
            r'![\1](\2)',
            markdown_text
        )
        
        # Remove bolding from table rows
        md_lines = processed_md.split('\n')
        for idx, line in enumerate(md_lines):
            if line.strip().startswith('|'):
                md_lines[idx] = line.replace('**', '')
        processed_md = '\n'.join(md_lines)
        
        # Convert MD -> HTML
        html_content = md_module.markdown(processed_md, extensions=['extra', 'fenced_code', 'tables', 'toc'])
        
        def strip_emojis(text):
            result = []
            for char in text:
                code = ord(char)
                if (0x1F300 <= code <= 0x1FFFF or 0x2600 <= code <= 0x27BF or
                    0x2300 <= code <= 0x23FF or 0xFE00 <= code <= 0xFE0F):
                    continue
                result.append(char)
            return ''.join(result)
        
        html_content = strip_emojis(html_content)
        
        def fetch_and_encode_image(url):
            MIN_IMAGE_SIZE = 5 * 1024
            try:
                req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
                with urllib.request.urlopen(req, timeout=10) as response:
                    image_data = response.read()
                if len(image_data) < MIN_IMAGE_SIZE:
                    return None
                content_type = 'image/jpeg'
                if 'png' in url.lower():
                    content_type = 'image/png'
                elif 'gif' in url.lower():
                    content_type = 'image/gif'
                b64_data = base64.b64encode(image_data).decode('utf-8')
                return f'data:{content_type};base64,{b64_data}'
            except:
                return None
        
        def replace_src(match):
            url = match.group(1)
            data_uri = fetch_and_encode_image(url)
            if data_uri:
                return f'src="{data_uri}"'
            return match.group(0)
        
        html_content = re.sub(r'src="(https?://[^"]+)"', replace_src, html_content)
        
        full_html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
@page {{ size: a4 portrait; margin: 1.5cm; }}
body {{ font-family: Helvetica, sans-serif; font-size: 9pt; line-height: 1.4; }}
h1 {{ font-size: 18pt; color: #333; margin-top: 0.8em; margin-bottom: 0.4em; }}
h2 {{ font-size: 14pt; color: #444; margin-top: 0.6em; margin-bottom: 0.3em; }}
h3 {{ font-size: 12pt; color: #555; margin-top: 0.5em; margin-bottom: 0.2em; }}
pre {{ background: #f4f4f4; padding: 6px; font-family: Courier, monospace; font-size: 6pt; }}
code {{ background: #f0f0f0; padding: 1px 2px; font-family: Courier, monospace; font-size: 6pt; }}
table {{ border-collapse: collapse; width: 100%; font-size: 6pt; margin: 0.4em 0; }}
th, td {{ border: 1px solid #999; padding: 2px 4px; text-align: left; }}
th {{ background: #e8e8e8; font-weight: bold; }}
img {{ max-width: 100%; height: auto; }}
a {{ color: #0066cc; text-decoration: underline; }}
</style>
</head>
<body>
{html_content}
</body>
</html>"""
        
        buffer = io.BytesIO()
        pisa_status = pisa.CreatePDF(full_html, dest=buffer)
        if pisa_status.err:
            raise Exception("PDF conversion failed")
        return buffer.getvalue()
            
    else:
        # Fallback to MD
        return markdown_text.encode("utf-8")
