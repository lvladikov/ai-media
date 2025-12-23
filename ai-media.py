#!/usr/bin/env python3
"""
AI-Media: Local Offline Media Generation CLI

Generate Images, Videos, and Audio using local AI models (Flux, MusicGen, etc.).
Wraps 'diffusers' and 'transformers' libraries.

Usage:
  python ai-media.py -i -p "Prompt" -o image.png
  python ai-media.py -v -p "Prompt" -o video.mp4 -l 5s
  python ai-media.py -a -p "Prompt" -o audio.mp3 -l 30s
"""

import argparse
import warnings
import signal
import sys
import os
import time

# Suppress common library warnings
# (Some libraries use different warning categories or print directly)
warnings.filterwarnings("ignore", message="User provided device_type of 'cuda'", category=UserWarning)
warnings.filterwarnings("ignore", message=".*torch_dtype.*deprecated.*", category=FutureWarning)
warnings.filterwarnings("ignore", message=".*torch_dtype.*deprecated.*", category=DeprecationWarning)
warnings.filterwarnings("ignore", message=".*torch_dtype.*deprecated.*", category=UserWarning)
warnings.filterwarnings("ignore", message=".*torch_dtype.*deprecated.*")  # Catch all
warnings.filterwarnings("ignore", message=".*upcast_vae.*deprecated.*", category=FutureWarning)
import ast
import json
import logging
import os
import signal
import sys
import re
import argparse
import time
from datetime import datetime
import shutil
import subprocess
try:
    import psutil  # For resource checking
except ImportError:
    psutil = None
from pathlib import Path
from datetime import datetime
import PIL.Image
import PIL.ImageOps
import markdown
from bs4 import BeautifulSoup
from ddgs import DDGS
import docx
from xhtml2pdf import pisa

import threading
_loading_timer = None
_loading_shown = False
def _show_loading_message():
    global _loading_shown
    _loading_shown = True
    print("⏳ Loading... (May take a moment on first boot while modules initialize and cache)", flush=True)

# Only start timer if likely interactive mode (no args or just --interactive)
# We do this early so it can run while heavy modules load
if len(sys.argv) == 1 or (len(sys.argv) == 2 and sys.argv[1] in ('--interactive', '-I')):
    _loading_timer = threading.Timer(1.0, _show_loading_message) # Reduced to 1s for better UX
    _loading_timer.daemon = True
    _loading_timer.start()

# --- UI & Terminal Frameworks ---
from rich.console import Console
from rich.markdown import Markdown
from rich.status import Status
from prompt_toolkit import PromptSession
from prompt_toolkit.history import InMemoryHistory
from prompt_toolkit.formatted_text import HTML
from prompt_toolkit.completion import PathCompleter, NestedCompleter, FuzzyCompleter
from prompt_toolkit.lexers import Lexer
from prompt_toolkit.styles import Style

# --- Core AI Frameworks ---
try:
    import torch
except ImportError:
    torch = None

try:
    from transformers import pipeline, AutoTokenizer, AutoModelForCausalLM
except ImportError:
    pipeline = None

# --- Monkey Patch for basicsr/torchvision compatibility ---
try:
    from torchvision.transforms import functional_tensor
except ImportError:
    try:
        import torchvision.transforms.functional as F
        import sys
        from types import ModuleType
        sys.modules['torchvision.transforms.functional_tensor'] = ModuleType('torchvision.transforms.functional_tensor')
        sys.modules['torchvision.transforms.functional_tensor'].rgb_to_grayscale = F.rgb_to_grayscale
    except Exception:
        pass
# ----------------------------------------------------------

try:
    from tqdm import tqdm
except ImportError:
    tqdm = None

def _check_ffmpeg_encoder(encoder_name, w=256, h=256):
    """
    Check if FFmpeg can actually initialize the given encoder at target resolution.
    Used for probing hardware limits (e.g. NVENC max resolution).
    """
    try:
        # Run a tiny 1-frame test encoding to null at target resolution
        cmd = [
            'ffmpeg', '-y', '-f', 'lavfi', '-i', f'nullsrc=s={w}x{h}', 
            '-c:v', encoder_name, '-t', '0.1', '-f', 'null', '-'
        ]
        
        # Suppress output unless verbose debugging is needed
        subprocess.check_output(cmd, stderr=subprocess.STDOUT, timeout=5)
        return True
    except:
        return False

try:
    from basicsr.archs.rrdbnet_arch import RRDBNet
    from realesrgan import RealESRGANer
    HAS_REALESRGAN = True
except ImportError:
    HAS_REALESRGAN = False

# Suppress warnings
warnings.filterwarnings("ignore")

# Suppress verbose logging from transformers/diffusers
logging.getLogger("transformers").setLevel(logging.ERROR)
logging.getLogger("diffusers").setLevel(logging.ERROR)

# Set environment variable to suppress transformers warnings (must be before import)
os.environ["TRANSFORMERS_VERBOSITY"] = "error"
os.environ["DIFFUSERS_VERBOSITY"] = "error"
# Suppress resource_tracker warnings in subprocesses (caused by os._exit)
os.environ["PYTHONWARNINGS"] = "ignore::UserWarning:multiprocessing.resource_tracker"
os.environ["TOKENIZERS_PARALLELISM"] = "false" # Fix for deadlock warning

# CUDA Memory Optimization - Reduce fragmentation on Windows/NVIDIA
# This helps prevent "CUDA out of memory" errors even when GPU has free memory
os.environ["PYTORCH_ALLOC_CONF"] = "expandable_segments:True"


# Global arguments holder for JSON reporting
args = None

try:
    import ftfy
except ImportError:
    ftfy = None

# --- Constants ---
DEFAULT_IMAGE_SIZE = "720p"  # Maps to 1280x720
DEFAULT_AUDIO_SAMPLING = "32000" # Hz (MusicGen default is usually 32k)
DEFAULT_AUDIO_BITDEPTH = 16
DEFAULT_DURATION = "15s"

# Model Mappings (Short codes to Hugging Face Hub IDs)
IMAGE_MODELS = {
    "flux": "black-forest-labs/FLUX.1-schnell",        # State of the art, fast
    "flux-dev": "black-forest-labs/FLUX.1-dev",        # Higher quality, slower
    "sdxl": "stabilityai/sdxl-turbo",                  # Fast, good quality (DEFAULT)
    "sd-1.5": "runwayml/stable-diffusion-v1-5",        # Classic, lightweight
    "upscaler": "stabilityai/stable-diffusion-x4-upscaler", # 4x Upscaling
    "upscaler_x2": "stabilityai/sd-x2-latent-upscaler",     # 2x Latent Upscaling
    "default": "stabilityai/sdxl-turbo"
}

EDIT_MODELS = {
    "instruct-pix2pix": "timbrooks/instruct-pix2pix",
    "instruct-pix2pix-sdxl": "diffusers/sdxl-instructpix2pix-768",
    "remove-bg": "briaai/RMBG-1.4",
    "default": "timbrooks/instruct-pix2pix"
}

AUDIO_MODELS = {
    "musicgen-small": "facebook/musicgen-small",       # Fast, good for music
    "musicgen-medium": "facebook/musicgen-medium",     # Better quality music
    "musicgen-large": "facebook/musicgen-large",       # Best quality music
    "audioldm2": "cvssp/audioldm2",                    # General audio/SFX
    "stable-audio": "stabilityai/stable-audio-open-1.0", # Variable length, high quality (Gated)
    "bark": "suno/bark",                               # TTS / Audio (Transformer)
    "default": "facebook/musicgen-medium"
}

VIDEO_MODELS = {
    "ms-1.7b": "damo-vilab/text-to-video-ms-1.7b",     # Has watermark issues
    "zeroscope": "cerspense/zeroscope_v2_576w",        # 576x320 optimized (default)
    "zeroscope-xl": "cerspense/zeroscope_v2_XL",       # 1024x576 V2V upscaler (internal use)
    "cogvideox": "THUDM/CogVideoX-5b",                 # High quality (requires high VRAM)
    "wan2.2": "Wan-AI/Wan2.2-T2V-A14B-Diffusers",          # Alibaba Wan 2.2 (14B)
    "ltx-video": "Lightricks/LTX-Video",               # Lightricks LTX-Video (Fast, High Res)
    "mochi-1": "genmo/mochi-1-preview",                # Mochi 1 (Physics/Motion SOTA)
    "hunyuan": "hunyuanvideo-community/HunyuanVideo",                 # HunyuanVideo (13B, Cinematic)
    "svd": "stabilityai/stable-video-diffusion-img2vid-xt", # SVD Image-to-Video
    "default": "cerspense/zeroscope_v2_576w"
}

TEXT_MODELS = {
    # Reasoning-focused (Chain-of-Thought) - DeepSeek R1 Distilled
    # These show step-by-step reasoning before the final answer
    "deepseek-r1-qwen-7b": "deepseek-ai/DeepSeek-R1-Distill-Qwen-7B",      # ~7GB VRAM
    "deepseek-r1-qwen-14b": "deepseek-ai/DeepSeek-R1-Distill-Qwen-14B",    # ~14GB VRAM
    "deepseek-r1-qwen-32b": "deepseek-ai/DeepSeek-R1-Distill-Qwen-32B",    # ~24GB VRAM
    "deepseek-r1-llama-8b": "deepseek-ai/DeepSeek-R1-Distill-Llama-8B",    # ~8GB VRAM
    "deepseek-r1-llama-70b": "deepseek-ai/DeepSeek-R1-Distill-Llama-70B",  # ~40GB VRAM (requires high-end GPU)
    # General-purpose (Newer knowledge cutoffs)
    "qwen3-8b": "Qwen/Qwen3-8B",  # Note: May have MPS issues on Apple Silicon
    "qwen-2.5-14b": "Qwen/Qwen2.5-14B-Instruct",
    # Established models
    "llama-3.1-8b": "meta-llama/Meta-Llama-3.1-8B-Instruct",
    "mistral-nemo-12b": "mistralai/Mistral-Nemo-Instruct-2407",
    # Default: Llama 3.1 (stable on all platforms)
    "default": "meta-llama/Meta-Llama-3.1-8B-Instruct"
}

# Resolution Presets
RESOLUTIONS = {
    "480p": (854, 480),
    "576p": (1024, 576),
    "720p": (1280, 720),
    "900p": (1600, 900),
    "1080p": (1920, 1080),
    "1440p": (2560, 1440),
    "2k": (2048, 1080), # approx
    "3k": (3072, 1728),
    "2160p": (3840, 2160),
    "4k": (3840, 2160),
    "5k": (5120, 2880),
    "6k": (6144, 3456),
    "7k": (7168, 4032),
    "4320p": (7680, 4320),
    "8k": (7680, 4320),
    "9k": (9216, 5184),
    "10k": (10240, 5760),
    "hd": (1280, 720),
    "fhd": (1920, 1080),
    "uhd": (3840, 2160),
    "sd": (640, 480),
    "vga": (640, 480)
}

# Model resource requirements (estimated RAM/VRAM in GB)
# Based on model training specs, Hugging Face model cards, and practical testing.
# Format: { model_id: { "vram": X, "ram": Y, "max_resolution": (W, H) } }
MODEL_REQUIREMENTS = {
    # Image Models (max_resolution based on training data and VRAM constraints)
    "runwayml/stable-diffusion-v1-5": {"vram": 4, "ram": 8, "max_resolution": (1280, 1280)},  # Trained 512x512, works to ~1280
    "stabilityai/sdxl-turbo": {"vram": 8, "ram": 16, "max_resolution": (1536, 1536)},  # Trained 1024x1024
    "black-forest-labs/FLUX.1-schnell": {"vram": 16, "ram": 70, "max_resolution": (2048, 2048)}, # ~70GB on Mac (float32)
    "black-forest-labs/FLUX.1-dev": {"vram": 24, "ram": 80, "max_resolution": (2048, 2048)},
    # Audio Models (max_duration in seconds, based on model architecture limits)
    "facebook/musicgen-small": {"vram": 4, "ram": 8, "max_duration": 30},
    "facebook/musicgen-medium": {"vram": 8, "ram": 12, "max_duration": 60},
    "facebook/musicgen-large": {"vram": 16, "ram": 24, "max_duration": 120},
    "cvssp/audioldm2": {"vram": 8, "ram": 12, "max_duration": 60},
    "stabilityai/stable-audio-open-1.0": {"vram": 10, "ram": 16, "max_duration": 47}, # Max 47s training
    "suno/bark": {"vram": 4, "ram": 12, "max_duration": 30}, # Small/Large split, conservative est
    # Video Models (max_resolution based on training data)
    "damo-vilab/text-to-video-ms-1.7b": {"vram": 12, "ram": 16, "max_resolution": (1280, 720)},
    "cerspense/zeroscope_v2_576w": {"vram": 8, "ram": 12, "max_resolution": (576, 320)},
    "cerspense/zeroscope_v2_XL": {"vram": 10, "ram": 16, "max_resolution": (1024, 576)},  # V2V Upscaler
    "THUDM/CogVideoX-5b": {"vram": 32, "ram": 48, "max_resolution": (1920, 1080)}, # ~50GB on Mac (float32)
    "stabilityai/stable-video-diffusion-img2vid-xt": {"vram": 8, "ram": 12, "max_resolution": (1024, 576)},
    "Wan-AI/Wan2.2-T2V-A14B-Diffusers": {"vram": 24, "ram": 64, "max_resolution": (1280, 720)}, # ~14B params, high usage
    "Wan-AI/Wan2.2-I2V-A14B-Diffusers": {"vram": 24, "ram": 64, "max_resolution": (1280, 720)},
    "Lightricks/LTX-Video": {"vram": 16, "ram": 32, "max_resolution": (1216, 704)},      # Efficient DiT
    "genmo/mochi-1-preview": {"vram": 19, "ram": 48, "max_resolution": (848, 480)},      # High VRAM (10B DiT)
    "hunyuanvideo-community/HunyuanVideo": {"vram": 24, "ram": 64, "max_resolution": (1280, 720)},      # Huge 13B model
    "hunyuanvideo-community/HunyuanVideo-I2V": {"vram": 24, "ram": 64, "max_resolution": (1280, 720)},
    "stabilityai/stable-diffusion-x4-upscaler": {"vram": 8, "ram": 16, "max_resolution": (4096, 4096)},
    "stabilityai/sd-x2-latent-upscaler": {"vram": 4, "ram": 8, "max_resolution": (2048, 2048)},
    "timbrooks/instruct-pix2pix": {"vram": 8, "ram": 12, "max_resolution": (1024, 1024)},
    "diffusers/sdxl-instructpix2pix-768": {"vram": 10, "ram": 16, "max_resolution": (1024, 1024)},
    "briaai/RMBG-1.4": {"vram": 4, "ram": 8, "max_resolution": (2048, 2048)},
    # Text Models
    "meta-llama/Meta-Llama-3.1-8B-Instruct": {"vram": 16, "ram": 24, "max_resolution": None},
    "mistralai/Mistral-Nemo-Instruct-2407": {"vram": 24, "ram": 32, "max_resolution": None},
    "Qwen/Qwen2.5-14B-Instruct": {"vram": 28, "ram": 48, "max_resolution": None},
}


def clear_gpu_memory():
    """Clear GPU memory cache to reduce fragmentation and prevent OOM errors.
    
    Call this between heavy operations to free unused memory.
    """
    import gc
    gc.collect()
    
    try:
        import torch
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
            torch.cuda.synchronize()
        elif torch.backends.mps.is_available():
            # MPS doesn't have explicit cache clearing, but gc helps
            pass
    except ImportError:
        pass


def get_system_resources():
    """Get available system RAM and VRAM."""
    ram_available = 0
    vram_available = 0
    vram_total = 0
    
    try:
        mem = psutil.virtual_memory()
        ram_available = mem.available / (1024**3)  # GB
    except ImportError:
        pass
    
    try:
        import torch
        if torch.cuda.is_available():
            vram_total = torch.cuda.get_device_properties(0).total_memory / (1024**3)
            vram_used = torch.cuda.memory_allocated(0) / (1024**3)
            vram_available = vram_total - vram_used
        elif torch.backends.mps.is_available():
            # MPS uses unified memory - estimate as 75% of RAM for GPU tasks
            try:
                vram_available = psutil.virtual_memory().available / (1024**3) * 0.75
            except:
                vram_available = 8  # Conservative default
    except ImportError:
        pass
    
    return ram_available, vram_available


def check_resources_and_warn(model_id, width=None, height=None, duration=None, force=False):
    """
    Check if system resources are sufficient for the requested task.
    Returns True to proceed, False to abort.
    """
    reqs = MODEL_REQUIREMENTS.get(model_id)
    if not reqs:
        return True  # Unknown model, proceed with caution
    
    ram_available, vram_available = get_system_resources()
    warnings = []
    
    # Check RAM
    if ram_available > 0 and ram_available < reqs.get("ram", 0):
        warnings.append(f"RAM: {ram_available:.1f}GB available, {reqs['ram']}GB recommended")
    
    # Check VRAM (if available)
    if vram_available > 0 and vram_available < reqs.get("vram", 0):
        warnings.append(f"VRAM: {vram_available:.1f}GB available, {reqs['vram']}GB recommended")
    
    # Check resolution limits
    max_res = reqs.get("max_resolution")
    is_zeroscope = "zeroscope" in model_id.lower() and "xl" not in model_id.lower()
    
    if max_res and width and height:
        if width > max_res[0] or height > max_res[1]:
            if is_zeroscope:
                # Zeroscope has dynamic upscaling - this is informational, not a warning
                warnings.append(f"Resolution: {width}x{height} exceeds native {max_res[0]}x{max_res[1]} (Dynamic Upscaling will be used)")
            else:
                warnings.append(f"Resolution: {width}x{height} exceeds recommended max {max_res[0]}x{max_res[1]}")
    
    # Check duration limits (for audio/video)
    max_dur = reqs.get("max_duration")
    if max_dur and duration and duration > max_dur:
        warnings.append(f"Duration: {duration}s exceeds recommended max {max_dur}s")
    
    if not warnings:
        return True
    
    # Display warnings - use different style for zeroscope upscaling info
    if is_zeroscope and len(warnings) == 1 and "Dynamic Upscaling" in warnings[0]:
        # Zeroscope-specific informational message (not a scary warning)
        # Check if MPS (XL is skipped on Apple Silicon)
        import torch
        is_mps = torch.backends.mps.is_available() and not torch.cuda.is_available()
        
        print("\n📐 Dynamic Upscaling Pipeline:\n")
        print(f"   Target:   {width}x{height}")
        print(f"   Native:   {max_res[0]}x{max_res[1]}")
        if is_mps:
            print(f"   Method:   Zeroscope 576w → Real-ESRGAN → Target (XL skipped on Mac)")
        else:
            print(f"   Method:   Zeroscope 576w → XL (1024x576) → Real-ESRGAN → Target")
        print(f"\n   Model: {model_id}")
        print(f"   ℹ️  This is the optimal workflow for high-res zeroscope output.\n")
        
        if force:
            return True
        
        try:
            choice = input("   Proceed with dynamic upscaling? [Y/n]: ").lower().strip()
            if choice in ['', 'y', 'yes']:
                print("")  # Spacer
                return True
            print("\n💡 Tip: Use -s 576x320 for native resolution (fastest).\n")
            sys.exit(0)
        except KeyboardInterrupt:
            print("\n❌ Operation cancelled.\n")
            sys.exit(0)
    
    # Standard warning display for other cases
    print("\n⚠️  Resource Warning:\n")
    for w in warnings:
        print(f"   • {w}")
    print(f"\n   Model: {model_id}")
    
    # Check for VAE Tiling condition (Resolution > 1536x1536)
    if width and height:
        total_pixels = width * height
        if total_pixels > 3072 * 3072: # ~9.4MP
             print(f"\n   ⚠️  CRITICAL WARNING: Resolution {width}x{height} is extremely high.")
             print(f"      Standard generation will likely fail with 'Invalid buffer size'.")
             print(f"      Recommended: Generate at 2K/4K and use an external upscaler.")
             print(f"      💡 Or try: -s 720p --upscale -uf 4x (to get 5K)\n")
        elif total_pixels > 1536 * 1536:
            print(f"\n   ℹ️  Note: VAE Tiling will be enabled to reduce memory usage.\n")
        
    print(f"   This job may cause slowdowns, swapping, or crashes.\n")
    
    if force:
        print("   (Proceeding due to --force flag)\n")
        return True
    
    try:
        choice = input("   Continue anyway? [y/N]: ").lower().strip()
        if choice in ['y', 'yes']:
            print("")  # Spacer
            return True
        print("\n💡 Tip: Try a smaller resolution, shorter duration, or lighter model.\n")
        sys.exit(0)
    except KeyboardInterrupt:
        print("\n❌ Operation cancelled.\n")
        sys.exit(0)

# --- Signal Handling ---
# Global test state for CTRL+C handling
_test_state = {
    'active': False,
    'passed': 0,
    'failed': 0,
    'total': 0,
    'current_process': None  # Track subprocess for cleanup
}
_interrupted = False
_force_kill_timer = None
_first_interrupt_time = None

def _force_exit():
    """Forcefully terminate the process after timeout."""
    print("\n💀 Force-killing process (GPU operations did not respond in time)...")
    os._exit(1)  # Immediate exit, bypasses cleanup

def _kill_test_subprocess():
    """Kill the current test subprocess and its children."""
    proc = _test_state.get('current_process')
    if proc:
        pid = proc.pid
        print(f"\n   ⚠️  Killing subprocess tree (PID {pid})...")
        if os.name == 'nt':
            try:
                import subprocess
                subprocess.run(['taskkill', '/T', '/F', '/PID', str(pid)], 
                               capture_output=True, timeout=5)
                print(f"   ✅ Process tree terminated.")
            except Exception as e:
                print(f"   ⚠️  taskkill failed: {e}")
                try:
                    proc.kill()
                except:
                    pass
        else:
            try:
                os.killpg(os.getpgid(pid), signal.SIGKILL)
            except:
                try:
                    proc.kill()
                except:
                    pass

def signal_handler(sig, frame):
    global _interrupted, _force_kill_timer, _first_interrupt_time
    
    if _test_state['active']:
        # Kill subprocess first
        _kill_test_subprocess()
        
        completed = _test_state['passed'] + _test_state['failed']
        print(f"\n\n{'='*60}")
        print(f"❌ Test suite interrupted by user (CTRL+C)")
        print(f"{'='*60}")
        print(f"   Completed: {completed}/{_test_state['total']}")
        print(f"   Passed: {_test_state['passed']} ✅")
        if _test_state['failed'] > 0:
            print(f"   Failed: {_test_state['failed']} ❌")
        else:
            print(f"   Failed: {_test_state['failed']}")
        print(f"   Skipped: {_test_state['total'] - completed}")
        print(f"{'='*60}")
        sys.exit(130)
    else:
        _interrupted = True
        
        # Immediately exit - GPU/multiprocessing cleanup causes hangs
        print("\n\n⚠️  Interrupted!")
        os._exit(0)  # Immediate exit, no cleanup (cleanup hangs with model_cpu_offload)

signal.signal(signal.SIGINT, signal_handler)
signal.signal(signal.SIGTERM, signal_handler)

# --- Helper Parsing Functions ---

def get_optimal_device_and_dtype(quiet=False):
    """
    Detect the best available hardware (CUDA, MPS, or CPU) 
    and return the device string and optimal torch dtype.
    """
    try:
        import torch
        if torch.cuda.is_available():
            if not quiet: print(f"🚀 Detected NVIDIA GPU: Using CUDA\n")
            return torch.device("cuda"), torch.float16
            
        if torch.backends.mps.is_available():
            if not quiet: print(f"🍎 Detected Apple Silicon: Using MPS (Metal Performance Shaders)\n")
            return torch.device("mps"), torch.float16
    except ImportError:
        pass
        
    if not quiet: print(f"💻 Using CPU (Slow): CUDA or MPS not detected (or torch missing)\n")
    return torch.device("cpu"), torch.float32

def parse_size(value):
    """
    Parse size string or object into (width, height).
    Accepts:
      - "480p", "720p", "1080p", "1440p"
      - "1k", "2k", "3k", "4k", "5k", ... "10k"
      - "1280x720"
      - "w: 1280, h: 720" (Braces {} are optional)
    """
    if not value:
        return RESOLUTIONS[DEFAULT_IMAGE_SIZE]
        
    normalized = value.strip().lower()
    
    # Check presets
    if normalized in RESOLUTIONS:
        return RESOLUTIONS[normalized]
        
    # Check WxH format
    if 'x' in normalized:
        try:
            w, h = map(int, normalized.split('x'))
            return (w, h)
        except ValueError:
            pass
            
    # Check Object/JSON-like format
    if '{' in normalized and '}' in normalized:
        try:
            w = None
            h = None
            
            # Remove braces
            content = normalized.strip("{}")
            parts = content.split(',')
            
            for part in parts:
                if ':' not in part: continue
                k, v = part.split(':', 1)
                k = k.strip()
                v = int(re.sub(r'[^0-9]', '', v)) # extract number
                
                if k in ['w', 'width']:
                    w = v
                elif k in ['h', 'height']:
                    h = v
            
            if w and h:
                return (w, h)
        except Exception as e:
            print(f"Warning: Failed to parse object size '{value}': {e}")
    
    # Fallback default if parsing fails
    print(f"Warning: Could not parse size '{value}'. Using default 720p.")
    return RESOLUTIONS["720p"]


def parse_duration(value):
    """
    Parse duration into seconds (float).
    Accepts:
      - Numeric (seconds)
      - Strings: "15s", "1m", "1h50m", "50s"
      - Objects: "{h:1, m:25, s:10}"
    """
    if not value:
        return 15.0
        
    # If standard number strings "15", "15.5"
    try:
        return float(value)
    except ValueError:
        pass
        
    normalized = str(value).strip().lower()
    
    total_seconds = 0
    
    # Check Object format
    if '{' in normalized:
        try:
            content = normalized.strip("{}")
            parts = content.split(',')
            for part in parts:
                if ':' not in part: continue
                k, v = part.split(':', 1)
                k = k.strip()
                try:
                    val = float(re.sub(r'[^0-9\.]', '', v))
                except: continue
                
                if k in ['h', 'hours', 'hour']:
                    total_seconds += val * 3600
                elif k in ['m', 'mins', 'min', 'minutes', 'minute']:
                    total_seconds += val * 60
                elif k in ['s', 'sec', 'secs', 'seconds', 'second']:
                    total_seconds += val
            return total_seconds
        except Exception:
            pass

    # Check String format "1h50m10s"
    # Regex to find pairs of number+unit
    pattern = r'(\d+(?:\.\d+)?)\s*([hms])'
    matches = re.findall(pattern, normalized)
    
    if matches:
        for val_str, unit in matches:
            val = float(val_str)
            if unit == 'h':
                total_seconds += val * 3600
            elif unit == 'm':
                total_seconds += val * 60
            elif unit == 's':
                total_seconds += val
        return total_seconds
        
    # Fallback logic for "M:S" or "H:M:S" or just "50s"
    if ':' in normalized:
        parts = normalized.split(':')
        parts = [float(p) for p in parts]
        if len(parts) == 3: # H:M:S
            return parts[0]*3600 + parts[1]*60 + parts[2]
        elif len(parts) == 2: # M:S
            return parts[0]*60 + parts[1]
            
    return 15.0


def parse_sampling_rate(value):
    """Parse sampling rate string to integer Hz."""
    if not value:
        return 32000
    
    normalized = str(value).strip().lower()
    
    # Handle "44.1khz" -> 44100
    if 'k' in normalized:
        try:
            num = float(re.sub(r'[^0-9\.]', '', normalized))
            return int(num * 1000)
        except:
             pass # Fallthrough to default behavior
    
    try:
        return int(re.sub(r'[^0-9]', '', normalized))
    except:
        return 32000


# --- Format Helpers ---
def format_time(seconds):
    """Convert seconds to human readable string (e.g. 2w 1d 1h 2m 3.5s)."""
    if not seconds:
        return "0s"
        
    current = float(seconds)
    intervals = (
        ('w', 604800),  # 60 * 60 * 24 * 7
        ('d', 86400),   # 60 * 60 * 24
        ('h', 3600),    # 60 * 60
        ('m', 60),
    )
    
    result = []
    for name, count in intervals:
        value = int(current // count)
        if value:
            current -= value * count
            result.append(f"{value}{name}")
            
    # Remaining seconds with precision
    if current > 0 or not result:
        # If integer-ish, show int, else float
        if current % 1 == 0:
            result.append(f"{int(current)}s")
        else:
            result.append(f"{current:.1f}s")
            
    return " ".join(result)


def parse_bitrate(value):
    """Return bitrate string in standardized format or passed through if complex."""
    if not value:
        return None
    return value.strip()


# --- Generation Logic Wrappers ---

def write_report_json(path, stats):
    """Write generation stats to a JSON file."""
    try:
        import json
        with open(path, 'w') as f:
            json.dump(stats, f, indent=2)
    except Exception as e:
        print(f"⚠️ Failed to write report JSON: {e}")

def generate_image(prompt, output_file, width, height, model_name="default", steps=30, guidance_scale=7.5, unsafe=False):
    """Generate image using Diffusers (Flux/SDXL)."""
    
    # Resolve Model ID
    model_id = IMAGE_MODELS.get(model_name.lower(), model_name) # Allow raw ID if not in map
    if model_name.lower() == "default": model_id = IMAGE_MODELS["default"]
    
    print(f"🎨 Generating Image")
    print(f"   Model:  {model_id}")
    print(f"   Prompt: '{prompt}'")
    print(f"   Size:   {width}x{height}")
    print(f"   Output: {output_file}")
    print("") # Spacer
    
    try:
        import torch
        from diffusers import FluxPipeline, AutoPipelineForText2Image
        
        # Determine device
        device, dtype = get_optimal_device_and_dtype(quiet=True)
        

        
        # Determine Pipeline Class based on model
        if "flux" in model_id.lower():
            # FLUX on MPS requires float32 to avoid dtype mismatch errors
            flux_dtype = torch.float32 if device.type == "mps" else dtype
            pipe = FluxPipeline.from_pretrained(
                model_id, 
                torch_dtype=flux_dtype
            )
            # Flux parameters
            extra_kwargs = {
                "guidance_scale": 0.0, 
                "num_inference_steps": 4, 
                "max_sequence_length": 256
            }
        elif "sdxl-turbo" in model_id.lower() or "turbo" in model_id.lower():
            # SDXL Turbo requires very few steps and specific settings
            # Use float32 on MPS to avoid black images (VAE produces NaN in float16)
            sdxl_dtype = torch.float32 if device.type == "mps" else dtype
            pipe = AutoPipelineForText2Image.from_pretrained(
                model_id, 
                torch_dtype=sdxl_dtype,
                variant="fp16" if sdxl_dtype == torch.float16 else None
            )
            # SDXL Turbo needs low steps and no guidance
            extra_kwargs = {
                "guidance_scale": 0.0,
                "num_inference_steps": 4
            }
        else:
            # Generic Text2Image (SD1.5, etc.)
            # On MPS, generic models are safer in float32 to prevent dtype mismatches
            run_dtype = torch.float32 if device.type == "mps" else dtype
            
            pipe = AutoPipelineForText2Image.from_pretrained(
                model_id, 
                torch_dtype=run_dtype,
                variant="fp16" if run_dtype == torch.float16 else None
            )
            extra_kwargs = {} # Use defaults
            
        pipe.to(device)
        
        # Disable NSFW safety checker if requested
        if unsafe:
            pipe.safety_checker = None
            pipe.requires_safety_checker = False
        
        # Memory optimization and VAE fix for MPS/Mac
        if device.type == "mps":
            pipe.enable_attention_slicing()
            # Fix black images on MPS: VAE produces NaN in float16, use float32
            if hasattr(pipe, 'vae'):
                pipe.vae = pipe.vae.to(torch.float32)
            # Fix Safety Checker dtype mismatch on MPS (Input Half vs Bias Float)
            if hasattr(pipe, 'safety_checker') and pipe.safety_checker is not None:
                pipe.safety_checker = pipe.safety_checker.to(torch.float32)

        # High-Resolution Memory Optimization (VAE Tiling)
        # 4K images (3840x2160 = ~8.3MP) cause massive VRAM spikes during decoding without tiling
        # Trigger tiling if pixels > 1536x1536 (~2.3MP)
        total_pixels = width * height
        if total_pixels > 1536 * 1536:
            print(f"ℹ️  High resolution detected ({width}x{height}).")
            print(f"   ✓ Enabling VAE Tiling (Memory Optimization)\n")
            if hasattr(pipe, 'vae') and hasattr(pipe.vae, 'enable_tiling'):
                pipe.vae.enable_tiling()
            else:
                pipe.enable_vae_tiling()
        
        # --- Performance Tracking ---
        tracker = PerformanceTracker()
        
        print(f"🎨 Generating {width}x{height} image... (This may take a moment)")
        
        # Suppress RuntimeWarning from diffusers image_processor during NSFW filtering
        start_time = time.time()
        with warnings.catch_warnings():
            warnings.filterwarnings("ignore", category=RuntimeWarning, message="invalid value encountered in cast")
            
            # Start Resource Monitoring
            with ResourceMonitor() as monitor:
                output = pipe(
                    prompt, 
                    height=height, 
                    width=width,
                    **extra_kwargs
                )
            
            # Collect metrics
            duration = time.time() - start_time
            avg_cpu, avg_ram, avg_vram, avg_gpu = monitor.get_averages()
            
            # Record Performance
            tracker.record_image(model_id, width, height, device, duration, cpu=avg_cpu, ram=avg_ram, vram=avg_vram, gpu=avg_gpu)
            print(f"   ✓ Generated in {format_time(duration)} (RAM: {avg_ram:.1f}GB | VRAM: {avg_vram:.1f}GB | CPU: {avg_cpu:.1f}% | GPU: {avg_gpu:.1f}%)")
            
            # Write JSON report if requested
            if 'args' in globals() and hasattr(globals()['args'], 'report_json') and globals()['args'].report_json:
                 stats = {
                     "time": duration,
                     "ram": avg_ram,
                     "vram": avg_vram,
                     "cpu": avg_cpu,
                     "gpu": avg_gpu,
                     "width": width,
                     "height": height
                 }
                 write_report_json(globals()['args'].report_json, stats)
            
        image = output.images[0]
        
        # Check for NSFW content interception
        if hasattr(output, "nsfw_content_detected") and output.nsfw_content_detected:
            # output.nsfw_content_detected is a list of booleans
            if output.nsfw_content_detected[0]:
                print(f"⚠️  Warning: Potential NSFW content detected.\n")
                print(f"The model's safety checker has blocked the image (returning a black frame).")
                print(f"👉 Please modify your prompt and try again.")
                print(f"💡 If your prompt is appropriate, try again with --unsafe to disable the safety checker.\n")
        
        image.save(output_file)
        print(f"✅ Image saved to {output_file}")
        return True
        
    except ImportError as e:
        print(f"❌ Error: Missing dependencies. {e}")
        return False
    except Exception as e:
        err_str = str(e)
        if "401" in err_str or "restricted" in err_str.lower():
            print(f"❌ Access Denied (401). Test model '{model_id}' is gated.")
            print("   👉 Solution 1: Run 'huggingface-cli login' with your token.")
            print("   👉 Solution 2: Use an open model like '--image-model sd-1.5'.")
        elif "divisible by 8" in err_str:
            print(f"❌ Resolution Error: {e}")
            
            # Smart Correction
            new_w = round(width / 8) * 8
            new_h = round(height / 8) * 8
            
            print(f"\n💡 Tip: Dimensions must be multiples of 8.")
            print(f"   Closest valid size: {new_w}x{new_h}")
            
            try:
                choice = input(f"   🔄 Retry with {new_w}x{new_h}? [y/N]: ").lower().strip()
                if choice in ['y', 'yes']:
                    print("") # Spacer
                    return generate_image(prompt, output_file, new_w, new_h, model_name=model_name, unsafe=unsafe)
            except KeyboardInterrupt:
                pass
            print("")
        elif "Invalid buffer size" in err_str:
            print(f"\n❌ Hardware Limitation Reached (Single Buffer Limit)")
            print(f"   Error: {e}")
            print(f"\n   Explanation:")
            print(f"   • Native {width}x{height} generation requires calculating a massive Attention Matrix.")
            print(f"   • This exceeded the maximum allowed size for a single tensor (usually ~4GB on MPS/Metal).")
            print(f"   • This is a hardware/driver limit, not a VRAM limit.")
            print(f"\n   💡 Solution: Use a lower resolution (e.g. 4k or 2k).")
            print(f"      (Native 5K generation requires 'MultiDiffusion' tiling which is not currently supported.)\n")
            
            # Auto-Upscale Fallback for 5K requests
            # If target was roughly 5K (5120x2880), offering 1280x720 (720p) -> x4 Upscale = 5120x2880
            # 1280 * 4 = 5120
            
            try:
                print(f"   ✨ Alternative: Generate at 1280x720 and Auto-Upscale x4?")
                print(f"      This produces a 5120x2880 (5K) image using the Upscaler model.")
                choice = input(f"   🔄 Try Auto-Upscale workflow? [y/N]: ").lower().strip()
                if choice in ['y', 'yes']:
                    print("\n📉 Switching to base resolution: 1280x720...")
                    # 1. Generate Base Image
                    success = generate_image(prompt, output_file, 1280, 720, model_name=model_name, unsafe=unsafe)
                    if success:
                        # 2. Upscale Result
                        print("")
                        return upscale_image_file(output_file, output_file, strength=0.0, factor=4.0) # Overwrite
            except KeyboardInterrupt:
                pass
            print("")
        else:
            print(f"❌ Generation failed: {e}")
        return False


def generate_caption(input_path, device, quiet=False, model_type="florence"):
    """
    Generate a text description for an image or video.
    Models: 
      - 'florence' (Microsoft Florence-2-Large, SOTA)
      - 'blip' (Salesforce BLIP-Large, Classic)
    """
    try:
        if not quiet: print(f"👁️  Analyzing input: {input_path}")
        
        # ----------------------------------------------------------------
        # MODEL: BLIP
        # ----------------------------------------------------------------
        if model_type == "blip":
            from transformers import BlipProcessor, BlipForConditionalGeneration
            from diffusers.utils import load_image
            import torch
            from PIL import Image
            import cv2
            import numpy as np

            caption_model_id = "Salesforce/blip-image-captioning-large"
            if not quiet: print(f"   Loading Caption Model: {caption_model_id}...")
            
            processor = BlipProcessor.from_pretrained(caption_model_id)
            model = BlipForConditionalGeneration.from_pretrained(caption_model_id).to(device)
            
            # Check if video
            ext = input_path.lower().split('.')[-1]
            is_video = ext in ['mp4', 'mov', 'avi', 'mkv', 'webm', 'flv', 'gif']
            
            if is_video:
                 # Video Logic (Same as before)
                cap = cv2.VideoCapture(input_path)
                if not cap.isOpened():
                    return "Unknown video content"
                total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
                fps = cap.get(cv2.CAP_PROP_FPS)
                num_samples = 10
                indices = np.linspace(0, total_frames - 1, num_samples, dtype=int)
                captions = []
                for i, idx in enumerate(indices):
                    cap.set(cv2.CAP_PROP_POS_FRAMES, idx)
                    ret, frame = cap.read()
                    if not ret: continue
                    rgb_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                    pil_image = Image.fromarray(rgb_frame)
                    
                    inputs = processor(pil_image, return_tensors="pt").to(device)
                    out = model.generate(**inputs)
                    frame_caption = processor.decode(out[0], skip_special_tokens=True)
                    captions.append(frame_caption)
                cap.release()
                return ", ".join(captions)
            else:
                # Image Logic
                raw_image = load_image(input_path).convert('RGB')
                inputs = processor(raw_image, return_tensors="pt").to(device)
                out = model.generate(**inputs)
                caption = processor.decode(out[0], skip_special_tokens=True)
                if not quiet: print(f"   Detected: '{caption}'")
                return caption

        # ----------------------------------------------------------------
        # MODEL: Florence-2 (Default/SOTA)
        # ----------------------------------------------------------------
        else:
            from transformers import AutoProcessor, AutoModelForCausalLM
            from diffusers.utils import load_image
            import cv2
            import numpy as np
            import torch
            from PIL import Image
            
            # Load Florence-2 (SOTA Captioning, ~1.5GB)
            # Note: Requires timm and trust_remote_code=True
            caption_model_id = "microsoft/Florence-2-large"
            
            if not quiet: print(f"   Loading Caption Model: {caption_model_id}...")
            
            # Florence-2 needs device_map or manual to(device). 
            # On MPS, manual to(device) is safer for now.
            processor = AutoProcessor.from_pretrained(caption_model_id, trust_remote_code=True)
            
            # Use eager attention to avoid SDPA crashes on MPS/Mac with recent transformers
            model = AutoModelForCausalLM.from_pretrained(
                caption_model_id, 
                trust_remote_code=True,
                attn_implementation="eager"
            ).to(device) 
            
            # Task prompt for Florence-2
            task_prompt = "<MORE_DETAILED_CAPTION>"
            
            # Check if video
            ext = input_path.lower().split('.')[-1]
            is_video = ext in ['mp4', 'mov', 'avi', 'mkv', 'webm', 'flv', 'gif']
            
            if is_video:
                cap = cv2.VideoCapture(input_path)
                if not cap.isOpened():
                    if not quiet: print(f"⚠️  Could not open video: {input_path}")
                    return "Unknown video content"
                    
                total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
                fps = cap.get(cv2.CAP_PROP_FPS)
                duration = total_frames / fps if fps > 0 else 0
                
                # Select 10 evenly distributed frames
                num_samples = 10
                indices = np.linspace(0, total_frames - 1, num_samples, dtype=int)
                
                captions = []
                if not quiet: print(f"   Analyzing {num_samples} frames from video ({duration:.1f}s)...")
                
                for i, idx in enumerate(indices):
                    cap.set(cv2.CAP_PROP_POS_FRAMES, idx)
                    ret, frame = cap.read()
                    if not ret: continue
                    
                    # Convert BGR (OpenCV) to RGB (PIL)
                    rgb_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                    pil_image = Image.fromarray(rgb_frame)
                    
                    # Generate caption (Florence-2)
                    inputs = processor(text=task_prompt, images=pil_image, return_tensors="pt") # .to(device) moved below
                    
                    # Ensure pixel_values are correct dtype if needed, though from_pretrained handles it usually.
                    # On MPS, float32 is safest.
                    if device.type == "mps":
                        inputs["pixel_values"] = inputs["pixel_values"].to(device, torch.float32)
                        inputs["input_ids"] = inputs["input_ids"].to(device)
                    else:
                        inputs = inputs.to(device)

                    # Disable cache to avoid MPS past_key_values crash
                    generated_ids = model.generate(
                        input_ids=inputs["input_ids"],
                        pixel_values=inputs["pixel_values"],
                        max_new_tokens=1024,
                        do_sample=False,
                        num_beams=3,
                        use_cache=False,
                    )
                    
                    generated_text = processor.batch_decode(generated_ids, skip_special_tokens=False)[0]
                    parsed_answer = processor.post_process_generation(generated_text, task=task_prompt, image_size=(pil_image.width, pil_image.height))
                    frame_caption = parsed_answer[task_prompt]
                    
                    timestamp = idx / fps if fps > 0 else 0
                    captions.append(f"[{timestamp:.1f}s]: {frame_caption}")
                    if not quiet: print(f"   Frame {i+1}/{num_samples} ({timestamp:.1f}s): {frame_caption}")
                    
                cap.release()
                
                # Consolidated description
                summary = ", ".join([c.split(": ")[1] for c in captions])
                return summary 
                
            else:
                # Image handling
                raw_image = load_image(input_path).convert('RGB')
                
                inputs = processor(text=task_prompt, images=raw_image, return_tensors="pt") # .to(device) moved below
                
                # Ensure pixel_values are correct dtype if needed, though from_pretrained handles it usually.
                # On MPS, float32 is safest.
                if device.type == "mps":
                    inputs["pixel_values"] = inputs["pixel_values"].to(device, torch.float32)
                    inputs["input_ids"] = inputs["input_ids"].to(device)
                else:
                    inputs = inputs.to(device)

                # Disable cache to avoid MPS past_key_values crash
                generated_ids = model.generate(
                    input_ids=inputs["input_ids"],
                    pixel_values=inputs["pixel_values"],
                    max_new_tokens=1024,
                    do_sample=False,
                    num_beams=3,
                    use_cache=False,
                )
                
                generated_text = processor.batch_decode(generated_ids, skip_special_tokens=False)[0]
                parsed_answer = processor.post_process_generation(generated_text, task=task_prompt, image_size=(raw_image.width, raw_image.height))
                caption = parsed_answer[task_prompt]
                
                if not quiet: print(f"   Detected: '{caption}'")
                return caption
                
    except ImportError as e:
        print(f"❌ Error: Missing dependencies for captioning. {e}")
        return None
    except Exception as e:
        print(f"❌ Caption generation failed: {e}")
        return None


def generate_edit(input_path, prompt, output_path, model_name="default", guidance_scale=7.5, image_guidance_scale=1.5, steps=50, unsafe=False):
    """
    Edit an existing image based on instructions using InstructPix2Pix.
    """
    import torch
    from diffusers import StableDiffusionInstructPix2PixPipeline, StableDiffusionXLInstructPix2PixPipeline
    from diffusers.utils import load_image
    
    # Resolve Model ID
    model_id = EDIT_MODELS.get(model_name.lower(), model_name)
    if model_name.lower() == "default": model_id = EDIT_MODELS["default"]
    
    print(f"🎨 Editing Image")
    print(f"   Model:     {model_id}")
    print(f"   Input:     {input_path}")
    print(f"   Instruct:  '{prompt}'")
    print(f"   Output:    {output_path}")
    print("") 

    try:
        device, dtype = get_optimal_device_and_dtype(quiet=True)
        
        # CRITICAL FIX: InstructPix2Pix (SD1.5 based) often produces black images on MPS with float16.
        # We force float32 for this specific pipeline on MPS to ensure valid output.
        if device.type == "mps":
            print(f"   ℹ️  MPS Detected: Forcing float32 for InstructPix2Pix to prevent black images.")
            dtype = torch.float32

        # Load Input Image
        image = load_image(input_path)
        image = PIL.ImageOps.exif_transpose(image)
        image = image.convert("RGB")
        
        # Initialize Pipeline
        if "sdxl" in model_id.lower():
            # SDXL InstructPix2Pix
            # Note: This model uses torch_dtype for precision, not separate variant files
            pipe = StableDiffusionXLInstructPix2PixPipeline.from_pretrained(
                model_id,
                torch_dtype=dtype
            ).to(device)
            # Default scales for SDXL are different
            if guidance_scale == 7.5: guidance_scale = 7.0 
            if image_guidance_scale == 1.5: image_guidance_scale = 1.25
            
        else:
            # Standard InstructPix2Pix (SD 1.5 based)
            kwargs = {}
            if unsafe:
                kwargs["safety_checker"] = None
                
            pipe = StableDiffusionInstructPix2PixPipeline.from_pretrained(
                model_id, 
                torch_dtype=dtype,
                **kwargs
            ).to(device)
            
        # Optimization
        if device.type == "mps":
             pipe.enable_attention_slicing()
        
        # Generate
        print(f"✨ Applying edits... (Steps: {steps}, Text Guide: {guidance_scale}, Image Guide: {image_guidance_scale})")
        with torch.inference_mode():
            output = pipe(
                prompt,
                image=image,
                num_inference_steps=steps,
                guidance_scale=guidance_scale,
                image_guidance_scale=image_guidance_scale
            )
            
        result = output.images[0]
        
        # Safety Check
        if hasattr(output, "nsfw_content_detected") and output.nsfw_content_detected:
             if output.nsfw_content_detected[0]:
                print(f"⚠️  Warning: Potential NSFW content detected. Black image returned.")
                print(f"    ℹ️  This is frequently a false positive on Mac/MPS/CPU.")
                print(f"    👉  Please retry this command with '--unsafe' to see the image.")
                
        result.save(output_path)
        print(f"✅ Edited image saved to {output_path}")
        return True

    except Exception as e:
        print(f"❌ Edit failed: {e}")
        return False


def remove_background(input_path, output_path, model_name="remove-bg", silhouette=False):
    """
    Remove background using RMBG-1.4 (Transformer based).
    """
    import torch
    from transformers import AutoModelForImageSegmentation
    from torchvision.transforms.functional import normalize
    import numpy as np
    
    print(f"✂️  Removing Background")
    print(f"   Input:  {input_path}")
    print(f"   Output: {output_path}")
    if silhouette: print(f"   Mode:   Silhouette Maker")
    print("")

    try:
        device, dtype = get_optimal_device_and_dtype(quiet=True)
        
        # Load Model
        model_id = EDIT_MODELS.get(model_name, "briaai/RMBG-1.4")
        model = AutoModelForImageSegmentation.from_pretrained(model_id, trust_remote_code=True)
        model.to(device)
        model.eval()
        
        # Load Image
        image = PIL.Image.open(input_path).convert("RGB")
        image = PIL.ImageOps.exif_transpose(image)
        original_size = image.size
        
        # Preprocess (Model expects 1024x1024)
        model_input_size = (1024, 1024)
        image_scaled = image.resize(model_input_size, PIL.Image.BILINEAR)
        im_tensor = torch.tensor(np.array(image_scaled)).permute(2, 0, 1).unsqueeze(0).float() / 255.0
        im_tensor = normalize(im_tensor, [0.5, 0.5, 0.5], [1.0, 1.0, 1.0]).to(device)
        
        # Inference
        print(f"🧠 Processing...")
        with torch.inference_mode():
            result = model(im_tensor)
        
        # Post-process Mask
        result = torch.nn.functional.interpolate(result[0][0], size=original_size[::-1], mode='bilinear')
        ma = torch.sigmoid(result)
        ma = (ma - ma.min()) / (ma.max() - ma.min())
        
        # Create PIL Mask from Tensor
        mask = ma.cpu().data.numpy()[0, 0]
        mask_pil = PIL.Image.fromarray((mask * 255).astype(np.uint8))
        
        # Apply Mask
        if silhouette:
            # Create black foreground
            foreground = PIL.Image.new("RGBA", original_size, (0, 0, 0, 255))
            final_image = PIL.Image.new("RGBA", original_size, (255, 255, 255, 0)) # Transparent
            final_image.paste(foreground, (0, 0), mask_pil)
        else:
            final_image = image.copy()
            final_image.putalpha(mask_pil)
            
        final_image.save(output_path, "PNG")
        print(f"✅ Saved to {output_path}")
        return True
        
    except Exception as e:
        print(f"❌ Background removal failed: {e}")
        return False

def generate_long_bark(prompt, processor, model, device, voice_preset, sample_rate=24000):
    """
    Generate long-form audio with Bark by splitting text into sentences.
    Avoids 'history' chaining to prevent hallucinations/degradation.
    Concatenates independent chunks with the same voice preset.
    """
    import numpy as np
    
    # Smart split by sentence ending punctuation
    # Splits on: . ? ! or newline, keeping the punctuation
    sentences = re.split(r'([.?!]+|\n+)', prompt)
    
    # Recombine split text (sentence + punctuation)
    chunks = []
    current_chunk = ""
    
    for s in sentences:
        s = s.strip()
        if not s: continue
        
        # If it's punctuation, attach to previous
        if re.match(r'^[.?!]+$', s) or re.match(r'^\n+$', s):
            if chunks:
                chunks[-1] += s
            else:
                current_chunk += s # edge case: starts with punctuation
        else:
            # If current chunk is getting too long (Bark limit ~14s is roughly ~20-30 words depending on speed)
            # Heuristic: ~200 chars or ~30 words is safer upper bound
            if len(current_chunk) + len(s) > 200:
                chunks.append(current_chunk)
                current_chunk = s
            else:
                # Basic accumulation
                if current_chunk:
                    # If current didn't end with proper punctuation, add space
                    chunks.append(current_chunk)
                    current_chunk = s
                else:
                    current_chunk = s
    
    if current_chunk:
        chunks.append(current_chunk)
        
    print(f"   ✂️  Splitting long text into {len(chunks)} chunks for stable generation...")
    full_audio = []
    
    for i, text_chunk in enumerate(chunks):
        if not text_chunk.strip(): continue
        print(f"   ▶️  Generating chunk {i+1}/{len(chunks)}: '{text_chunk[:30]}...'")
        
        # Independent generation (Best stability)
        inputs = processor(text_chunk, voice_preset=voice_preset).to(device)
        audio_array = model.generate(**inputs, do_sample=True)
        audio_array = audio_array.cpu().numpy().squeeze()
        full_audio.append(audio_array)
        
        # Add a short silence between sentences for natural pacing (0.25s)
        silence_len = int(sample_rate * 0.25)
        full_audio.append(np.zeros(silence_len))

    # Concatenate all
    if not full_audio: return np.array([])
    return np.concatenate(full_audio)

def generate_audio(prompt, output_path, duration, sampling_rate, model_name="default", image_input=None, caption_model="florence", voice_preset="v2/en_speaker_6"):
    """Generate audio using MusicGen or AudioLDM (supports Image-to-Audio via captioning)."""
    
    model_id = AUDIO_MODELS.get(model_name.lower(), model_name)
    if model_name.lower() == "default": model_id = AUDIO_MODELS["default"]
    
    import sys
    sys.setrecursionlimit(50000) # Fix for Stable Audio / torchsde recursion on MPS
    
    device, dtype = get_optimal_device_and_dtype(quiet=True)
    
    # --- Image-to-Audio Logic (Captioning) ---
    if image_input:
        caption = generate_caption(image_input, device, model_type=caption_model)
        if caption:
            # If no user prompt is provided, use the caption directly
            if not prompt:
                print(f"   ℹ️  No prompt provided. Using generated caption as prompt.")
                full_prompt = caption
            else:
                # Combine User Prompt + Image Caption
                # Prompt is the "Action" or "Style", Caption is the "Content"
                full_prompt = f"{prompt}, inspired by {caption}"
                
            print(f"   Full Prompt: '{full_prompt}'")
            # Update prompt for downstream models
            prompt = full_prompt
        else:
            print(f"⚠️  Image analysis failed. Proceeding with text prompt only.")

    print(f"🎵 Generating Audio")
    print(f"   Model:    {model_id}")
    print(f"   Prompt:   '{prompt}'")
    if image_input: print(f"   Input Img: {image_input}")
    
    if "bark" in model_id.lower():
        print(f"   Duration: Auto (Text-based)")
    else:
        print(f"   Duration: {duration}s")
        
    print(f"   Sampling: {sampling_rate}Hz")
    print(f"   Output:   {output_path}")
    print("") # Spacer
    
    try:
        import torch
        import scipy.io.wavfile
        from transformers import pipeline
        from diffusers import AudioLDM2Pipeline
        
        # Logic for Different Model Types
        if "musicgen" in model_id.lower():
            # Use Transformers Pipeline
            print(f"   Loading MusicGen pipeline...")
            synthesizer = pipeline("text-to-audio", model_id, device=device)
            
            # MusicGen: ~50 tokens per sec estimate, strict max_new_tokens
            max_tokens = int(duration * 50) 
            print(f"🎵 Synthesizing audio... (MusicGen)")
            
            # --- Performance Tracking ---
            tracker = PerformanceTracker()

            start_time = time.time()
            with ResourceMonitor() as monitor:
                music = synthesizer(prompt, forward_params={"max_new_tokens": max_tokens})
            
            # Collect & Record
            gen_duration = time.time() - start_time
            avg_cpu, avg_ram, avg_vram, avg_gpu = monitor.get_averages()
            tracker.record_linear("audio", model_id, device, duration, gen_duration, cpu=avg_cpu, ram=avg_ram, vram=avg_vram, gpu=avg_gpu)
            print(f"   ✓ Generated in {format_time(gen_duration)} (RAM: {avg_ram:.1f}GB | VRAM: {avg_vram:.1f}GB | CPU: {avg_cpu:.1f}% | GPU: {avg_gpu:.1f}%)")
            
            if 'args' in globals() and hasattr(globals()['args'], 'report_json') and globals()['args'].report_json:
                 stats = {
                     "time": gen_duration,
                     "ram": avg_ram,
                     "vram": avg_vram,
                     "cpu": avg_cpu,
                     "gpu": avg_gpu
                 }
                 write_report_json(globals()['args'].report_json, stats)
            
            rate = music["sampling_rate"]
            audio_data = music["audio"]
            
            # Save raw WAV (transpose for scipy)
            scipy.io.wavfile.write(output_path + ".tmp.wav", rate, audio_data.T)
            src_path = output_path + ".tmp.wav"
            
        elif "audioldm" in model_id.lower():
            # Use Diffusers Pipeline for AudioLDM2
            # Workaround for transformers compatibility: explicit load of language_model
            from transformers import GPT2LMHeadModel
            print(f"   Loading AudioLDM2 pipeline components...")
            language_model = GPT2LMHeadModel.from_pretrained(model_id, subfolder="language_model").to(dtype=dtype)
            
            pipe = AudioLDM2Pipeline.from_pretrained(
                model_id, 
                language_model=language_model,
                torch_dtype=dtype
            )
            pipe.to(device)
            
            print(f"🎵 Synthesizing audio... (AudioLDM2)")
            
            # --- Performance Tracking ---
            tracker = PerformanceTracker()

            start_time = time.time()
            with ResourceMonitor() as monitor:
                audio = pipe(prompt, audio_length_in_s=duration).audios[0]
            
            # Collect & Record
            gen_duration = time.time() - start_time
            avg_cpu, avg_ram, avg_vram, avg_gpu = monitor.get_averages()
            tracker.record_linear("audio", model_id, device, duration, gen_duration, cpu=avg_cpu, ram=avg_ram, vram=avg_vram, gpu=avg_gpu)
            print(f"   ✓ Generated in {format_time(gen_duration)} (RAM: {avg_ram:.1f}GB | VRAM: {avg_vram:.1f}GB | CPU: {avg_cpu:.1f}% | GPU: {avg_gpu:.1f}%)")
            
            if 'args' in globals() and hasattr(globals()['args'], 'report_json') and globals()['args'].report_json:
                 stats = {
                     "time": gen_duration,
                     "ram": avg_ram,
                     "vram": avg_vram,
                     "cpu": avg_cpu,
                     "gpu": avg_gpu
                 }
                 write_report_json(globals()['args'].report_json, stats)
            rate = 16000 # AudioLDM default usually
            
            scipy.io.wavfile.write(output_path + ".tmp.wav", rate, audio.T)
            src_path = output_path + ".tmp.wav"
            
        elif "stable-audio" in model_id.lower():
            # Use Diffusers Pipeline for Stable Audio
            from diffusers import StableAudioPipeline, EulerDiscreteScheduler
            
            print(f"   Loading StableAudioPipeline...")
            pipe = StableAudioPipeline.from_pretrained(model_id, torch_dtype=dtype)
            pipe.to(device)
            
            # Switch scheduler to EulerDiscrete to avoid torchsde recursion error on MPS
            print(f"   ℹ️  Swapping scheduler to EulerDiscrete (MPS optimization)")
            pipe.scheduler = EulerDiscreteScheduler.from_config(pipe.scheduler.config)
            
            print(f"🎵 Synthesizing audio... (Stable Audio)")
            
            # --- Performance Tracking ---
            tracker = PerformanceTracker()

            start_time = time.time()
            with ResourceMonitor() as monitor:
                # Stable Audio takes 'audio_end_in_s'
                audio = pipe(prompt, audio_start_in_s=0.0, audio_end_in_s=duration, num_inference_steps=50).audios[0]

            # Collect & Record
            gen_duration = time.time() - start_time
            avg_cpu, avg_ram, avg_vram, avg_gpu = monitor.get_averages()
            tracker.record_linear("audio", model_id, device, duration, gen_duration, cpu=avg_cpu, ram=avg_ram, vram=avg_vram, gpu=avg_gpu)
            print(f"   ✓ Generated in {format_time(gen_duration)} (RAM: {avg_ram:.1f}GB | VRAM: {avg_vram:.1f}GB | CPU: {avg_cpu:.1f}% | GPU: {avg_gpu:.1f}%)")
            rate = 44100 # Standard for Stable Audio Open
            
            # Ensure numpy
            import torch
            if isinstance(audio, torch.Tensor):
                audio = audio.cpu().float().numpy()
            
            scipy.io.wavfile.write(output_path + ".tmp.wav", rate, audio.T)
            src_path = output_path + ".tmp.wav"
            
        elif "bark" in model_id.lower():
            # Use Transformers for Bark
            from transformers import BarkModel, AutoProcessor
            print(f"   Loading Bark models...")
            
            # Bark requires float32 on all platforms (float16 causes "Unsupported data type" errors)
            # This applies to both MPS and CUDA despite CUDA usually supporting float16
            bark_dtype = torch.float32
                
            processor = AutoProcessor.from_pretrained(model_id)
            model = BarkModel.from_pretrained(model_id, torch_dtype=bark_dtype).to(device)
            
            print(f"🎵 Synthesizing audio... (Bark)")
            if duration > 14:
                 print(f"   (Note: Bark generates max ~14s sequences per history block. Output will be shorter than {duration}s)")
            
            print(f"""   💡 Tip:
   *  Lyrics: Use '♪' for singing (e.g., `♪ Hello World ♪`).
   *  Effects: Use tags like `[laughter]`, `[cheers]`, `[music]`, `[sighs]`, `[gasps]`, `[clears throat]`, `—` (hesitation).
   *  Plain text without these tokens will usually be spoken as speech.
   *  Voice: Using preset '{voice_preset}'. Change with --voice-preset (e.g. 'v2/fr_speaker_1').
   *  Example: `python ai-media.py -a --audio-model bark -p "♪ Hello World ♪ [laughter]"`""")
            
            # Decide if Long-Form is needed
            # Heuristic: Text > 150 chars OR user requested > 15 seconds
            # Bark roughly does ~14s max.
            is_long = len(prompt) > 150 or duration > 15.0
            
            # --- Performance Tracking ---
            tracker = PerformanceTracker()
            
            start_time = time.time()
            with ResourceMonitor() as monitor:
                if is_long:
                    print(f"   📜 Long text detected. Using chunked generation.")
                    audio_array = generate_long_bark(prompt, processor, model, device, voice_preset)
                else:
                    # Use user-specified voice preset
                    inputs = processor(prompt, voice_preset=voice_preset).to(device)
                    # Bark output shape: [1, length]
                    audio_array = model.generate(**inputs) 
                    audio_array = audio_array.cpu().numpy().squeeze()
            
            # Collect & Record
            gen_duration = time.time() - start_time
            avg_cpu, avg_ram, avg_vram, avg_gpu = monitor.get_averages()
            tracker.record_linear("audio", model_id, device, duration, gen_duration, cpu=avg_cpu, ram=avg_ram, vram=avg_vram, gpu=avg_gpu)
            print(f"   ✓ Generated in {format_time(gen_duration)} (RAM: {avg_ram:.1f}GB | VRAM: {avg_vram:.1f}GB | CPU: {avg_cpu:.1f}% | GPU: {avg_gpu:.1f}%)")
            
            if 'args' in globals() and hasattr(globals()['args'], 'report_json') and globals()['args'].report_json:
                 stats = {
                     "time": gen_duration,
                     "ram": avg_ram,
                     "vram": avg_vram,
                     "cpu": avg_cpu,
                     "gpu": avg_gpu
                 }
                 write_report_json(globals()['args'].report_json, stats)
            
            rate = model.generation_config.sample_rate # 24000
            
            # Scipy write
            scipy.io.wavfile.write(output_path + ".tmp.wav", rate, audio_array)
            src_path = output_path + ".tmp.wav"
            
        else:
            print(f"❌ Unknown model type for audio: {model_id}")
            return False

        # Conversion / Final Save
        # If user asked for .mp3, convert. If .wav, rename.
        must_convert = not output_path.lower().endswith(".wav")
        
        if must_convert:
            import subprocess
            cmd = ["ffmpeg", "-y", "-i", src_path, output_path, "-loglevel", "error"]
            try:
                subprocess.run(cmd, check=True)
                os.remove(src_path)
                print(f"✅ Converted and saved to {output_path}")
            except subprocess.CalledProcessError:
                print(f"⚠️  FFmpeg conversion failed. Saved as WAV: {src_path}")
        else:
            os.rename(src_path, output_path)
            print(f"✅ Audio saved to {output_path}")
            
        return True
        
    except ImportError as e:
        print(f"❌ Error: Missing dependencies or import failed. {e}")
        return False
    except Exception as e:
        print(f"❌ Audio generation failed: {e}")
        return False


def get_video_encoding_params(output_path):
    """Get FFmpeg encoding parameters based on output file extension.
    
    Returns a list of FFmpeg arguments for video codec, pixel format, and audio codec.
    Supports: mp4, mkv, mov, webm, wmv, avi.
    Utilizes hardware acceleration (NVENC/VideoToolbox) if available.
    """
    import torch
    ext = os.path.splitext(output_path)[1].lower()
    
    # Platform detection
    has_cuda = torch.cuda.is_available()
    has_mps = torch.backends.mps.is_available()
    
    # 1. Video Codec Selection (Default to H.264 for widest compatibility)
    vcodec = "libx264"
    if ext in ['.mp4', '.m4v', '.mkv', '.mov']:
        if has_cuda:
            vcodec = "h264_nvenc"
        elif has_mps:
            vcodec = "h264_videotoolbox"
    elif ext == '.webm':
        vcodec = "libvpx-vp9"
    elif ext == '.wmv':
        vcodec = "wmv2"
    elif ext == '.avi':
        vcodec = "mpeg4"
        
    # 2. Audio Codec Selection
    acodec = "aac"
    if ext == '.webm':
        acodec = "libopus"
    elif ext == '.wmv':
        acodec = "wmav2"
    elif ext == '.avi':
        acodec = "mp3"
        
    # 3. Parameters
    params = ["-c:v", vcodec, "-pix_fmt", "yuv420p", "-c:a", acodec]
    
    # Add bitrate for less efficient formats
    if ext in ['.webm', '.wmv', '.avi']:
        params.extend(["-b:v", "2M"])
        
    return params


def ffmpeg_resize_video(input_path, output_path, target_w, target_h):
    """Resize video to exact target dimensions using FFmpeg Lanczos.
    
    Used as a final step when AI upscalers produce dimensions that don't match
    the target exactly (e.g., Real-ESRGAN's fixed 4x scale).
    
    Args:
        input_path: Path to input video
        output_path: Path for output video  
        target_w: Target width (will be made even for codec compatibility)
        target_h: Target height (will be made even for codec compatibility)
        
    Returns:
        True on success, False on failure
    """
    # Ensure dimensions are even (required by most codecs)
    target_w = target_w if target_w % 2 == 0 else target_w + 1
    target_h = target_h if target_h % 2 == 0 else target_h + 1
    
    print(f"   📐 FFmpeg resize: {target_w}x{target_h}")
    
    try:
        cmd = [
            "ffmpeg", "-y", "-i", input_path,
            "-vf", f"scale={target_w}:{target_h}:flags=lanczos",
            "-c:a", "copy",  # Copy audio stream unchanged
            output_path
        ]
        subprocess.run(cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        return True
    except Exception as e:
        print(f"   ⚠️  FFmpeg resize failed: {e}")
        return False


def upscale_video_zeroscope_xl(video_frames, prompt, device=None, dtype=None, strength=0.6):
    """Upscale video frames using zeroscope_v2_XL (Video-to-Video diffusion).
    
    This function takes video frames from zeroscope_v2_576w (576x320) and upscales
    them to 1024x576 using the XL model's video-to-video pipeline with temporal
    consistency.
    
    Args:
        video_frames: List of PIL Images or numpy arrays from 576w generation
        prompt: The original text prompt (used for guidance during upscaling)
        device: torch device (auto-detected if None)
        dtype: torch dtype (auto-detected if None)
        strength: Denoise strength 0.0-1.0. Higher = more creative, lower = faithful.
                  Default 0.6 as recommended by model docs.
                  
    Returns:
        List of upscaled PIL Images at 1024x576, or None on failure
    """
    print(f"\n🔄 Upscaling with Zeroscope XL (Video-to-Video)...")
    print(f"   Input:    576x320 ({len(video_frames)} frames)")
    print(f"   Output:   1024x576")
    print(f"   Strength: {strength}")
    
    try:
        import torch
        from diffusers import DiffusionPipeline, DPMSolverMultistepScheduler
        from PIL import Image
        import numpy as np
        
        # Auto-detect device/dtype if not provided
        if device is None or dtype is None:
            device, dtype = get_optimal_device_and_dtype(quiet=True)
        
        # MPS limitation: Force CPU for XL V2V upscale
        # MPS fails with "tensor dims larger than INT_MAX" on large video tensors
        original_device = device
        if device.type == "mps":
            print("   ⚠️  MPS Limitation: Forcing CPU for XL V2V upscale (MPS tensor limit).")
            device = torch.device("cpu")
            dtype = torch.float32
        
        # Load XL model
        xl_model_id = VIDEO_MODELS.get("zeroscope-xl", "cerspense/zeroscope_v2_XL")
        print(f"   Model:    {xl_model_id}")
        print(f"   Device:   {device}")
        
        pipe = DiffusionPipeline.from_pretrained(xl_model_id, torch_dtype=dtype)
        pipe.scheduler = DPMSolverMultistepScheduler.from_config(pipe.scheduler.config)
        
        # Memory optimization
        pipe.to(device)
        if device.type == "cpu":
            pipe.enable_attention_slicing()
        
        pipe.enable_vae_slicing()
        
        # Prepare input frames - resize to 1024x576 for XL input
        print(f"   Preparing frames...")
        resized_frames = []
        for frame in video_frames:
            if isinstance(frame, np.ndarray):
                # Handle float32/float64 arrays (0-1 range) - convert to uint8 (0-255)
                if frame.dtype in [np.float32, np.float64]:
                    frame = (frame * 255).clip(0, 255).astype(np.uint8)
                frame = Image.fromarray(frame)
            # Resize to XL's native resolution
            resized_frame = frame.resize((1024, 576), Image.Resampling.LANCZOS)
            resized_frames.append(resized_frame)
        
        # Run XL upscaling (Video-to-Video)
        print(f"   Running diffusion upscale... (This may take a while)")
        start_time = time.time()
        
        with ResourceMonitor() as monitor:
            output = pipe(
                prompt=prompt,
                video=resized_frames,
                strength=strength,
                num_inference_steps=25
            )
        
        duration = time.time() - start_time
        avg_cpu, avg_ram, avg_vram, avg_gpu = monitor.get_averages()
        print(f"   ✓ XL Upscale complete in {format_time(duration)} (RAM: {avg_ram:.1f}GB | VRAM: {avg_vram:.1f}GB)")
        
        # Cleanup XL pipeline
        del pipe
        clear_gpu_memory()
        
        # Extract output frames
        upscaled_frames = output.frames[0] if hasattr(output, 'frames') else output.images
        return upscaled_frames
        
    except Exception as e:
        print(f"   ❌ XL Upscale failed: {e}")
        import traceback
        traceback.print_exc()
        return None


def generate_video(prompt, output_path, duration, width, height, model_name="default", image_input=None, audio_prompt=None, audio_model="default"):
    """Generate video (Text-to-Video or Image-to-Video) with optional Audio.
    
    For Zeroscope: Implements dynamic upscaling pipeline when target resolution
    exceeds native 576x320. Uses zeroscope_v2_XL for V2V upscaling to 1024x576,
    then Real-ESRGAN for higher resolutions.
    """
    
    # Resolve Model ID
    base_model = VIDEO_MODELS.get(model_name.lower(), model_name)
    if model_name.lower() == "default": base_model = VIDEO_MODELS["default"]
    
    # --- Zeroscope Dynamic Upscaling Detection ---
    # Native resolution for zeroscope is 576x320
    # If target is higher, we generate at native and upscale dynamically
    is_zeroscope = "zeroscope" in base_model.lower() and "xl" not in base_model.lower()
    zeroscope_native_w, zeroscope_native_h = 576, 320
    zeroscope_xl_w, zeroscope_xl_h = 1024, 576
    
    needs_xl_upscale = False
    needs_esrgan_upscale = False
    target_width, target_height = width, height
    gen_width, gen_height = width, height
    
    if is_zeroscope and not image_input:  # Only for T2V zeroscope
        if width > zeroscope_native_w or height > zeroscope_native_h:
            # Generate at native resolution, upscale later
            gen_width, gen_height = zeroscope_native_w, zeroscope_native_h
            
            # Check if we're on MPS - skip XL (CPU diffusion is impractical)
            import torch
            is_mps = torch.backends.mps.is_available() and not torch.cuda.is_available()
            
            if is_mps:
                # MPS: Skip XL entirely, go straight to Real-ESRGAN
                needs_xl_upscale = False
                needs_esrgan_upscale = True
                print(f"📐 Dynamic Upscaling Pipeline (MPS Optimized):")
                print(f"   ⚠️  Skipping XL V2V (CPU diffusion too slow on Apple Silicon)")
                print(f"   Target:  {target_width}x{target_height}")
                print(f"   Step 1:  Generate at {gen_width}x{gen_height} (Zeroscope native)")
                print(f"   Step 2:  Real-ESRGAN to ~{target_width}x{target_height}")
                print(f"   Step 3:  FFmpeg resize to exact {target_width}x{target_height}")
            else:
                # CUDA/CPU: Use full pipeline with XL
                needs_xl_upscale = True
                if width > zeroscope_xl_w or height > zeroscope_xl_h:
                    needs_esrgan_upscale = True
                    
                print(f"📐 Dynamic Upscaling Pipeline Activated:")
                print(f"   Target:  {target_width}x{target_height}")
                print(f"   Step 1:  Generate at {gen_width}x{gen_height} (Zeroscope native)")
                if needs_xl_upscale:
                    print(f"   Step 2:  XL Upscale to {zeroscope_xl_w}x{zeroscope_xl_h} (V2V diffusion)")
                if needs_esrgan_upscale:
                    print(f"   Step 3:  Real-ESRGAN to ~{target_width}x{target_height}")
                    print(f"   Step 4:  FFmpeg resize to exact {target_width}x{target_height}")
            print("")
    
    # Handle Image-to-Video Logic
    is_i2v = True if image_input else False
    
    if is_i2v:
        # Check if we need to switch to an I2V variant
        if "cogvideox" in base_model.lower():
            model_id = "THUDM/CogVideoX-5b-I2V"
        elif "wan2.2" in base_model.lower() or "wan2.2" in model_name.lower():
            model_id = "Wan-AI/Wan2.2-I2V-A14B-Diffusers"
        elif "hunyuan" in base_model.lower() or "hunyuan" in model_name.lower():
            model_id = "hunyuanvideo-community/HunyuanVideo-I2V"
        elif "stable-video-diffusion" in base_model.lower() or model_name.lower() == "svd":
            model_id = "stabilityai/stable-video-diffusion-img2vid-xt"
        else:
            print(f"⚠️  Warning: Model '{model_name}' ({base_model}) may not support Image-to-Video.")
            print(f"   Switching to 'svd' (Stable Video Diffusion) as fallback.")
            model_id = "stabilityai/stable-video-diffusion-img2vid-xt"
    else:
        # Text-to-Video
        model_id = base_model

    print(f"{'='*60}")
    if needs_xl_upscale:
        print(f"📐 Step 1: Generate at {gen_width}x{gen_height} (Zeroscope native)")
    else:
        print(f"🎬 Generating Video ({'Image-to-Video' if is_i2v else 'Text-to-Video'})")
    print(f"{'='*60}")
    print(f"   Model:    {model_id}")
    print(f"   Prompt:   '{prompt}'")
    if is_i2v: print(f"   Input Img: {image_input}")
    if audio_prompt: print(f"   Audio:    '{audio_prompt}' (Will generate and mux)")
    print(f"   Duration: {duration}s")
    print("") # Spacer
    
    # Determine actual video output path (temp if mixing audio)
    video_out = output_path
    if audio_prompt:
        video_out = output_path + ".temp_video.mp4"
        audio_out = output_path + ".temp_audio.wav"
        # Clean up any leftover temp files from previous runs
        for temp_file in [video_out, audio_out, audio_out + ".tmp.wav"]:
            if os.path.exists(temp_file):
                try:
                    os.remove(temp_file)
                except OSError:
                    pass  # Ignore if can't delete
    
    # Ensure a clean slate
    clear_gpu_memory()
    pipe = None
    cuda_was_enabled = None
    
    try:
        import torch
        from diffusers import (
            DiffusionPipeline, 
            DPMSolverMultistepScheduler, 
            CogVideoXImageToVideoPipeline,
            StableVideoDiffusionPipeline,
            WanPipeline,
            WanImageToVideoPipeline,
            HunyuanVideoPipeline,
            HunyuanVideoImageToVideoPipeline,
            utils
        )
        from diffusers.utils import export_to_video, load_image
        
        device, dtype = get_optimal_device_and_dtype(quiet=True)
        
        # MPS FIX: These models need Float32/CPU on MPS
        # - Text-to-Video models: Float16 corrupts output
        # - SVD: 3D convolutions cause "Invalid buffer size" errors on MPS
        mps_incompatible_models = ["ms-1.7b", "text-to-video-ms-1.7b", "zeroscope", "stable-video-diffusion", "cogvideox"]
        is_mps_incompatible = any(m in model_id.lower() for m in mps_incompatible_models)
        
        if device.type == "mps" and is_mps_incompatible:
            # SVD needs CPU entirely (3D convolutions fail on MPS)
            if "stable-video-diffusion" in model_id.lower():
                print("⚠️  MPS Compatibility: SVD requires CPU on Apple Silicon.")
                print("   (3D convolutions cause 'Invalid buffer size' on MPS)")
                device = torch.device("cpu")
            else:
                print("⚠️  MPS Compatibility: Using Float32 for correct video output.")
            dtype = torch.float32
        
        print("⚠️  Video generation is resource intensive.")
        
        # cuDNN workaround: Disable cuDNN for video generation to prevent "GET was unable to find an engine" errors
        # This is more reliable but slightly slower. The error occurs during VAE decoding with certain CUDA versions.
        if torch.cuda.is_available():
            cuda_was_enabled = torch.backends.cudnn.enabled
            torch.backends.cudnn.enabled = False
        
        # --- Stage 1: Video Generation ---
        
        # Load Pipeline
        if "cogvideox" in model_id.lower() and is_i2v:
            pipe = CogVideoXImageToVideoPipeline.from_pretrained(model_id, torch_dtype=dtype)
            
            print(f"   ℹ️  Applying Memory Optimizations for CogVideoX...")
            pipe.enable_sequential_cpu_offload() 
            pipe.vae.enable_tiling()
            pipe.vae.enable_slicing()

        elif "wan2.2" in model_id.lower():
            # Wan 2.2 (Alibaba)
            if is_i2v:
                print(f"   ℹ️  Loading Wan 2.2 Image-to-Video Pipeline...")
                pipe = WanImageToVideoPipeline.from_pretrained(model_id, torch_dtype=dtype)
            else:
                print(f"   ℹ️  Loading Wan 2.2 Text-to-Video Pipeline...")
                pipe = WanPipeline.from_pretrained(model_id, torch_dtype=dtype)
            
            # Massive model optimizations (14B params)
            # MPS: Use sequential CPU offload for memory-constrained unified memory
            if device.type == "mps":
                print("   ℹ️  MPS: Enabling Sequential CPU Offload for Wan 2.2 (memory-safe)...")
                pipe.enable_sequential_cpu_offload()
            else:
                print("   ℹ️  Enabling Model CPU Offload for Wan 2.2...")
                pipe.enable_model_cpu_offload()
            pipe.vae.enable_tiling()
            
        elif "ltx-video" in model_id.lower():
            # LTX-Video (Lightricks)
            from diffusers import LTXPipeline
            print(f"   ℹ️  Loading LTX-Video Pipeline...")
            pipe = LTXPipeline.from_pretrained(model_id, torch_dtype=dtype)
            
            # LTX Optimizations (Fast DiT)
            # Default to CPU offload to be safe on consumer GPUs
            pipe.enable_model_cpu_offload()
            pipe.vae.enable_tiling()

        elif "mochi-1" in model_id.lower():
            # Mochi 1 (Genmo)
            from diffusers import MochiPipeline
            print(f"   ℹ️  Loading Mochi 1 Pipeline...")
            pipe = MochiPipeline.from_pretrained(model_id, torch_dtype=dtype)
            
            # Very heavy model (10B params), needs heavy offloading
            # MPS: Use sequential CPU offload for memory-constrained unified memory
            if device.type == "mps":
                print("   ℹ️  MPS: Enabling Sequential CPU Offload for Mochi 1 (memory-safe)...")
                pipe.enable_sequential_cpu_offload()
            else:
                print("   ℹ️  Enabling Model CPU Offload for Mochi 1...")
                pipe.enable_model_cpu_offload()
            pipe.vae.enable_tiling()

        elif "hunyuan" in model_id.lower():
            # HunyuanVideo (Tencent)
            if is_i2v:
                print(f"   ℹ️  Loading HunyuanVideo Image-to-Video Pipeline...")
                pipe = HunyuanVideoImageToVideoPipeline.from_pretrained(model_id, torch_dtype=dtype)
            else:
                print(f"   ℹ️  Loading HunyuanVideo Text-to-Video Pipeline...")
                pipe = HunyuanVideoPipeline.from_pretrained(model_id, torch_dtype=dtype)
            
            # Massive model (13B), similar to Wan 2.2
            # MPS: Use sequential CPU offload for memory-constrained unified memory
            if device.type == "mps":
                print("   ℹ️  MPS: Enabling Sequential CPU Offload for HunyuanVideo (memory-safe)...")
                pipe.enable_sequential_cpu_offload()
                # MPS Optimization: Force VAE to float16 to save memory (even if unstable? worth a try)
                # and enable slicing + tiling
                try:
                    pipe.vae.enable_tiling()
                    pipe.vae.enable_slicing()
                    print("   ℹ️  Enabled VAE Tiling & Slicing for HunyuanVideo")
                except Exception as e:
                    print(f"   ⚠️  Could not enable VAE optimizations: {e}")
            else:
                print("   ℹ️  Enabling Model CPU Offload for HunyuanVideo...")
                pipe.enable_model_cpu_offload()
                pipe.vae.enable_tiling()

        elif "stable-video-diffusion" in model_id.lower():
            pipe = StableVideoDiffusionPipeline.from_pretrained(model_id, torch_dtype=dtype, variant="fp16" if dtype == torch.float16 else None)
        else:
            # Generic / Text-to-Video
            try:
                if dtype == torch.float16:
                    pipe = DiffusionPipeline.from_pretrained(model_id, torch_dtype=dtype, variant="fp16")
                else:
                    pipe = DiffusionPipeline.from_pretrained(model_id, torch_dtype=dtype)
            except Exception:
                pipe = DiffusionPipeline.from_pretrained(model_id, torch_dtype=dtype)
        
        # Scheduler Optimization (Skip if handled or problematic)
        if hasattr(pipe, "scheduler"):
            # SVD, Mochi, and LTX may have specific schedulers they prefer
            is_sensitive_scheduler = any(x in model_id.lower() for x in ["stable-video-diffusion", "mochi", "ltx", "wan", "hunyuan"])
            if not is_sensitive_scheduler:
                try:
                    pipe.scheduler = DPMSolverMultistepScheduler.from_config(pipe.scheduler.config)
                except: pass 
        
        # Device/Memory Optimization
        # (CogVideoX and new models handled specifically above)
        special_handling = ["cogvideox", "wan", "ltx", "mochi", "hunyuan"]
        if not any(x in model_id.lower() for x in special_handling):
            if device.type == "cpu":
                pipe.to(device)
            else:
                pipe.enable_model_cpu_offload()
                if device.type == "mps":
                    pipe.enable_attention_slicing()
        
        # Generate Frames
        # --- Performance Tracking ---
        tracker = PerformanceTracker()
        
        # Fix dimensions for specific models if needed (some enforce strict aspect ratios or resolution buckets)
        render_width, render_height = gen_width, gen_height
        
        # LTX-Video prefers multiples of 32
        if "ltx-video" in model_id.lower():
             render_width = (render_width // 32) * 32
             render_height = (render_height // 32) * 32
        # Mochi prefers multiples of 16 (often 848x480 native)
        elif "mochi" in model_id.lower():
             render_width = (render_width // 16) * 16
             render_height = (render_height // 16) * 16
        # Wan/Hunyuan have flexible bucket resolution but benefit from standard ratios (1280x720)
        
        print(f"🎬 Rendering video frames at {render_width}x{render_height}... (This might be slow)")
        
        start_time = time.time()
        with ResourceMonitor() as monitor:
            if is_i2v:
                init_image = load_image(image_input)
                init_image = init_image.resize((gen_width, gen_height))
                
                if "stable-video-diffusion" in model_id.lower():
                    video_frames = pipe(init_image).frames[0]
                elif "wan2.2" in model_id.lower():
                    # Wan 2.2 I2V
                    video_frames = pipe(prompt=prompt, image=init_image, num_frames=81, num_inference_steps=50).frames[0]
                elif "hunyuan" in model_id.lower():
                    # HunyuanVideo I2V
                    video_frames = pipe(prompt=prompt, image=init_image, num_frames=61, num_inference_steps=50).frames[0]
                else:
                    # CogVideoX
                    video_frames = pipe(prompt=prompt, image=init_image, num_frames=49, guidance_scale=6.0, num_inference_steps=50).frames[0]
            else:
                num_frames = int(duration * 16)
                video_frames = pipe(prompt, num_inference_steps=25, num_frames=num_frames).frames[0]
        
        # Collect & Record
        gen_duration = time.time() - start_time
        avg_cpu, avg_ram, avg_vram, avg_gpu = monitor.get_averages()
        tracker.record_linear("video", model_id, device, duration, gen_duration, gen_width, gen_height, cpu=avg_cpu, ram=avg_ram, vram=avg_vram, gpu=avg_gpu)
        print(f"   ✓ Rendered in {format_time(gen_duration)} (RAM: {avg_ram:.1f}GB | VRAM: {avg_vram:.1f}GB | CPU: {avg_cpu:.1f}% | GPU: {avg_gpu:.1f}%)")
        
        # --- Zeroscope Dynamic Upscaling Pipeline ---
        if needs_xl_upscale:
            # Cleanup base pipeline before loading XL
            del pipe
            pipe = None
            clear_gpu_memory()
            
            # Step 2: XL Upscale (576x320 -> 1024x576)
            print(f"\n{'='*60}")
            print(f"📐 Step 2: XL Upscale to {zeroscope_xl_w}x{zeroscope_xl_h} (V2V diffusion)")
            print(f"{'='*60}")
            
            upscaled_frames = upscale_video_zeroscope_xl(
                video_frames, 
                prompt, 
                device=device, 
                dtype=dtype,
                strength=0.6
            )
            
            if upscaled_frames is not None:
                video_frames = upscaled_frames
                gen_width, gen_height = zeroscope_xl_w, zeroscope_xl_h
            else:
                print("   ⚠️  XL upscale failed, continuing with base resolution")
                needs_esrgan_upscale = False  # Skip further upscaling
        
        if 'args' in globals() and hasattr(globals()['args'], 'report_json') and globals()['args'].report_json:
                stats = {
                    "time": gen_duration,
                    "ram": avg_ram,
                    "vram": avg_vram,
                    "cpu": avg_cpu,
                    "gpu": avg_gpu,
                    "width": target_width,
                    "height": target_height
                }
                write_report_json(globals()['args'].report_json, stats)
        
        # Save Video (raw export - may need re-encoding for compatibility)
        temp_raw_video = video_out + ".raw.mp4"
        export_to_video(video_frames, temp_raw_video, fps=7 if "stable-video-diffusion" in model_id.lower() else 8) # CogVideoX is usually 8fps
        
        # Re-encode with FFmpeg for universal playback (format-aware)
        import subprocess
        encoding_params = get_video_encoding_params(video_out)
        try:
            subprocess.run([
                "ffmpeg", "-y", "-i", temp_raw_video,
                *encoding_params,
                video_out
            ], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            os.remove(temp_raw_video)
            print(f"✅ Video saved to {video_out} ({gen_width}x{gen_height})")
        except Exception as e:
            # Fallback: use raw video if FFmpeg fails
            os.rename(temp_raw_video, video_out)
            print(f"⚠️  Video saved (may require VLC to play): {video_out}")
        
        # Step 3 & 4: Real-ESRGAN + FFmpeg resize for target > 1024x576
        if needs_esrgan_upscale:
            print(f"\n{'='*60}")
            print(f"📐 Step 3: Real-ESRGAN Upscale to ~{target_width}x{target_height}")
            print(f"{'='*60}")
            
            # Calculate upscale factor from 1024x576 to target
            # Real-ESRGAN uses 4x model, so we may need FFmpeg for final sizing
            esrgan_factor = max(target_width / zeroscope_xl_w, target_height / zeroscope_xl_h)
            esrgan_factor = min(esrgan_factor, 4.0)  # Cap at 4x for single pass
            
            temp_esrgan_input = video_out
            temp_esrgan_output = video_out + ".esrgan.mp4"
            
            # Run Real-ESRGAN upscale
            esrgan_success = upscale_video_fast(
                temp_esrgan_input, 
                temp_esrgan_output, 
                factor=esrgan_factor
            )
            
            if esrgan_success and os.path.exists(temp_esrgan_output):
                # Check if we need FFmpeg resize to exact target
                # (Real-ESRGAN produces dimensions that are multiples of its scale)
                import cv2
                cap = cv2.VideoCapture(temp_esrgan_output)
                esrgan_w = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
                esrgan_h = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
                cap.release()
                
                if esrgan_w != target_width or esrgan_h != target_height:
                    # Step 4: FFmpeg resize to exact dimensions
                    print(f"\n{'='*60}")
                    print(f"📐 Step 4: FFmpeg resize to exact {target_width}x{target_height}")
                    print(f"{'='*60}")
                    
                    temp_final = video_out + ".final.mp4"
                    if ffmpeg_resize_video(temp_esrgan_output, temp_final, target_width, target_height):
                        os.remove(temp_esrgan_output)
                        os.remove(video_out)  # Remove intermediate XL output
                        os.rename(temp_final, video_out)
                        print(f"✅ Final video: {video_out} ({target_width}x{target_height})")
                    else:
                        # FFmpeg resize failed, use ESRGAN output as-is
                        os.remove(video_out)
                        os.rename(temp_esrgan_output, video_out)
                        print(f"✅ Video saved: {video_out} ({esrgan_w}x{esrgan_h})")
                else:
                    # ESRGAN output matches target
                    os.remove(video_out)
                    os.rename(temp_esrgan_output, video_out)
                    print(f"✅ Final video: {video_out} ({target_width}x{target_height})")
            else:
                print(f"   ⚠️  Real-ESRGAN failed, keeping XL output at {gen_width}x{gen_height}")
        
        # --- Stage 2 & 3: Audio Generation & Muxing ---
        
        # FORCE CLEANUP of Video Model before Audio Model loads
        del pipe
        pipe = None
        clear_gpu_memory()
        
        if audio_prompt:
            print("🔊 Generating Audio track...")
            audio_out = output_path + ".temp_audio.wav"
            # Use default music model or let user specify? Use default for now or expose args if needed.
            # We reuse generate_audio function!
            # Note: We need to pass sampling rate, let's default to standard 32k or use global default
            audio_success = generate_audio(audio_prompt, audio_out, duration, 32000, model_name=audio_model)
            
            if audio_success:
                print("🔗 Muxing Video and Audio...")
                # FFmpeg merge
                # -c:v copy (copy video stream)
                # -c:a aac (encode audio to aac for mp4 compatibility)
                # -shortest (cut to shortest stream, usually video)
                import subprocess
                cmd = [
                    "ffmpeg", "-y",
                    "-i", video_out,
                    "-i", audio_out,
                    "-c:v", "copy",
                    "-c:a", "aac",
                    "-shortest",
                    output_path,
                    "-loglevel", "error"
                ]
                try:
                    subprocess.run(cmd, check=True)
                    print(f"✅ Final merged video saved to {output_path}")
                    # Cleanup temp files after successful mux
                    for temp_file in [video_out, audio_out]:
                        if os.path.exists(temp_file):
                            try: os.remove(temp_file)
                            except: pass
                except subprocess.CalledProcessError:
                    print(f"❌ Muxing failed. Check FFmpeg.")
            else:
                print("❌ Audio generation failed. Returning silent video (renaming temp).")
                os.rename(video_out, output_path)
                # Cleanup audio temp if it exists
                if os.path.exists(audio_out):
                    try: os.remove(audio_out)
                    except: pass
                
        return True
        
    except Exception as e:
        print(f"❌ Video generation failed: {e}")
        # Clean temp if exists
        if audio_prompt and os.path.exists(video_out):
            try: os.remove(video_out)
            except: pass
        return False
        
    finally:
        # --- Final Cleanup (Crucial for VRAM) ---
        if pipe is not None:
            del pipe
        
        if cuda_was_enabled is not None:
             torch.backends.cudnn.enabled = cuda_was_enabled
             
        clear_gpu_memory()


def ensure_paths(output_path):
    """Create parent directories if needed."""
    if output_path:
        p = Path(output_path)
        if p.parent:
            p.parent.mkdir(parents=True, exist_ok=True)


# --- Performance Tracking ---

class PerformanceTracker:
    def __init__(self, filepath="performance.json"):
        self.filepath = filepath
        self.data = self._load()

    def _load(self):
        if os.path.exists(self.filepath):
            try:
                with open(self.filepath, 'r') as f:
                    return json.load(f)
            except:
                return {}
        return {}

    def _save(self):
        with open(self.filepath, 'w') as f:
            json.dump(self.data, f, indent=2)

    def record_image(self, model, width, height, device, time_taken, cpu=0, ram=0, vram=0, gpu=0):
        dev_str = device.type if hasattr(device, 'type') else str(device)
        key = f"{model}|{dev_str}|{width}x{height}"
        if "image" not in self.data: self.data["image"] = {}
        
        # New Logic: Keep only average_time
        # Re-average strategy: (last_average + new_time) / 2
        entry = self.data["image"].get(key, {})
        
        if "average_time" in entry:
            new_avg = (entry["average_time"] + time_taken) / 2.0
            # Reconstruct to ensure order: Time -> RAM -> VRAM -> CPU -> GPU
            entry = {
                "average_time": new_avg,
                "average_ram": (entry.get("average_ram", ram) + ram) / 2.0,
                "average_vram": (entry.get("average_vram", vram) + vram) / 2.0,
                "average_cpu": (entry.get("average_cpu", cpu) + cpu) / 2.0,
                "average_gpu": (entry.get("average_gpu", gpu) + gpu) / 2.0
            }
        else:
            entry = {
                "average_time": time_taken,
                "average_ram": ram,
                "average_vram": vram,
                "average_cpu": cpu,
                "average_gpu": gpu
            }
            
        self.data["image"][key] = entry
        self._save()

    def record_linear(self, category, model, device, duration, time_taken, width=None, height=None, cpu=0, ram=0, vram=0, gpu=0):
        """Record Audio/Video generation using rolling average rate (seconds to gen / seconds of content)."""
        dev_str = device.type if hasattr(device, 'type') else str(device)
        # For video, resolution also matters, so we include it in key
        if category == "video":
            key = f"{model}|{dev_str}|{width}x{height}"
        else:
            key = f"{model}|{dev_str}"
            
        if category not in self.data: self.data[category] = {}
        
        current_rate = time_taken / duration if duration > 0 else 0
        entry = self.data[category].get(key, {})
        
        if "average_rate" in entry:
            new_rate = (entry["average_rate"] + current_rate) / 2.0
            # Reconstruct to ensure order: Rate -> RAM -> VRAM -> CPU -> GPU
            entry = {
                "average_rate": new_rate,
                "average_ram": (entry.get("average_ram", ram) + ram) / 2.0,
                "average_vram": (entry.get("average_vram", vram) + vram) / 2.0,
                "average_cpu": (entry.get("average_cpu", cpu) + cpu) / 2.0,
                "average_gpu": (entry.get("average_gpu", gpu) + gpu) / 2.0
            }
        else:
            entry = {
                "average_rate": current_rate, 
                "average_ram": ram, 
                "average_vram": vram, 
                "average_cpu": cpu, 
                "average_gpu": gpu
            }
            
        self.data[category][key] = entry
        self._save()

    def estimate_image(self, model, width, height, device):
        dev_str = device.type if hasattr(device, 'type') else str(device)
        key = f"{model}|{dev_str}|{width}x{height}"
        stats = self.data.get("image", {}).get(key)
        if stats and "average_time" in stats:
            return stats["average_time"], stats.get("average_cpu", 0), stats.get("average_ram", 0), stats.get("average_vram", 0), stats.get("average_gpu", 0)
        return 0, 0, 0, 0, 0

    def estimate_linear(self, category, model, device, duration, width=None, height=None):
        dev_str = device.type if hasattr(device, 'type') else str(device)
        if category == "video":
            key = f"{model}|{dev_str}|{width}x{height}"
        else:
            key = f"{model}|{dev_str}"
            
        stats = self.data.get(category, {}).get(key)
        if stats and "average_rate" in stats:
            return stats["average_rate"] * duration, stats.get("average_cpu", 0), stats.get("average_ram", 0), stats.get("average_vram", 0), stats.get("average_gpu", 0)
        return 0, 0, 0, 0, 0

class ResourceMonitor:
    """Monitors CPU, RAM, and GPU VRAM/Load usage in a background thread."""
    def __init__(self, interval=0.5):
        self.interval = interval
        self.running = False
        self.thread = None
        self.cpu_readings = []
        self.ram_readings = []
        self.vram_readings = []
        self.gpu_readings = [] # GPU Load %
        
        try:
            import psutil
            self.psutil = psutil
        except ImportError:
            self.psutil = None
            print("⚠️  'psutil' not found. Resource monitoring disabled.")
            
        # Check for torch to monitor VRAM
        try:
            import torch
            self.torch = torch
            self.has_cuda = torch.cuda.is_available()
            self.has_mps = hasattr(torch.backends, 'mps') and torch.backends.mps.is_available()
        except ImportError:
            self.torch = None
            self.has_cuda = False
            self.has_mps = False

    def _monitor(self):
        import time
        
        while self.running:
            if self.psutil:
                cpu = self.psutil.cpu_percent(interval=None)
                ram = self.psutil.virtual_memory().used / (1024**3) # GB
                self.cpu_readings.append(cpu)
                self.ram_readings.append(ram)
                
            # VRAM Monitoring
            vram = 0
            if self.torch:
                if self.has_cuda:
                    vram = self.torch.cuda.memory_allocated() / (1024**3) # GB
                elif self.has_mps:
                     if hasattr(self.torch, 'mps') and hasattr(self.torch.mps, 'current_allocated_memory'):
                         vram = self.torch.mps.current_allocated_memory() / (1024**3)
                     elif hasattr(self.torch.mps, 'driver_allocated_memory'):
                         vram = self.torch.mps.driver_allocated_memory() / (1024**3)
            self.vram_readings.append(vram)
            
            # GPU Load Monitoring
            gpu_load = 0
            if self.has_cuda:
                try:
                    # Windows/Linux with NVIDIA drivers
                    result = subprocess.run(
                        ['nvidia-smi', '--query-gpu=utilization.gpu', '--format=csv,noheader,nounits'],
                        capture_output=True, text=True, check=False, timeout=1.0
                    )
                    if result.returncode == 0:
                        gpu_load = float(result.stdout.strip())
                except Exception:
                    pass
            elif self.has_mps:
                try:
                    # Apple Silicon: Query AGXAccelerator via ioreg for GPU utilization
                    import re
                    # Create a sanitized environment for the subprocess to avoid inheritance 
                    # of MallocStackLogging state which causes warnings on macOS.
                    env = os.environ.copy()
                    env["MallocStackLogging"] = "0" 
                    
                    result = subprocess.run(
                        ['ioreg', '-r', '-d', '1', '-w', '0', '-c', 'AGXAccelerator'],
                        capture_output=True, text=True, check=False,
                        env=env, timeout=1.0
                    )
                    if result.returncode == 0:
                        # Extract "Device Utilization %" from PerformanceStatistics
                        match = re.search(r'"Device Utilization %"=(\d+)', result.stdout)
                        if match:
                            gpu_load = float(match.group(1))
                except Exception:
                    pass
            self.gpu_readings.append(gpu_load)
            
            time.sleep(self.interval)

    def __enter__(self):
        if self.psutil:
            self.psutil.cpu_percent(interval=None) # Prime CPU
            self.running = True
            import threading
            self.thread = threading.Thread(target=self._monitor, daemon=True)
            self.thread.start()
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.running = False
        if self.thread:
            self.thread.join(timeout=1.0)
            
    def get_averages(self):
        avg_cpu = sum(self.cpu_readings) / len(self.cpu_readings) if self.cpu_readings else 0
        avg_ram = sum(self.ram_readings) / len(self.ram_readings) if self.ram_readings else 0
        avg_vram = sum(self.vram_readings) / len(self.vram_readings) if self.vram_readings else 0
        avg_gpu = sum(self.gpu_readings) / len(self.gpu_readings) if self.gpu_readings else 0
        return avg_cpu, avg_ram, avg_vram, avg_gpu

# --- Upscaling Logic ---

# --- Upscaling Logic ---

def check_resources_and_confirm(w, h, f, dev):
    """
    Checks if the target upscale resolution is safe for the current system resources.
    Returns True if safe/confirmed, False if user aborts.
    """
    try:
        import psutil
    except ImportError:
        return True # Cannot check without psutil
        
    target_w = int(w * f)
    target_h = int(h * f)
    target_pixels = target_w * target_h
    megapixels = target_pixels / 1_000_000
    
    # Estimate RAM (Very rough heuristic for Float32 Latent Pipeline)
    # Empirical Rule of Thumb: 1MP output needs ~1GB RAM on CPU for safe execution.
    estimated_ram_gb = (megapixels * 0.8) if dev == "cpu" else (megapixels * 0.4) 
    
    vm = psutil.virtual_memory()
    available_gb = vm.available / (1024**3)
    
    is_huge = megapixels > 25  # > 5K/6K image
    is_tight = estimated_ram_gb > (available_gb * 0.9)
    
    if is_huge or is_tight:
        print("\n⚠️  RESOURCE WARNING: High-Resolution Upscale Detected")
        print(f"   Input:  {w}x{h}")
        print(f"   Target: {target_w}x{target_h} ({megapixels:.1f} MP)")
        print(f"   Device: {dev.upper()}")
        print(f"   Est. RAM Required: ~{estimated_ram_gb:.1f} GB")
        print(f"   Available RAM:      {available_gb:.1f} GB")
        
        if is_tight:
             print("   🔴 WARNING: This may cause massive swapping or system freeze.")
        elif is_huge:
             print("   🟠 WARNING: This resolution is extremely high (Billboard size).")
        
        if os.environ.get("AI_MEDIA_FORCE", "0") == "1":
             print("   (Proceeding due to --force flag)")
             return True
             
        confirm = input("\n   Do you want to proceed? [y/N]: ").strip().lower()
        return confirm == 'y'
    return True


def simple_upscale_image(image_path, output_path, factor=2.0):
    """Simple non-AI image upscaling using PIL Lanczos interpolation.
    
    Fast, preserves original quality without AI hallucination.
    Uses Lanczos algorithm which is considered the best for downscaling/upscaling.
    """
    from PIL import Image
    
    print(f"🔍 Simple Upscaling Image: {image_path}")
    print(f"   Method: PIL Lanczos (No AI)")
    print(f"   Factor: {factor}x")
    
    # Pre-flight check: Avoid wasting processing time if output exists
    if Path(output_path).exists():
        if os.environ.get("AI_MEDIA_FORCE", "0") != "1":
            print(f"\n⚠️  Output file already exists: {output_path}")
            confirm = input("   Overwrite? [y/N]: ").strip().lower()
            if confirm != 'y':
                print("❌ Aborted to prevent overwrite.")
                return False
        else:
            print(f"   ⚠️  Overwriting existing file (--force).")
    
    try:
        image = Image.open(image_path).convert("RGB")
        orig_w, orig_h = image.size
        target_w = int(orig_w * factor)
        target_h = int(orig_h * factor)
        
        print(f"   {orig_w}x{orig_h} → {target_w}x{target_h}")
        
        upscaled = image.resize((target_w, target_h), Image.LANCZOS)
        
        # Ensure output directory exists
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        upscaled.save(output_path)
        
        print(f"✅ Simple upscaled image saved to {output_path}")
        return True
        
    except Exception as e:
        print(f"❌ Simple upscaling failed: {e}")
        return False


def simple_upscale_video(video_path, output_path, factor=2.0):
    """Simple non-AI video upscaling using FFmpeg scale filter.
    
    Fast, uses FFmpeg's high-quality Lanczos scaling.
    """
    import subprocess
    
    print(f"🔍 Simple Upscaling Video: {video_path}")
    print(f"   Method: FFmpeg Lanczos (No AI)")
    print(f"   Factor: {factor}x")
    
    # Pre-flight check: Avoid wasting processing time if output exists
    if Path(output_path).exists():
        if os.environ.get("AI_MEDIA_FORCE", "0") != "1":
            print(f"\n⚠️  Output file already exists: {output_path}")
            confirm = input("   Overwrite? [y/N]: ").strip().lower()
            if confirm != 'y':
                print("❌ Aborted to prevent overwrite.")
                return False
        else:
            print(f"   ⚠️  Overwriting existing file (--force).")
    
    try:
        # Get video dimensions first
        probe_cmd = [
            "ffprobe", "-v", "error",
            "-select_streams", "v:0",
            "-show_entries", "stream=width,height",
            "-of", "csv=p=0:s=x",
            str(video_path)
        ]
        result = subprocess.run(probe_cmd, capture_output=True, text=True)
        if result.returncode != 0:
            print(f"❌ Failed to probe video: {result.stderr}")
            return False
        
        orig_w, orig_h = map(int, result.stdout.strip().split('x'))
        target_w = int(orig_w * factor)
        target_h = int(orig_h * factor)
        
        # FFmpeg requires even dimensions
        target_w = target_w if target_w % 2 == 0 else target_w + 1
        target_h = target_h if target_h % 2 == 0 else target_h + 1
        
        print(f"   {orig_w}x{orig_h} → {target_w}x{target_h}")
        
        # Ensure output directory exists
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        # FFmpeg upscale with Lanczos (format-aware encoding)
        encoding_params = get_video_encoding_params(str(output_path))
        # For upscaling, we want to preserve audio and use high quality
        # Filter out audio codec from encoding_params and use -c:a copy
        video_params = [p for i, p in enumerate(encoding_params) if not (encoding_params[i-1:i] == ["-c:a"] or p in ["aac", "libopus", "wmav2", "mp3"])]
        video_params = [p for p in video_params if p != "-c:a"]
        
        cmd = [
            "ffmpeg", "-y", "-i", str(video_path),
            "-vf", f"scale={target_w}:{target_h}:flags=lanczos",
            *video_params,
            "-c:a", "copy",  # Preserve audio
            str(output_path),
            "-loglevel", "warning"
        ]
        
        print("   ⏳ Processing with FFmpeg...")
        result = subprocess.run(cmd, capture_output=True, text=True)
        
        if result.returncode != 0:
            print(f"❌ FFmpeg failed: {result.stderr}")
            return False
        
        print(f"✅ Simple upscaled video saved to {output_path}")
        return True
        
    except Exception as e:
        print(f"❌ Simple upscaling failed: {e}")
        return False

def upscale_image_file(image_path, output_path, strength=0.0, factor=2.0):
    """Upscale an image using smart multi-stage AI upscaling.
       
       Uses optimal combination of x4, x2 AI passes + final Lanczos resize:
       - 6x  → 4x AI + 1.5x Lanczos
       - 8x  → 4x AI + 2x AI
       - 10x → 4x AI + 2x AI + 1.25x Lanczos
       
       Args:
           image_path: Path to image to upscale
           output_path: Output file path
           strength: Noise strength 0.0-1.0 (x4 upscaler only). Higher values add more
                     noise, allowing the model to generate more details/texture but
                     potentially diverging from the original. Default 0.0 keeps original.
           factor: Upscale factor (e.g., 2.0, 4.0, 6.0, 8.0)
    """
    
    
    # Select Model based on factor
    # <= 2.0x -> use x2 Latent Upscaler (Fast, Faithful)
    # > 2.0x  -> use x4 Upscaler (Detailed)
    use_x2_model = (factor <= 2.0)
    model_id = IMAGE_MODELS['upscaler_x2'] if use_x2_model else IMAGE_MODELS['upscaler']
    
    print(f"🚀 Upscaling Image: {image_path}")
    print(f"   Model: {model_id}")
    print(f"   Target Factor: {factor}x")
    
    # Noise level for x4 upscaler (maps 0.0-1.0 to 0-100)
    # Default to 0 (faithful to original) - user can increase for more detail generation
    noise_level = int(strength * 100)
    
    if use_x2_model:
        if strength > 0:
            print(f"   ⚠️  Note: x2 Latent Upscaler does not support --upscale-strength. Ignored.")
    else:
        if strength > 0:
            print(f"   Noise Level: {noise_level} (strength={strength} - more creative/details)")
        else:
            print(f"   Noise Level: {noise_level} (faithful to original)")
    
    # Pre-flight check: Avoid wasting processing time if output exists
    if Path(output_path).exists():
        if os.environ.get("AI_MEDIA_FORCE", "0") != "1":
            print(f"\n⚠️  Output file already exists: {output_path}")
            confirm = input("   Overwrite? [y/N]: ").strip().lower()
            if confirm != 'y':
                print("❌ Aborted to prevent overwrite.")
                return False
        else:
            print(f"   ⚠️  Overwriting existing file (--force).")
    
    if factor > 4.0:
        print(f"   ℹ️  Factor > 4x detected. This will require multiple AI passes.")
    
    try:
        from diffusers import StableDiffusionUpscalePipeline, StableDiffusionLatentUpscalePipeline
        from PIL import Image
        import torch
        import time 
        
        start_time = time.time()


        device, dtype = get_optimal_device_and_dtype()
        
        # MPS GLOBAL FIX: Force CPU for ALL Upscaling (x2 and x4)
        # Reason 1 (x2): "View size not compatible" (Tensor Stride errors)
        # Reason 2 (x4): "MPSNDArrayMatrixMultiplication ... too large for kernel" (Driver Limit)
        # Only CPU can handle these massive tensors safely.
        if device.type == "mps":
            print("   ⚠️  MPS Compatibility: Switching to CPU for Upscaling (Avoids Kernel Crashes).")
            device = torch.device("cpu")
            dtype = torch.float32  # BFloat16 causes hangs on Apple Silicon CPU
        
        # Load Image
        try:
            image = Image.open(image_path).convert("RGB")
            orig_w, orig_h = image.size
        except Exception as e:
            print(f"❌ Error loading source image: {e}")
            return False

        # --- RESOURCE SAFETY CHECK ---
        if not check_resources_and_confirm(orig_w, orig_h, factor, device.type):
             print("❌ Upload aborted by user.")
             return False

        # ===== SMART MULTI-STAGE UPSCALING =====
        # Strategy: Use optimal combination of 4x, 2x AI passes + final Lanczos resize
        # Examples:
        #   6x  → 4x AI + 1.5x Lanczos
        #   8x  → 4x AI + 2x AI  
        #   10x → 4x AI + 2x AI + 1.25x Lanczos
        #   3x  → 2x AI + 1.5x Lanczos
        #   5x  → 4x AI + 1.25x Lanczos
        
        # Calculate optimal pass sequence
        passes = []
        remaining = factor
        
        while remaining >= 2.0:
            if remaining >= 4.0:
                passes.append(('x4', 4.0))
                remaining /= 4.0
            elif remaining >= 2.0:
                passes.append(('x2', 2.0))
                remaining /= 2.0
        
        # Remaining factor will be handled by Lanczos at the end
        final_lanczos_factor = remaining if remaining > 1.0 else None
        
        # Display the plan
        print(f"\n   📋 Upscale Plan:")
        for i, (model_type, scale) in enumerate(passes, 1):
            print(f"      Pass {i}: {model_type} AI ({scale}x)")
        if final_lanczos_factor:
            print(f"      Final: Lanczos resize ({final_lanczos_factor:.2f}x)")
        print("")
        
        # Load both pipelines if needed (lazy loading)
        pipe_x2 = None
        pipe_x4 = None
        
        def get_pipeline(model_type):
            nonlocal pipe_x2, pipe_x4
            
            if model_type == 'x2':
                if pipe_x2 is None:
                    print(f"   🔗 Loading x2 Latent Upscaler...")
                    pipe_x2 = StableDiffusionLatentUpscalePipeline.from_pretrained(
                        IMAGE_MODELS['upscaler_x2'],
                        torch_dtype=dtype,
                    )
                    if device.type != "cpu":
                        pipe_x2.enable_model_cpu_offload()
                    else:
                        pipe_x2.to(device)
                    if hasattr(pipe_x2, 'vae') and hasattr(pipe_x2.vae, 'enable_tiling'):
                        pipe_x2.vae.enable_tiling()
                return pipe_x2, 64  # alignment requirement
            else:
                if pipe_x4 is None:
                    print(f"   🔗 Loading x4 Upscaler...")
                    pipe_x4 = StableDiffusionUpscalePipeline.from_pretrained(
                        IMAGE_MODELS['upscaler'],
                        torch_dtype=dtype,
                        variant="fp16" if dtype == torch.float16 else None
                    )
                    if device.type != "cpu":
                        pipe_x4.enable_model_cpu_offload()
                    else:
                        pipe_x4.to(device)
                    if hasattr(pipe_x4, 'vae') and hasattr(pipe_x4.vae, 'enable_tiling'):
                        pipe_x4.vae.enable_tiling()
                return pipe_x4, 8  # alignment requirement
        
        # Upscaling prompts (neutral to avoid hallucination)
        upscale_prompt = "sharp, high resolution"
        negative_prompt = "blur, noise, artifacts, distortion, jpeg artifacts, oversaturated, low quality"
        
        current_image = image
        
        monitor = ResourceMonitor()
        monitor.__enter__()
        for pass_idx, (model_type, step_scale) in enumerate(passes, 1):
            print("")
            print("=" * 60)
            print(f"🎨 Pass {pass_idx}/{len(passes)}: {model_type} AI Upscaling ({step_scale}x)")
            print("=" * 60)
            print("")
            
            pipe, alignment = get_pipeline(model_type)
            
            # --- DIMENSION ALIGNMENT FIX ---
            img_w, img_h = current_image.size
            pad_w = (alignment - (img_w % alignment)) % alignment
            pad_h = (alignment - (img_h % alignment)) % alignment
            
            if pad_w > 0 or pad_h > 0:
                padded_w, padded_h = img_w + pad_w, img_h + pad_h
                final_w, final_h = int(img_w * step_scale), int(img_h * step_scale)
                print(f"   ℹ️  Temporarily padding {img_w}x{img_h} → {padded_w}x{padded_h} ({alignment}px alignment required)")
                print(f"       Will crop back to {final_w}x{final_h} after upscaling.")
                
                # Create padded image (reflect padding for seamless edges)
                padded_image = Image.new("RGB", (padded_w, padded_h))
                padded_image.paste(current_image, (0, 0))
                # Mirror-fill the padding area for better edge blending
                if pad_w > 0:
                    right_edge = current_image.crop((img_w - pad_w, 0, img_w, img_h))
                    padded_image.paste(right_edge.transpose(Image.FLIP_LEFT_RIGHT), (img_w, 0))
                if pad_h > 0:
                    bottom_edge = current_image.crop((0, img_h - pad_h, img_w, img_h))
                    padded_image.paste(bottom_edge.transpose(Image.FLIP_TOP_BOTTOM), (0, img_h))
            else:
                padded_image = current_image
            
            # Run upscaling with model-specific parameters
            print(f"   ⏳ Rendering Pass {pass_idx}...", flush=True)
            if model_type == 'x2':
                # x2 Latent Upscaler: doesn't support noise_level or negative_prompt
                upscaled_result = pipe(
                    prompt=upscale_prompt, 
                    image=padded_image, 
                    num_inference_steps=50,
                    show_progress_bar=True # Explicitly show progress for Step 2
                ).images[0]
            else:
                # x4 Upscaler: supports noise_level and negative_prompt
                upscaled_result = pipe(
                    prompt=upscale_prompt, 
                    image=padded_image, 
                    negative_prompt=negative_prompt,
                    noise_level=noise_level,
                    num_inference_steps=75,
                    show_progress_bar=True # Explicitly show progress for Step 2
                ).images[0]
            
            # Crop back to target dimensions (remove padding effect)
            target_w_pass = int(img_w * step_scale)
            target_h_pass = int(img_h * step_scale)
            if upscaled_result.size != (target_w_pass, target_h_pass):
                current_image = upscaled_result.crop((0, 0, target_w_pass, target_h_pass))
            else:
                current_image = upscaled_result
            
            print(f"   ✓ Pass {pass_idx} complete: {current_image.size[0]}x{current_image.size[1]}")
        
        # Final Resize to exact factor using Lanczos (high-quality non-AI resize)
        target_w = int(orig_w * factor)
        target_h = int(orig_h * factor)
        
        if current_image.size != (target_w, target_h):
            actual_lanczos = target_w / current_image.size[0]
            print(f"\n   ↘️  Lanczos resize ({actual_lanczos:.2f}x) to exact target: {target_w}x{target_h}")
            current_image = current_image.resize((target_w, target_h), Image.LANCZOS)
        
        # Ensure output directory exists
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        current_image.save(output_path)
        print(f"\n✅ Upscaled image saved to {output_path}")
        
        # Performance Summary
        monitor.__exit__(None, None, None)
        duration = time.time() - start_time
        cpu_p, ram_avail, vram_avail, gpu_p = monitor.get_averages()
        
        mem = (ram_avail, vram_avail, cpu_p, gpu_p)
        print(f"   ✓ Processed in {duration:.1f}s (RAM: {mem[0]:.1f}GB | VRAM: {mem[1]:.1f}GB | CPU: {mem[2]:.1f}% | GPU: {mem[3]:.1f}%)")
        
        # Write JSON Report if requested
        try:
            g_args = globals().get("args")
            if g_args and hasattr(g_args, "report_json") and g_args.report_json:
                import json
                stats = {
                    "time": duration,
                    "ram": mem[0],
                    "vram": mem[1],
                    "cpu": mem[2],
                    "gpu": mem[3]
                }
                with open(g_args.report_json, 'w') as f:
                    json.dump(stats, f)
        except Exception as e:
            print(f"   ⚠️  Failed to write report JSON: {e}")

        return True
        
    except Exception as e:
        print(f"❌ Upscaling failed: {e}")
        return False

def upscale_image_fast(input_path, output_path, factor=4.0):
    """Upscale image using Real-ESRGAN (Fast, single pass)."""
    if not HAS_REALESRGAN:
        print("❌ Real-ESRGAN not installed. Cannot run fast upscale.")
        print("   Please install: pip install realesrgan")
        return False

    print(f"🚀 Upscaling Image (Fast Mode): {input_path}", flush=True)
    print(f"   Factor: {factor}x", flush=True)
    
    try:
        import cv2
        from PIL import Image
        import numpy as np
        import time
        
        start_time = time.time()
        
        if not os.path.exists(input_path):
            print(f"❌ Input file not found: {input_path}")
            return False

        # Ensure output directory exists
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)

        # Check for CUDA/MPS availability
        device, _ = get_optimal_device_and_dtype(quiet=True)
        print(f"   Device: {device}", flush=True)

        # Load default model (RealESRGAN_x4plus)
        model_name = 'RealESRGAN_x4plus'
        
        # Model configuration
        model = RRDBNet(num_in_ch=3, num_out_ch=3, num_feat=64, num_block=23, num_grow_ch=32, scale=4)
        netscale = 4

        # Determine model path
        weights_dir = Path("weights")
        weights_dir.mkdir(exist_ok=True)
        model_path = weights_dir / f"{model_name}.pth"
        
        if not model_path.exists():
            print(f"   ⬇️  Downloading model: {model_name}...")
            model_path_str = f'https://github.com/xinntao/Real-ESRGAN/releases/download/v0.1.0/{model_name}.pth'
        else:
            model_path_str = str(model_path)

        upsampler = RealESRGANer(
            scale=netscale,
            model_path=model_path_str,
            model=model,
            tile=0,
            tile_pad=10,
            pre_pad=0,
            half=True if device.type != 'cpu' else False,
            device=device,
        )

        # Load image with OpenCV (BGR format)
        img = cv2.imread(input_path, cv2.IMREAD_UNCHANGED)
        if img is None:
            print(f"❌ Failed to load image: {input_path}")
            return False

        orig_h, orig_w = img.shape[:2]
        print(f"   Input: {orig_w}x{orig_h}")

        # Start resource monitoring
        monitor = ResourceMonitor()
        monitor.__enter__()
        
        try:
            # Upscale with Real-ESRGAN
            print(f"   🎨 Enhancing details with Real-ESRGAN...", flush=True)
            
            # Since Real-ESRGAN's .enhance() is a single atomic call, we show a waiting indicator
            # or a synthetic progress message for high resolutions.
            if factor > 4.0 or orig_w > 4000:
                print(f"   ⏳ This is a very high resolution upscale. GPU is working...", flush=True)
                
            output, _ = upsampler.enhance(img, outscale=factor)
            
            out_h, out_w = output.shape[:2]
            print(f"   ✅ Enhancement complete. Final size: {out_w}x{out_h}", flush=True)
            
            # Save output
            cv2.imwrite(output_path, output)
            print(f"\n✅ Fast upscaled image saved to {output_path}")
            
        finally:
            monitor.__exit__(None, None, None)

        # Performance Summary
        duration = time.time() - start_time
        cpu_p, ram_avail, vram_avail, gpu_p = monitor.get_averages()
        
        mem = (ram_avail, vram_avail, cpu_p, gpu_p)
        print(f"   ✓ Processed in {duration:.1f}s (RAM: {mem[0]:.1f}GB | VRAM: {mem[1]:.1f}GB | CPU: {mem[2]:.1f}% | GPU: {mem[3]:.1f}%)")
        
        # Write JSON Report if requested
        try:
            g_args = globals().get("args")
            if g_args and hasattr(g_args, "report_json") and g_args.report_json:
                import json
                stats = {
                    "time": duration,
                    "ram": mem[0],
                    "vram": mem[1],
                    "cpu": mem[2],
                    "gpu": mem[3]
                }
                with open(g_args.report_json, 'w') as f:
                    json.dump(stats, f)
        except Exception as e:
            print(f"   ⚠️  Failed to write report JSON: {e}")

        return True

    except Exception as e:
        print(f"❌ Fast upscaling failed: {e}")
        import traceback
        traceback.print_exc()
        return False

def upscale_video_fast(video_path, output_path, factor=4.0, device=None, model_name=None, force_realesrgan=False, upscale_half=False, video_upscaler=None, codec="auto"):
    """Upscale video using Real-ESRGAN (Fast, single pass per frame)."""
    if not HAS_REALESRGAN:
        print("❌ Real-ESRGAN not installed. Cannot run fast upscale.")
        print("   Please install: pip install realesrgan")
        return False

    print(f"🚀 Upscaling Video (Fast Mode): {video_path}", flush=True)
    print(f"   Factor: {factor}x (fixed or scaled)", flush=True)
    
    # Real-ESRGAN typically uses x4 models. If factor != 4, we might need resizing.
    # But usually x4plus is a 4x model.
    # If user asks for 2x, we could upscale 4x then downscale 0.5x, or find a 2x model.
    # For now, we will use the standard x4 model and Lanczos resize final output if needed, 
    # OR rely on RealESRGANer's outscale param if it supports arbitrary scaling (it does).
    
    try:
        # --- OpenH264 DLL Setup (Windows) - BEFORE cv2 import ---
        # Python 3.8+ requires explicit DLL directory registration
        # The DLL must be in place BEFORE cv2 is imported
        if os.name == 'nt':
            dll_name = "openh264-1.8.0-win64.dll"
            install_dir = os.path.join(sys.prefix, 'Scripts')
            if not os.path.exists(install_dir):
                install_dir = os.getcwd()
            
            # Determine cv2 package directory (without importing cv2)
            cv2_dir = os.path.join(sys.prefix, 'Lib', 'site-packages', 'cv2')
            cv2_dll_path = os.path.join(cv2_dir, dll_name) if os.path.isdir(cv2_dir) else None
            
            dll_path = os.path.join(install_dir, dll_name)
            cwd_dll = os.path.join(os.getcwd(), dll_name)
            
            # Download DLL if not present anywhere
            if not os.path.exists(dll_path) and not os.path.exists(cwd_dll) and (not cv2_dll_path or not os.path.exists(cv2_dll_path)):
                try:
                    print(f"   ⬇️  Downloading OpenH264 library for H.264 support...")
                    import urllib.request
                    import bz2
                    import shutil
                    
                    url = "https://github.com/cisco/openh264/releases/download/v1.8.0/openh264-1.8.0-win64.dll.bz2"
                    bz2_path = dll_path + ".bz2"
                    
                    urllib.request.urlretrieve(url, bz2_path)
                    
                    with bz2.BZ2File(bz2_path, 'rb') as fr, open(dll_path, 'wb') as fw:
                        shutil.copyfileobj(fr, fw)
                    
                    os.remove(bz2_path)
                    print(f"   ✅ OpenH264 downloaded: {dll_name}")
                except Exception as e:
                    print(f"   ⚠️  Failed to download OpenH264: {e}")
            
            # Copy DLL to cv2 package directory (where OpenCV looks for it)
            if cv2_dll_path and not os.path.exists(cv2_dll_path):
                src_dll = dll_path if os.path.exists(dll_path) else cwd_dll
                if os.path.exists(src_dll):
                    try:
                        import shutil
                        shutil.copy2(src_dll, cv2_dll_path)
                        print(f"   ✅ OpenH264 installed to cv2 package")
                    except Exception:
                        pass
            
            # Register DLL directory as fallback
            if hasattr(os, 'add_dll_directory'):
                try:
                    os.add_dll_directory(install_dir)
                    if cv2_dll_path:
                        os.add_dll_directory(os.path.dirname(cv2_dll_path))
                except Exception:
                    pass
            os.environ['PATH'] = install_dir + os.pathsep + os.environ.get('PATH', '')
        # ----------------------------------------------------
        
        import cv2
        import torch
        
        if not os.path.exists(video_path):
            print(f"❌ Input file not found: {video_path}")
            return False

        # Ensure output directory exists
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)

        # Check for CUDA/MPS availability
        device, _ = get_optimal_device_and_dtype(quiet=True)
        # RealESRGANer expects device object
        
        print(f"   Device: {device}", flush=True)

        # Load default model (RealESRGAN_x4plus)
        model_name = 'RealESRGAN_x4plus'
        
        # Model configuration
        model = RRDBNet(num_in_ch=3, num_out_ch=3, num_feat=64, num_block=23, num_grow_ch=32, scale=4)
        netscale = 4

        # Determine model path
        # Check standard weights folder or download
        weights_dir = Path("weights")
        weights_dir.mkdir(exist_ok=True)
        model_path = weights_dir / f"{model_name}.pth"
        
        if not model_path.exists():
            print(f"   ⬇️  Downloading model: {model_name}...")
            model_url = f'https://github.com/xinntao/Real-ESRGAN/releases/download/v0.1.0/{model_name}.pth'
            # We pass the URL string to RealESRGANer if file doesn't exist? 
            # Actually RealESRGANer takes 'model_path' as string. If it's a local path it uses it.
            # If we want auto-download, we might need to do it ourselves or rely on library.
            # The library logic usually downloads if we pass a name, but main class takes path.
            # Let's provide the URL to a downloader helper or just pass the path and hope user has it?
            # Prototype succeeded, but prototype printed "Downloading...". 
            # In prototype we set model_path to URL if local file missing?
            # Let's check prototype code logic again.
            # Ah, checked previous turn: 
            # if not os.path.isfile(model_path): model_path = f'https...'. 
            # So RealESRGANer handles URL download if passed as path!
            model_path_str = str(model_path) if model_path.exists() else f'https://github.com/xinntao/Real-ESRGAN/releases/download/v0.1.0/{model_name}.pth'
        else:
            model_path_str = str(model_path)

        upsampler = RealESRGANer(
            scale=netscale,
            model_path=model_path_str,
            model=model,
            tile=0, # Auto-tile? 0 means no tile. Use >0 for low VRAM but checks needed.
            tile_pad=10,
            pre_pad=0,
            half=True if device.type != 'cpu' else False, # fp16 on GPU
            device=device,
        )

        # Video Capture
        cap = cv2.VideoCapture(video_path)
        if not cap.isOpened():
            print(f"❌ Failed to open input video.")
            return False

        width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
        height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
        fps = cap.get(cv2.CAP_PROP_FPS)
        total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))

        # (OpenH264 DLL setup is handled before cv2 import above)

        print(f"   Input: {width}x{height} @ {fps}fps")
        
        # Calculate Output dimensions
        target_w = int(width * factor)
        target_h = int(height * factor)
        
        # FFmpeg requires even dimensions
        if target_w % 2 != 0: target_w += 1
        if target_h % 2 != 0: target_h += 1

        print(f"   Output: {target_w}x{target_h} (Target Codec: {codec.upper()})")
        
        # Resolution limit check - HEVC Level 6.2 maxes out at 8192x4320
        # libx265 can handle higher but may fail on extreme resolutions
        # Maximum practical limit is around 16K (16384x16384)
        # Beyond that, we auto-switch to AV1
        # HARD LIMIT: 15K max - tested on Windows, 16K (432MB/frame) exceeds OS pipe buffer limits
        # 15K (380MB/frame) works reliably
        MAX_DIMENSION = 15360
        if target_w > MAX_DIMENSION or target_h > MAX_DIMENSION:
            print(f"   ❌ Target Resolution {target_w}x{target_h} exceeds the stable 15K limit ({MAX_DIMENSION}px).")
            print(f"   ℹ️  The input video is {width}x{height}. A factor of {factor}x results in ~{max(target_w, target_h)/1024:.1f}K.")
            print(f"   ℹ️  At 16K+, frames exceed the Windows pipe buffer (432MB+), causing FFmpeg to crash.")
            print(f"   ℹ️  Maximum safe upscale factor for this video: {MAX_DIMENSION / max(width, height):.1f}x (to reach 15K).")
            return False


        # We will write to a temp file first (no audio), then mux audio
        temp_video_out = output_path + ".temp.mp4"
        
        # Video Writer
        # Try prioritized list of codecs for 8K Support
        # 1. avc1 (H.264) - Best compatibility, but needs OpenH264 on Windows
        # 2. hevc (H.265) - Best efficiency, high-res
        # 3. mp4v (MPEG-4) - Fallback, but size limit at 4K/8K?
        
        codecs_to_try = []
        if target_w > 4096 or target_h > 2304:
             codecs_to_try = ['avc1', 'hevc', 'mp4v']
        else:
             codecs_to_try = ['mp4v', 'avc1'] # Default behavior
        # Try to suppress OpenCV/FFmpeg backend errors during this probe
        from contextlib import contextmanager
        # os and sys are already imported globally


        @contextmanager
        def suppress_stderr():
            with open(os.devnull, "w") as devnull:
                old_stderr = os.dup(2)
                sys.stderr.flush()
                try:
                    os.dup2(devnull.fileno(), 2)
                    yield
                finally:
                    os.dup2(old_stderr, 2)
                    os.close(old_stderr)

        # Determine OpenCV backend preference
        api_preference = cv2.CAP_ANY # Default
        if sys.platform == 'win32':
            api_preference = cv2.CAP_DSHOW # DirectShow for Windows
        elif sys.platform == 'darwin':
            api_preference = cv2.CAP_AVFOUNDATION # AVFoundation for macOS
        else: # Linux and others
            api_preference = cv2.CAP_FFMPEG
    
        # 4. Initialize Video Writer
        # Try native OpenCV writers first (faster if they work), fallback to Pipe.
        # HEVC is preferred for 8K but harder to get working natively.
        
        fourcc = None
        fallback_pipe = False
        
        # Determine preferred codec based on arg and resolution
        target_codecs = []
        
        if codec == 'hevc':
             target_codecs = ['hevc', 'hvc1'] # Try HEVC variants
        elif codec == 'av1':
             target_codecs = ['av01'] # Rare in OpenCV, but try anyway
        elif codec == 'h264':
             # H.264 max is 8K (8192x4320). Force HEVC if user requests h264 but resolution exceeds spec.
             if target_w > 8192 or target_h > 4320:
                 print(f"   ℹ️  Resolution {target_w}x{target_h} exceeds H.264 limits. Using HEVC instead.")
                 target_codecs = ['hevc', 'hvc1']
             else:
                 target_codecs = ['avc1']
        else: # auto
            # Auto-select codec based on resolution
            if target_w > 8192 or target_h > 4320:
                # 8K+ requires HEVC
                target_codecs = ['hevc', 'hvc1', 'mp4v']
            elif target_w > 3840 or target_h > 2160:
                # 4K-8K: try H.264 first, fallback to HEVC
                target_codecs = ['avc1', 'hevc', 'mp4v']
            else:
                # Standard resolutions
                target_codecs = ['mp4v', 'avc1']
                
        # Add fail-safes
        if 'mp4v' not in target_codecs: target_codecs.append('mp4v')
        
        out = None
        active_codec = "unknown"
        
        # Attempt native initialization
        with suppress_stderr():
            for c in target_codecs:
                try:
                    fourcc = cv2.VideoWriter_fourcc(*c)
                    out = cv2.VideoWriter(temp_video_out, api_preference, fourcc, fps, (target_w, target_h))
                    if out.isOpened():
                        active_codec = c
                        # Basic check if it accepts frames of this size?
                        # Some writers open but fail on write. Hard to test without writing.
                        # We will assume success and catch write errors in the loop.
                        break
                except Exception:
                    pass # Try next codec
        
        if not out or not out.isOpened():
             print(f"   ℹ️  OpenCV writers failed. Falling back to robust FFmpeg Pipe...")
             fallback_pipe = True
             out = "FFMPEG_PIPE"
        else:
             print(f"   🎥 Video Writer initialized with codec: '{active_codec}'")
             
        if out == "FFMPEG_PIPE":
            # Raw Pipe Fallback
            # Determine encoder lib based on request
            vcodec_lib = "libx264"
            
            if codec == 'hevc':
                 # Try hardware HEVC first at target resolution
                 if device.type == 'cuda' and _check_ffmpeg_encoder("hevc_nvenc", target_w, target_h):
                     vcodec_lib = "hevc_nvenc"
                 elif device.type == 'mps' and _check_ffmpeg_encoder("hevc_videotoolbox", target_w, target_h):
                     vcodec_lib = "hevc_videotoolbox"
                 else:
                     vcodec_lib = "libx265"
                     
            elif codec == 'av1':
                 # Try hardware AV1 first at target resolution
                 if device.type == 'cuda' and _check_ffmpeg_encoder("av1_nvenc", target_w, target_h):
                     vcodec_lib = "av1_nvenc"
                     print(f"   ℹ️  Hardware AV1 Encoding supported (av1_nvenc).")
                 else:
                     # Switch to HEVC as fallback, checking HW then SW
                     print(f"   ℹ️  Hardware AV1 not supported at {target_w}x{target_h}. Fallback protocol initiated.")
                     if device.type == 'cuda' and _check_ffmpeg_encoder("hevc_nvenc", target_w, target_h):
                         vcodec_lib = "hevc_nvenc" 
                         print(f"   ℹ️  Fallback: Using Hardware HEVC (hevc_nvenc).")
                     elif device.type == 'mps' and _check_ffmpeg_encoder("hevc_videotoolbox", target_w, target_h):
                         vcodec_lib = "hevc_videotoolbox"
                         print(f"   ℹ️  Fallback: Using Hardware HEVC (hevc_videotoolbox).")
                     else:
                         vcodec_lib = "libx265"
                         print(f"   ℹ️  Fallback: Using Software HEVC (libx265).")
            
            else: # h264 or auto
                # For H.264, check limits. If target > 8K (or even 4K depending on card), switch to HEVC
                # Actually, check_ffmpeg_encoder("h264_nvenc", w, h) will tell us if H.264 is viable.
                use_h264 = True
                if device.type == 'cuda' and _check_ffmpeg_encoder("h264_nvenc", target_w, target_h):
                     vcodec_lib = "h264_nvenc"
                elif device.type == 'mps' and _check_ffmpeg_encoder("h264_videotoolbox", target_w, target_h):
                     vcodec_lib = "h264_videotoolbox"
                else:
                     # Fallback to software h264
                     # But software h264 also fails at extreme high res (level limits).
                     # Probe software h264
                     if _check_ffmpeg_encoder("libx264", target_w, target_h):
                         vcodec_lib = "libx264"
                     else:
                         use_h264 = False
                
                # If H.264 failed or if target is visibly too large for efficient H.264
                if not use_h264 or target_w > 8192 or target_h > 4320:
                     # Switch to HEVC
                     msg = "H.264 limits exceeded" if use_h264 else "H.264 encoder failed"
                     if device.type == 'cuda' and _check_ffmpeg_encoder("hevc_nvenc", target_w, target_h):
                          vcodec_lib = "hevc_nvenc"
                     elif device.type == 'mps' and _check_ffmpeg_encoder("hevc_videotoolbox", target_w, target_h):
                          vcodec_lib = "hevc_videotoolbox"
                     else:
                          vcodec_lib = "libx265"
                     print(f"   ℹ️  Resolution {target_w}x{target_h}: {msg}. Switching to HEVC ({vcodec_lib}).")

            # Final check for > 8K (Extreme resolutions)
            # Ensure we didn't pick something that will fail, though probing should have caught it.
            # If we are on software HEVC, we are usually good up to 16K.
            pass
            
            # Build FFmpeg command based on codec
            if vcodec_lib in ["av1_nvenc", "av1_vulkan"]:
                # AV1 encoder command
                if vcodec_lib == "av1_nvenc":
                    ffmpeg_cmd = [
                        'ffmpeg', '-y',
                        '-f', 'rawvideo',
                        '-vcodec', 'rawvideo',
                        '-s', f'{target_w}x{target_h}',
                        '-pix_fmt', 'bgr24',
                        '-r', str(fps),
                        '-i', '-',
                        '-c:v', 'av1_nvenc',
                        '-preset', 'p4',  # p1=fastest, p7=slowest. p4 is balanced.
                        '-cq', '30',      # Constant quality mode
                        '-pix_fmt', 'yuv420p',
                        temp_video_out
                    ]
                else:
                    ffmpeg_cmd = [
                        'ffmpeg', '-y',
                        '-f', 'rawvideo',
                        '-vcodec', 'rawvideo',
                        '-s', f'{target_w}x{target_h}',
                        '-pix_fmt', 'bgr24',
                        '-r', str(fps),
                        '-i', '-',
                        '-c:v', 'libsvtav1',
                        '-preset', '8',  # 0=slowest/best, 13=fastest. 8 is a good balance.
                        '-crf', '30',    # AV1 CRF: 0-63, lower=better. 30 is visually lossless.
                        '-pix_fmt', 'yuv420p',
                        temp_video_out
                    ]
            else:
                ffmpeg_cmd = [
                    'ffmpeg', '-y',
                    '-f', 'rawvideo',
                    '-vcodec', 'rawvideo',
                    '-s', f'{target_w}x{target_h}',
                    '-pix_fmt', 'bgr24',
                    '-r', str(fps),
                    '-i', '-',
                    '-c:v', vcodec_lib,
                    '-pix_fmt', 'yuv420p',
                ]
                
                # Dynamic Bitrate/Quality
                if "nvenc" in vcodec_lib:
                    ffmpeg_cmd.extend(['-preset', 'p4', '-cq', '20'])
                elif "videotoolbox" in vcodec_lib:
                    # VideoToolbox doesn't use -crf or -cq for high-res reliably
                    # We use -allow_sw 1 to permit software fallback if HW limits hit (rare but safe)
                    ffmpeg_cmd.extend(['-realtime', '1', '-allow_sw', '1'])
                else: # libx264/265
                    ffmpeg_cmd.extend(['-preset', 'fast', '-crf', '18' if codec == 'hevc' else '20'])
                
                ffmpeg_cmd.append(temp_video_out)
            
            import subprocess
            import tempfile
            
            # Use a NamedTemporaryFile for stderr to avoid pipe deadlocks
            # delete=False is required on Windows to allow the subprocess to open it
            err_file = tempfile.NamedTemporaryFile(delete=False, mode='w+')
            err_path = err_file.name
            
            try:
                 ffmpeg_proc = subprocess.Popen(
                    ffmpeg_cmd, 
                    stdin=subprocess.PIPE, 
                    stdout=subprocess.DEVNULL, 
                    stderr=err_file # Capture stderr in file to prevent buffer deadlock
                 )
                 # Quick check for immediate startup failure
                 time.sleep(0.5)
                 if ffmpeg_proc.poll() is not None:
                     # Process exited immediately
                     err_file.seek(0)
                     err_content = err_file.read()
                     print(f"   ❌ FFmpeg Pipe failed to start ({vcodec_lib}): {err_content.strip()}")
                     err_file.close() # Close handle
                     try: os.remove(err_path)
                     except: pass
                     return False
                 
                 print(f"   ✅ FFmpeg Pipe initialized ({vcodec_lib}).")
                 
            except Exception as e:
                 print(f"   ❌ FFmpeg Pipe failed to launch: {e}")
                 err_file.close()
                 try: os.remove(err_path)
                 except: pass
                 return False

        if 'start_time' not in locals():
            # Use the global time module (imported at top of file)
            start_time = time.time()
        
        print("   🎨 Upscaling frames...", flush=True)
        
        frame_idx = 0
        monitor = ResourceMonitor()
        if psutil: 
            monitor.__enter__() # Only start thread if psutil available

        try:
            while cap.isOpened():
                ret, frame = cap.read()
                if not ret:
                    break
                    
                frame_idx += 1
                
                # Process with RealESRGANer
                # It handles tiled processing internally if needed, very memory efficient
                try:
                    output, _ = upsampler.enhance(frame, outscale=factor)
                except Exception as e:
                    print(f"\n   ❌ Real-ESRGAN failed on frame {frame_idx}: {e}")
                    print(f"   ℹ️  This may indicate the resolution ({target_w}x{target_h}) exceeds GPU memory limits.")
                    raise
                
                # Resize if needed (for non-integer checks above, though RealESRGAN usually matches factor)
                # But we forced target_w/h to be even, so let's resize to match exactly
                if output.shape[1] != target_w or output.shape[0] != target_h:
                    output = cv2.resize(output, (target_w, target_h))
                
                if output.shape[1] != target_w or output.shape[0] != target_h:
                    output = cv2.resize(output, (target_w, target_h))
                
                if out == "FFMPEG_PIPE":
                    # Ensure uint8 data type for raw piping (fixes "Black Video" issue)
                    import numpy as np
                    if output.dtype != np.uint8:
                        output = output.astype(np.uint8)
                    try:
                        # Chunked write to avoid 'Errno 22' on massive frames (Windows pipe limits)
                        raw_data = output.tobytes()
                        chunk_size = 64 * 1024 * 1024 # 64MB chunks
                        for i in range(0, len(raw_data), chunk_size):
                            ffmpeg_proc.stdin.write(raw_data[i:i+chunk_size])
                    except Exception as e:
                        print(f"\n   ❌ FFmpeg pipe write failed on frame {frame_idx}: {e}")
                        print(f"   ℹ️  Frame size: {output.nbytes / 1024 / 1024:.1f}MB. Target Resolution: {target_w}x{target_h}")
                        raise
                else:
                    out.write(output)
                    
                if frame_idx % 10 == 0:
                    print(f"   Frame {frame_idx}/{total_frames}...", flush=True)
                    
        finally:
            cap.release()
            
            non_error_exit = True
            try:
                if hasattr(monitor, 'thread') and monitor.thread and monitor.thread.is_alive():
                    monitor.__exit__(None, None, None)
            except:
                non_error_exit = False
                
            if out == "FFMPEG_PIPE":
                try:
                    ffmpeg_proc.stdin.close()
                except:
                    pass
                
                # Robust wait with timeout (Increased to 30s for MP4 moov writing)
                try:
                    ffmpeg_proc.wait(timeout=30)
                except subprocess.TimeoutExpired:
                    print("   ⚠️  FFmpeg process hung. Forcing termination...", flush=True)
                    ffmpeg_proc.terminate()
                    try:
                        ffmpeg_proc.wait(timeout=5)
                    except subprocess.TimeoutExpired:
                        ffmpeg_proc.kill()
                
                # Close and remove stderr log
                try:
                    err_file.close()
                    if os.path.exists(err_path):
                        os.remove(err_path)
                except:
                    pass
                
                # Check exit code - if non-zero, it failed (or was killed)
                if ffmpeg_proc.returncode != 0:
                     print(f"   ❌ FFmpeg process failed/killed (Exit Code: {ffmpeg_proc.returncode}).")
                     if os.path.exists(temp_video_out):
                         try: os.remove(temp_video_out)
                         except: pass
                     return False

            # Friendly exit message
            if '_interrupted' in globals() and _interrupted:
                print("👋 Goodbye!")
                
        print(f"\n   ✅ Video track processing complete.")
        
        # Performance Summary
        duration = time.time() - start_time
        cpu_p, ram_avail, vram_avail, gpu_p = monitor.get_averages()
        
        mem = (ram_avail, vram_avail, cpu_p, gpu_p)
        print(f"   ✓ Processed in {duration:.1f}s (RAM: {mem[0]:.1f}GB | VRAM: {mem[1]:.1f}GB | CPU: {mem[2]:.1f}% | GPU: {mem[3]:.1f}%)")
        
        # Write JSON Report if requested (for Test Runner Stats)
        try:
            # We need to access 'args' from the global scope (main module args)
            # Or assume it was passed? Standard pattern in this file relies on global 'args' often,
            # but cleaner to check if 'args' is in globals.
            # Using 'globals().get("args")' as this function is top-level.
            g_args = globals().get("args")
            if g_args and hasattr(g_args, "report_json") and g_args.report_json:
                import json
                stats = {
                    "time": duration,
                    "ram": mem[0],
                    "vram": mem[1],
                    "cpu": mem[2],
                    "gpu": mem[3]
                }
                with open(g_args.report_json, 'w') as f:
                    json.dump(stats, f)
        except Exception as e:
            print(f"   ⚠️  Failed to write report JSON: {e}")

        # Check if source video has audio before attempting mux
        import subprocess
        
        def has_audio_track(video_file):
            """Check if video file has an audio track using ffprobe."""
            try:
                cmd = [
                    "ffprobe", "-v", "error", "-select_streams", "a",
                    "-show_entries", "stream=codec_type", "-of", "csv=p=0",
                    video_file
                ]
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=10)
                return "audio" in result.stdout
            except:
                return False
        
        if has_audio_track(video_path):
            # Mux Audio from source
            print(f"   🔗 Muxing audio from source...", flush=True)
            
            cmd = [
                 "ffmpeg", "-y",
                 "-i", temp_video_out,    # Input 0: Upscaled video (silent, already encoded)
                 "-i", video_path,        # Input 1: Original video (audio source)
                 "-map", "0:v",           # Use video from input 0
                 "-map", "1:a",           # Use audio from input 1
                 "-c:v", "copy",          # COPY video stream (preserves HEVC/H.264)
                 "-c:a", "aac",           # Encode audio to AAC
                 "-shortest",             # Match shortest
                 output_path,
                 "-loglevel", "error"
            ]
            
            try:
                subprocess.run(cmd, check=True)
                print(f"✅ Fast upscaled video saved to {output_path}")
                os.remove(temp_video_out)
                return True
            except subprocess.CalledProcessError:
                print("⚠️  Audio muxing failed. Saving silent video.")
                if os.path.exists(temp_video_out):
                      if os.path.exists(output_path): os.remove(output_path)
                      os.rename(temp_video_out, output_path)
                return True
        else:
            # No audio track in source - just rename temp to output
            print(f"   ℹ️  No audio track in source, skipping mux.")
            if os.path.exists(output_path): os.remove(output_path)
            os.rename(temp_video_out, output_path)
            print(f"✅ Fast upscaled video saved to {output_path}")
            return True

    except Exception as e:
        print(f"❌ Fast upscaling failed: {e}")
        return False

def upscale_video_file(video_path, output_path, strength=0.0, factor=2.0):
    """Upscale video by extracting frames, upscaling them (recursively if needed), and stitching back."""
    print(f"🚀 Upscaling Video: {video_path}")
    print(f"   Factor: {factor}x")
    
    # Pre-flight check: Avoid wasting processing time if output exists
    if Path(output_path).exists():
        if os.environ.get("AI_MEDIA_FORCE", "0") != "1":
            print(f"\n⚠️  Output file already exists: {output_path}")
            confirm = input("   Overwrite? [y/N]: ").strip().lower()
            if confirm != 'y':
                print("❌ Aborted to prevent overwrite.")
                return False
        else:
            print(f"   ⚠️  Overwriting existing file (--force).")
    
    if factor > 4.0:
        print(f"   ℹ️  Factor > 4x detected. This will require multiple AI passes per frame.")
    
    try:
        import cv2
        import shutil
        import subprocess
        
        # --- RESOURCE SAFETY CHECK (Get dims from video metadata first) ---
        cap_chk = cv2.VideoCapture(str(video_path))
        if cap_chk.isOpened():
             v_w = int(cap_chk.get(cv2.CAP_PROP_FRAME_WIDTH))
             v_h = int(cap_chk.get(cv2.CAP_PROP_FRAME_HEIGHT))
             cap_chk.release()
             
             # Reuse global checker
             # Note: For video, we might want to check device type, but we haven't loaded torch/device yet.
             # We can assume 'mps' if on Mac for the check warning purposes, or just check CPU RAM first.
             # Let's peek device quick
             import torch
             d_type = "mps" if torch.backends.mps.is_available() else ("cuda" if torch.cuda.is_available() else "cpu")
             
             if not check_resources_and_confirm(v_w, v_h, factor, d_type):
                  print("❌ Upload aborted by user.")
                  return False
        else:
             print("⚠️  Could not read video metadata for resource check.")
             
        # 1. Create temp directory
        temp_dir = Path("temp_upscale_frames")
        if temp_dir.exists(): shutil.rmtree(temp_dir)
        temp_dir.mkdir()
        
        # 2. Extract Frames
        print("🎞️  Extracting frames...")
        cam = cv2.VideoCapture(str(video_path))
        fps = cam.get(cv2.CAP_PROP_FPS)
        frame_count = 0
        
        while True:
            ret, frame = cam.read()
            if not ret: break
            
            frame_path = temp_dir / f"frame_{frame_count:05d}.png"
            cv2.imwrite(str(frame_path), frame)
            frame_count += 1
            
        cam.release()
        print(f"   Extracted {frame_count} frames.")
        
        # 3. Load Pipeline
        from diffusers import StableDiffusionUpscalePipeline, StableDiffusionLatentUpscalePipeline
        import torch
        from PIL import Image
        
        device, dtype = get_optimal_device_and_dtype()
        
        # MPS GLOBAL FIX: Force CPU for Video Upscaling too
        if device.type == "mps":
            print("   ⚠️  MPS Compatibility: Switching to CPU for Upscaling.")
            device = torch.device("cpu")
            dtype = torch.float32  # BFloat16 causes hangs on Apple Silicon CPU

        # Select Model based on factor
        use_x2_model = (factor <= 2.0)
        model_id = IMAGE_MODELS['upscaler_x2'] if use_x2_model else IMAGE_MODELS['upscaler']
        step_scale = 2.0 if use_x2_model else 4.0

        print(f"🔗 Loading Upscale Model ({'x2 Latent' if use_x2_model else 'x4 Std'})...")
        
        if use_x2_model:
            pipe = StableDiffusionLatentUpscalePipeline.from_pretrained(
                model_id, 
                torch_dtype=dtype, 
            )
        else:
            pipe = StableDiffusionUpscalePipeline.from_pretrained(
                model_id, 
                torch_dtype=dtype, 
                variant="fp16" if dtype == torch.float16 else None
            )
            
        # Memory Optimizations
        if device.type != "cpu":
             pipe.enable_model_cpu_offload() 
        else:
             pipe.to(device)
         
        # Enable tiling on CPU to survive high-res
        if hasattr(pipe, 'vae') and hasattr(pipe.vae, 'enable_tiling'):
            pipe.vae.enable_tiling()
        
        # 4. Process Frames
        print("🎨 Upscaling frames...")
        for i in range(frame_count):
            input_f = temp_dir / f"frame_{i:05d}.png"
            output_f = temp_dir / f"upscaled_{i:05d}.png"
            
            img = Image.open(input_f).convert("RGB")
            orig_w, orig_h = img.size
            
            # Recursive Loop per frame
            current_img = img
            current_scale = 1.0
            
            while current_scale < factor:
                # --- DIMENSION ALIGNMENT FIX ---
                # x2 latent upscaler: 64px, x4 upscaler: 8px
                alignment = 64 if use_x2_model else 8
                frame_w, frame_h = current_img.size
                pad_w = (alignment - (frame_w % alignment)) % alignment
                pad_h = (alignment - (frame_h % alignment)) % alignment
                
                if pad_w > 0 or pad_h > 0:
                    padded_w, padded_h = frame_w + pad_w, frame_h + pad_h
                    padded_img = Image.new("RGB", (padded_w, padded_h))
                    padded_img.paste(current_img, (0, 0))
                    if pad_w > 0:
                        right_edge = current_img.crop((frame_w - pad_w, 0, frame_w, frame_h))
                        padded_img.paste(right_edge.transpose(Image.FLIP_LEFT_RIGHT), (frame_w, 0))
                    if pad_h > 0:
                        bottom_edge = current_img.crop((0, frame_h - pad_h, frame_w, frame_h))
                        padded_img.paste(bottom_edge.transpose(Image.FLIP_TOP_BOTTOM), (0, frame_h))
                else:
                    padded_img = current_img
                
                upscaled_result = pipe(prompt="High quality", image=padded_img, num_inference_steps=15).images[0]
                
                # Crop back to target dimensions
                target_w_pass = int(frame_w * step_scale)
                target_h_pass = int(frame_h * step_scale)
                if upscaled_result.size != (target_w_pass, target_h_pass):
                    current_img = upscaled_result.crop((0, 0, target_w_pass, target_h_pass))
                else:
                    current_img = upscaled_result
                    
                current_scale *= step_scale
            
            # Final Resize
            target_w = int(orig_w * factor)
            target_h = int(orig_h * factor)
            if current_img.size != (target_w, target_h):
                current_img = current_img.resize((target_w, target_h), Image.LANCZOS)

            current_img.save(output_f)
            
            print(f"   Frame {i+1}/{frame_count} done.", end='\r')
            
        print(f"\n✅ All frames upscaled.")

        # 5. Stitch
        print("🔗 stitching video...")
        
        # Ensure output directory exists
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        # Get format-specific encoding params
        encoding_params = get_video_encoding_params(output_path)
        # Remove audio codec params since we're creating from images
        video_params = [p for i, p in enumerate(encoding_params) if not (encoding_params[i-1:i] == ["-c:a"] or p in ["aac", "libopus", "wmav2", "mp3"])]
        video_params = [p for p in video_params if p != "-c:a"]
        
        cmd = [
            "ffmpeg", "-y",
            "-framerate", str(fps),
            "-i", str(temp_dir / "upscaled_%05d.png"),
            *video_params,
            output_path
        ]
        subprocess.run(cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        
        shutil.rmtree(temp_dir)
        
        # Mux audio from source if present
        def has_audio_track(video_file):
            """Check if video file has an audio track using ffprobe."""
            try:
                result = subprocess.run([
                    "ffprobe", "-v", "error", "-select_streams", "a",
                    "-show_entries", "stream=codec_type", "-of", "csv=p=0",
                    video_file
                ], capture_output=True, text=True, timeout=10)
                return "audio" in result.stdout
            except:
                return False
        
        if has_audio_track(video_path):
            print(f"   🔗 Muxing audio from source...")
            temp_output = output_path + ".temp.mp4"
            os.rename(output_path, temp_output)
            
            mux_cmd = [
                "ffmpeg", "-y",
                "-i", temp_output,
                "-i", video_path,
                "-map", "0:v", "-map", "1:a",
                "-c:v", "copy", "-c:a", "aac", "-shortest",
                output_path, "-loglevel", "error"
            ]
            try:
                subprocess.run(mux_cmd, check=True)
                os.remove(temp_output)
            except:
                print("   ⚠️  Audio muxing failed, keeping silent video.")
                os.rename(temp_output, output_path)
        else:
            print(f"   ℹ️  No audio track in source, skipping mux.")
        
        print(f"✅ Upscaled video saved to {output_path}")
        return True

    except ImportError:
        print("❌ Missing dependencies (opencv-python, etc).")
        return False
    except Exception as e:
        print(f"❌ Video upscaling failed: {e}")
        return False

def parse_upscale_factor(val):
    """Parse upscale factor string (e.g., '2x', '4', '1.5') -> float."""
    if not val: return 2.0 # Default
    val = val.lower().strip().replace('x', '')
    try:
        f = float(val)
        if f <= 0: raise ValueError
        return f
    except:
        print(f"⚠️  Invalid upscale factor '{val}'. Using default 2.0x.")
        return 2.0

def convert_image_file(input_path, target):
    """Convert image format using PIL (no AI).
    
    Args:
        input_path: Source image file
        target: Output path, extension (.png), or format (png)
    """
    from PIL import Image
    from pathlib import Path
    
    # Parse target format
    target = target.strip()
    
    # Determine output path and format
    if '/' in target or '\\' in target or len(target) > 6:
        # It's a full path
        output_path = target
    elif target.startswith('.'):
        # It's an extension like ".png"
        name = Path(input_path).stem
        output_path = f"{name}{target}"
    else:
        # It's just a format like "png" or "PNG"
        name = Path(input_path).stem
        output_path = f"{name}.{target.lower()}"
    
    print(f"🔄 Converting Image: {input_path}")
    print(f"   Output: {output_path}")
    
    # Overwrite protection
    if Path(output_path).exists():
        if os.environ.get("AI_MEDIA_FORCE", "0") != "1":
            confirm = input(f"⚠️  '{output_path}' exists. Overwrite? [y/N]: ").strip().lower()
            if confirm != 'y':
                print("❌ Aborted.")
                return False
    
    try:
        img = Image.open(input_path)
        
        # Handle transparency for formats that don't support it
        output_ext = Path(output_path).suffix.lower()
        if output_ext in ['.jpg', '.jpeg'] and img.mode in ['RGBA', 'P']:
            print(f"   ℹ️  Converting RGBA → RGB (JPEG doesn't support transparency)")
            img = img.convert('RGB')
        
        # Ensure output directory exists
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        img.save(output_path)
        print(f"✅ Converted image saved to {output_path}")
        return True
    except Exception as e:
        print(f"❌ Conversion failed: {e}")
        return False

def convert_image_ffmpeg(input_path, target):
    """Convert image format using FFmpeg."""
    from pathlib import Path
    import subprocess
    
    target = target.strip()
    if '/' in target or '\\' in target or len(target) > 6:
        output_path = target
    elif target.startswith('.'):
        output_path = f"{Path(input_path).stem}{target}"
    else:
        output_path = f"{Path(input_path).stem}.{target.lower()}"
    
    print(f"🔄 Converting Image (FFmpeg): {input_path}")
    print(f"   Output: {output_path}")
    
    # Overwrite protection
    if Path(output_path).exists():
        if os.environ.get("AI_MEDIA_FORCE", "0") != "1":
            confirm = input(f"⚠️  '{output_path}' exists. Overwrite? [y/N]: ").strip().lower()
            if confirm != 'y':
                print("❌ Aborted.")
                return False
    
    try:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        subprocess.run(["ffmpeg", "-y", "-i", input_path, output_path], 
                      check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        print(f"✅ Converted image saved to {output_path}")
        return True
    except Exception as e:
        print(f"❌ Conversion failed: {e}")
        return False

def convert_video_file(input_path, target):
    """Convert video format using FFmpeg."""
    from pathlib import Path
    import subprocess
    
    target = target.strip()
    if '/' in target or '\\' in target or len(target) > 6:
        output_path = target
    elif target.startswith('.'):
        output_path = f"{Path(input_path).stem}{target}"
    else:
        output_path = f"{Path(input_path).stem}.{target.lower()}"
    
    print(f"🎬 Converting Video: {input_path}")
    print(f"   Output: {output_path}")
    
    # Overwrite protection
    if Path(output_path).exists():
        if os.environ.get("AI_MEDIA_FORCE", "0") != "1":
            confirm = input(f"⚠️  '{output_path}' exists. Overwrite? [y/N]: ").strip().lower()
            if confirm != 'y':
                print("❌ Aborted.")
                return False
    
    try:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        encoding_params = get_video_encoding_params(output_path)
        subprocess.run(["ffmpeg", "-y", "-i", input_path, *encoding_params, output_path], 
                      check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        print(f"✅ Converted video saved to {output_path}")
        return True
    except Exception as e:
        print(f"❌ Conversion failed: {e}")
        return False

def convert_audio_file(input_path, target):
    """Convert audio format using FFmpeg."""
    from pathlib import Path
    import subprocess
    
    target = target.strip()
    if '/' in target or '\\' in target or len(target) > 6:
        output_path = target
    elif target.startswith('.'):
        output_path = f"{Path(input_path).stem}{target}"
    else:
        output_path = f"{Path(input_path).stem}.{target.lower()}"
    
    print(f"🎵 Converting Audio: {input_path}")
    print(f"   Output: {output_path}")
    
    # Overwrite protection
    if Path(output_path).exists():
        if os.environ.get("AI_MEDIA_FORCE", "0") != "1":
            confirm = input(f"⚠️  '{output_path}' exists. Overwrite? [y/N]: ").strip().lower()
            if confirm != 'y':
                print("❌ Aborted.")
                return False
    
    try:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        subprocess.run(["ffmpeg", "-y", "-i", input_path, output_path], 
                      check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        print(f"✅ Converted audio saved to {output_path}")
        return True
    except Exception as e:
        print(f"❌ Conversion failed: {e}")
        return False


def convert_document_file(input_path, target):
    """Convert document format using MD as intermediate hub.
    
    Supported formats: md, html, pdf, docx, rtf, txt, json
    """
    from pathlib import Path
    import json as json_module
    
    SUPPORTED_FORMATS = ["md", "html", "pdf", "docx", "rtf", "txt", "json", "xhtml"]
    
    # Determine output path and format
    target = target.strip().lower()
    if '/' in target or '\\' in target:
        output_path = target
        output_format = Path(target).suffix.lstrip('.').lower()
    elif target.startswith('.'):
        output_path = f"{Path(input_path).stem}{target}"
        output_format = target.lstrip('.').lower()
    else:
        output_path = f"{Path(input_path).stem}.{target}"
        output_format = target.lower()
    
    if output_format not in SUPPORTED_FORMATS:
        print(f"❌ Unsupported output format: {output_format}")
        print(f"   Supported: {', '.join(SUPPORTED_FORMATS)}")
        return False
    
    # Determine input format
    input_format = Path(input_path).suffix.lstrip('.').lower()
    if input_format not in SUPPORTED_FORMATS:
        print(f"❌ Unsupported input format: {input_format}")
        print(f"   Supported: {', '.join(SUPPORTED_FORMATS)}")
        return False
    
    print(f"📄 Converting Document: {input_path}")
    print(f"   {input_format.upper()} → {output_format.upper()}")
    
    # Overwrite protection
    if Path(output_path).exists():
        if os.environ.get("AI_MEDIA_FORCE", "0") != "1":
            confirm = input(f"⚠️  '{output_path}' exists. Overwrite? [y/N]: ").strip().lower()
            if confirm != 'y':
                print("❌ Aborted.")
                return False
    
    try:
        # Step 1: Convert input to Markdown (intermediate format)
        markdown_content = ""
        
        if input_format == "md":
            with open(input_path, "r", encoding="utf-8") as f:
                markdown_content = f.read()
        
        elif input_format in ["html", "xhtml"]:
            # HTML to Markdown using html2text
            try:
                import html2text
                h = html2text.HTML2Text()
                h.ignore_links = False
                h.ignore_images = False
                with open(input_path, "r", encoding="utf-8") as f:
                    markdown_content = h.handle(f.read())
            except ImportError:
                # Fallback: strip HTML tags with BeautifulSoup
                with open(input_path, "r", encoding="utf-8") as f:
                    soup = BeautifulSoup(f.read(), "html.parser")
                    markdown_content = soup.get_text()
                print("   ⚠️ html2text not installed, using basic text extraction")
        
        elif input_format == "docx":
            doc = docx.Document(input_path)
            lines = []
            for para in doc.paragraphs:
                if para.style.name.startswith('Heading 1'):
                    lines.append(f"# {para.text}")
                elif para.style.name.startswith('Heading 2'):
                    lines.append(f"## {para.text}")
                elif para.style.name.startswith('Heading 3'):
                    lines.append(f"### {para.text}")
                else:
                    lines.append(para.text)
            markdown_content = "\n\n".join(lines)
        
        elif input_format == "txt":
            with open(input_path, "r", encoding="utf-8") as f:
                markdown_content = f.read()
        
        elif input_format == "json":
            with open(input_path, "r", encoding="utf-8") as f:
                data = json_module.load(f)
            # Try common JSON article structures
            if isinstance(data, dict):
                markdown_content = data.get("content", "") or data.get("markdown", "") or data.get("text", "") or str(data)
            else:
                markdown_content = str(data)
        
        elif input_format == "pdf":
            # PDF text extraction (limited)
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(input_path)
                text_parts = []
                for page in reader.pages:
                    text_parts.append(page.extract_text() or "")
                markdown_content = "\n\n".join(text_parts)
                print("   ⚠️ PDF conversion extracts text only (formatting/images lost)")
            except ImportError:
                print("❌ PyPDF2 required for PDF reading. Install: pip install PyPDF2")
                return False
        
        elif input_format == "rtf":
            # RTF is complex - basic text extraction only
            try:
                from striprtf.striprtf import rtf_to_text
                with open(input_path, "r", encoding="utf-8") as f:
                    markdown_content = rtf_to_text(f.read())
                print("   ⚠️ RTF conversion extracts text only (formatting lost)")
            except ImportError:
                print("❌ striprtf required for RTF reading. Install: pip install striprtf")
                return False
        
        if not markdown_content.strip():
            print("❌ No content extracted from input file")
            return False
        
        # Step 2: Convert Markdown to output format (reuse ArticleGenerator logic)
        # Create a minimal ArticleGenerator just for saving
        class DocConverter:
            pass
        
        converter = DocConverter()
        converter._save_formatted = ArticleGenerator._save_formatted.__get__(converter, DocConverter)
        
        # Ensure directory exists
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        converter._save_formatted(markdown_content, output_path, output_format)
        return True
        
    except Exception as e:
        print(f"❌ Conversion failed: {e}")
        import traceback
        traceback.print_exc()
        return False

# --- Interactive Mode ---

def emoji(emoji_char, fallback=""):
    """Return emoji if terminal supports it, otherwise return fallback text.
    
    Args:
        emoji_char: The emoji string to display (e.g., "🎨 ")
        fallback: Text to use if emoji encoding fails (e.g., "")
    
    Returns:
        emoji_char if terminal can encode it, otherwise fallback
    """
    try:
        emoji_char.encode(sys.stdout.encoding or 'utf-8')
        return emoji_char
    except (UnicodeEncodeError, LookupError, AttributeError):
        return fallback

def clear_screen():
    """Clear terminal screen."""
    os.system('cls' if os.name == 'nt' else 'clear')

def show_header(title="AI-Media"):
    """Show interactive mode header."""
    print(f"\n{'═'*60}")
    print(f"{emoji('🎨 ', '')}{title}")
    print(f"{'═'*60}\n")

def run_self_command(cmd_string):
    """Run ai-media.py with the given command arguments (cross-platform safe).
    
    Accepts a command string like: -gd "path/to/file" -cm model
    Properly handles quoted paths on both Windows and Unix.
    """
    import subprocess
    import shlex
    
    script_path = os.path.abspath(__file__)
    
    # Display what we're running
    print(f"\n🚀 Running: ai-media.py {cmd_string}\n")
    
    # Parse the command string properly
    # On Windows, shlex.split with posix=False preserves quotes, so we strip them manually
    if os.name == 'nt':
        # Use posix=True even on Windows to properly handle quotes
        # But escape backslashes first to prevent them being treated as escape chars
        escaped = cmd_string.replace('\\', '\\\\')
        try:
            args = shlex.split(escaped, posix=True)
            # Restore backslashes in paths
            args = [arg.replace('\\\\', '\\') for arg in args]
        except ValueError:
            # Fallback: manually parse
            args = []
            current = ""
            in_quotes = False
            for char in cmd_string:
                if char == '"' and not in_quotes:
                    in_quotes = True
                elif char == '"' and in_quotes:
                    in_quotes = False
                elif char == ' ' and not in_quotes:
                    if current:
                        args.append(current)
                        current = ""
                else:
                    current += char
            if current:
                args.append(current)
    else:
        # Unix/Mac - standard shlex parsing works fine
        args = shlex.split(cmd_string)
    
    # Run with subprocess (handles paths correctly on all platforms)
    subprocess.run([sys.executable, script_path] + args)

# --- Interactive Navigation Helpers ---


def get_cursor_position():
    """Query cursor position (row, col). Returns None if not supported/timeout."""
    if os.name == 'nt': return None
    import sys, tty, termios, select, re
    
    fd = sys.stdin.fileno()
    old_settings = termios.tcgetattr(fd)
    try:
        tty.setraw(fd)
        sys.stdout.write("\033[6n")
        sys.stdout.flush()
        
        # Wait for potential response
        if select.select([sys.stdin], [], [], 0.1)[0]:
            # Read until R
            resp = ""
            while True:
                ch = sys.stdin.read(1)
                resp += ch
                if ch == 'R': break
            
            # Parse \033[<row>;<col>R
            match = re.search(r'\x1b\[(\d+);(\d+)R', resp)
            if match:
                return int(match.group(1)), int(match.group(2))
    except (termios.error, IOError):
        pass
    finally:
        termios.tcsetattr(fd, termios.TCSADRAIN, old_settings)
    return None

def set_raw_mode(fd):
    """Enter raw mode, return old settings."""
    if os.name == 'nt' or not sys.stdin.isatty(): return None
    import termios, tty
    try:
        old = termios.tcgetattr(fd)
        tty.setraw(fd)
        return old
    except termios.error:
        return None

def restore_mode(fd, old_settings):
    """Restore terminal to old settings."""
    if os.name == 'nt' or not old_settings or fd is None: return
    import termios
    termios.tcsetattr(fd, termios.TCSADRAIN, old_settings)

def get_key():
    """Read a single key press from stdin (cross-platform)."""
    if os.name == 'nt':  # Windows
        import msvcrt
        ch = msvcrt.getch()
        
        # Handle special keys (arrow keys, etc.)
        if ch in (b'\x00', b'\xe0'):
            ch2 = msvcrt.getch()
            if ch2 == b'H': return 'UP'
            if ch2 == b'P': return 'DOWN'
            if ch2 == b'G': return 'HOME'
            if ch2 == b'O': return 'END'
            if ch2 == b'I': return 'PAGE_UP'
            if ch2 == b'Q': return 'PAGE_DOWN'
            return ch2.decode('utf-8', errors='ignore')
        
        if ch == b'\r': return 'ENTER'
        if ch == b'\x03': raise KeyboardInterrupt
        return ch.decode('utf-8', errors='ignore')
    else:  # Unix/Mac
        import sys, tty, termios, select
        fd = sys.stdin.fileno()
        
        # Only set raw mode if we are actually interactive (stdout is TTY)
        # Otherwise (e.g. tests) we might be sharing stdin with parent, and setting raw mode breaks parent output
        is_interactive_mode = sys.stdout.isatty()
        
        old_settings = None
        if is_interactive_mode:
             old_settings = termios.tcgetattr(fd)
        
        try:
            if is_interactive_mode:
                tty.setraw(fd)
            
            # Helper to read N bytes
            def read_bytes(n=1):
                return os.read(fd, n).decode('utf-8', errors='ignore')

            ch = read_bytes(1)
            
            if ch == '\x1b':  # Escape sequence
                # Use select to check if more characters are waiting
                if not select.select([sys.stdin], [], [], 0.05)[0]:
                    return 'ESC'
                    
                ch2 = read_bytes(1)
                
                # Handle [ sequences
                if ch2 == '[':
                    ch3 = read_bytes(1)
                    
                    # SGR Mouse: \033[<0;10;20M
                    if ch3 == '<':
                        mouse_seq = ""
                        while True:
                            char = read_bytes(1)
                            if char in ('m', 'M'):
                                end_char = char
                                break
                            mouse_seq += char
                        
                        parts = mouse_seq.split(';')
                        if len(parts) >= 3:
                            btn = parts[0]
                            x = int(parts[1])
                            y = int(parts[2])
                            # Standard left click is 0, Right is 2
                            # Scroll Up is 64, Scroll Down is 65
                            if end_char == 'M':
                                if btn == '0': return ('MOUSE', x, y)
                                if btn == '64': return 'SCROLL_UP'
                                if btn == '65': return 'SCROLL_DOWN'
                            return None

                    if ch3 == 'A': return 'UP'
                    if ch3 == 'B': return 'DOWN'
                    if ch3 == 'C': return 'RIGHT'
                    if ch3 == 'D': return 'LEFT'
                    if ch3 == 'H': return 'HOME'
                    if ch3 == 'F': return 'END'
                    if ch3 in ['1', '4', '5', '6']:
                        ch4 = read_bytes(1)
                        if ch4 == '~':
                            if ch3 == '1': return 'HOME'
                            if ch3 == '4': return 'END'
                            if ch3 == '5': return 'PAGE_UP'
                            if ch3 == '6': return 'PAGE_DOWN'
                
                # Handle O sequences
                if ch2 == 'O':
                    ch3 = read_bytes(1)
                    if ch3 == 'H': return 'HOME'
                    if ch3 == 'F': return 'END'
                    
        finally:
            if is_interactive_mode and old_settings:
                termios.tcsetattr(fd, termios.TCSADRAIN, old_settings)
        
        if ch == '\r' or ch == '\n': return 'ENTER'
        if ch == '\x03': raise KeyboardInterrupt
        return ch

def prompt_menu(prompt, options, allow_back=True, default_index=0, page_size=15):
    """
    Show interactive menu with arrow key navigation, pagination, & MOUSE support.
    Returns: value of selected option or None (if Back/Exit)
    """
    # ... (header omitted, unchanged)
    # Prepare items list
    items = list(options)
    if allow_back:
        items.append(("⬅️  Back", None))
    elif not options:
        # Fallback if empty
        return None

    current_idx = default_index if 0 <= default_index < len(items) else 0
    start_idx = 0

    # Hide cursor
    print("\033[?25l", end="")
    
    # --- MOUSE SUPPORT START ---
    mouse_enabled = False
    menu_start_row = None
    
    # Enter persistent raw mode
    fd_raw = None
    old_raw = None
    if os.name != 'nt' and sys.stdout.isatty():
         fd_raw = sys.stdin.fileno()
         old_raw = set_raw_mode(fd_raw)

    if os.name != 'nt' and sys.stdout.isatty():
        # Enable Mouse Reporting (1000: basic, 1006: SGR extended)
        # Note: In raw mode we need \r for new lines usually, but here print handles it via end="" or default
        # But wait, print() uses \n. In raw mode \n is just LF. We need \r\n.
        # Python's print() writes to stdout. We haven't set stdout to raw, only stdin?
        # tty.setraw() sets the file descriptor. If we set stdin, does it affect stdout processing?
        # Usually Output processing (OPOST) is part of termios.
        # tty.setraw() disables OPOST. So \n becomes just LF.
        # We need to explicitly print \r.
        print("\033[?1000h\033[?1006h", end="", flush=True)
        # Query start row if possible
        pos = get_cursor_position()
        if pos:
            menu_start_row = pos[0] # ABS row (1-based)
            mouse_enabled = True
    # --- MOUSE SUPPORT END ---

    if prompt:
        # In raw mode, \n is just LF (line feed) without carriage return
        # Replace \n with \r\n to ensure proper line formatting
        prompt_formatted = prompt.replace('\n', '\r\n')
        print(f"{prompt_formatted}\r")
        # Count actual lines in prompt for menu_start_row adjustment
        prompt_lines = prompt.count('\n') + 1
        if mouse_enabled: menu_start_row += prompt_lines

    # ANSI constants
    UP = "\033[F"
    CLEAR_LINE = "\033[K"
    CYAN = "\033[96m" 
    RESET = "\033[0m"
    DIM = "\033[90m"

    # Reserve space for menu
    max_view_lines = min(len(items), page_size) + 3
    for _ in range(max_view_lines):
        print(f"\r") # Just newline with CR
    
    # Move cursor back up to start of menu
    print(UP * max_view_lines, end="", flush=True)
    
    # Query cursor position NOW that we are at the top of the menu
    # This ensures we account for any scrolling that happened during reservation
    if mouse_enabled:
        pos = get_cursor_position()
        if pos:
            menu_start_row = pos[0] # ABS row (1-based)
            # Adjust if prompt was printed?
            # get_cursor_position returns current row.
            # If we printed PROMPT then reserved lines, then moved up:
            # We are at the line AFTER the prompt (start of the items list + indicators).
            # So menu_start_row is exactly where the FIRST line of our drawing area is.
            pass

    try:
        while True:
            # Pagination Logic
            if current_idx < start_idx:
                start_idx = current_idx
            elif current_idx >= start_idx + page_size:
                start_idx = current_idx - page_size + 1
            
            end_idx = min(len(items), start_idx + page_size)
            visible_items = items[start_idx:end_idx]
            
            lines_printed = 0
            
            # --- MOUSE TARGET MAP ---
            # Map absolute screen rows to item indices
            # Only valid if we know menu_start_row reliably
            # Since reserving lines might scroll, we should really update start_row 
            # by re-querying, but that's slow.
            # Alternative: Assume we are at top row of reserved block relative to cursor?
            # We are currently at top of reserved block.
            
            # Render Up Indicator
            if start_idx > 0:
                print(f"{DIM}   ⬆️  ... ({start_idx} more above){RESET}{CLEAR_LINE}\r")
                lines_printed += 1
            
            # Render Menu Items
            for i, (label, val) in enumerate(visible_items):
                abs_index = start_idx + i
                is_selected = (abs_index == current_idx)
                prefix = " > " if is_selected else "   "
                number = f"{abs_index+1}." if abs_index < len(options) else "0."
                
                if is_selected:
                    line = f"{CYAN}{prefix}{number:<4}  {label}{RESET}"
                else:
                    line = f"{prefix}{number:<4}  {label}"
                
                print(f"{line}{CLEAR_LINE}\r")
                lines_printed += 1
            
            # Render Down Indicator
            if end_idx < len(items):
                remaining = len(items) - end_idx
                print(f"{DIM}   ⬇️  ... ({remaining} more below){RESET}{CLEAR_LINE}\r")
                lines_printed += 1
                
            # Clear remaining reserved lines (leave 1 for hint)
            extra_lines = (max_view_lines - 1) - lines_printed
            for _ in range(extra_lines):
                print(f"{CLEAR_LINE}\r")
            
            # Render Hint Footer
            hint_back = ", '0' for Back" if allow_back else ""
            print(f"{DIM}(Tip: 'Home'/'End' or 'g'/'G' for top/bottom{hint_back}){RESET}{CLEAR_LINE}\r")
            
            # Move cursor back up
            print(UP * max_view_lines, end="", flush=True)

            # Handle Input
            key = get_key()
            
            if key == 'SCROLL_UP':
                 current_idx = max(0, current_idx - 1)
                 # Adjust start_idx to scroll smoothly
                 if current_idx < start_idx: start_idx = current_idx
                 continue
            
            if key == 'SCROLL_DOWN':
                 current_idx = min(len(items)-1, current_idx + 1)
                 # Adjust start_idx
                 if current_idx >= start_idx + page_size:
                     start_idx = current_idx - page_size + 1
                 continue
            
            if isinstance(key, tuple) and key[0] == 'MOUSE':
                 # ('MOUSE', x, y)
                 mx, my = key[1], key[2]
                 if mouse_enabled and menu_start_row:
                     # Calculate clicked line relative to menu top
                     # PROBLEM: If screen scrolled, menu_start_row is outdated.
                     # We might need to refresh menu_start_row occasionally or check relative?
                     # A robust way: query DSR only on Click?
                     # Let's try: get_cursor_position *now* matches our write cursor (top of menu)
                     # So we can just query it once per loop or on key press?
                     # Querying on every loop is flicker-prone.
                     # Querying only on click is better.
                     
                     # Wait, we are at top of menu block when reading key.
                     # So get_cursor_position() NOW returns the row of the top of the menu!
                     # Let's get it right now to calibrate.
                     # But get_key was blocking. We already have the key.
                     # We can't query DSR *after* the click because the mouse event already happened at Y.
                     # We need to know where we *were*.
                     
                     # Hack: Assume menu stays put unless we print newlines outside loop.
                     # But we can query it *if* we are unsure.
                     
                     # Let's try to just use menu_start_row calculate at start.
                     # If inaccurate, user might mis-click.
                     # But we can update it: when we enter key loop, we are at top.
                     # We could update menu_start_row = get_cursor_position() at top of while true?
                     # It adds default 30-100ms latency per loop. Might feel sluggish.
                     
                     # Compromise: Re-query start_row if we detect a click?
                     # No, because click Y is absolute.
                     # Let's try just using the initial one. If it breaks on scroll,
                     # we can add a check. OR we update it once every loop?
                     pass
                     
                 # Basic Relative Logic if we assume static:
                 if menu_start_row:
                     # Calculate row offset
                     # 1 line for prompt (maybe) - prompt printed *before* max_view loops
                     # prompt_message printed at: menu_start_row - 1 (if prompt existed)
                     # Actually we printed prompt, THEN reserved space.
                     # So we are at prompt + 1 (or 0 if no prompt).
                     
                     rel_y = my - menu_start_row
                     
                     # Account for Up indicator
                     header_offset = 1 if start_idx > 0 else 0
                     
                     # Item index = (rel_y - header_offset)
                     clicked_item_idx = rel_y - header_offset
                     
                     if 0 <= clicked_item_idx < len(visible_items):
                         # Clicked on valid item line
                         current_idx = start_idx + clicked_item_idx
                         break # Select!
                     
                     # Handle Indicators?
                     if start_idx > 0 and rel_y == 0:
                         # Clicked Up arrow area
                         current_idx = max(0, current_idx - page_size)
                         continue
                     
                     # Handle Down
                     footer_row = len(visible_items) + header_offset
                     if end_idx < len(items) and rel_y == footer_row:
                         # Clicked Down arrow
                         current_idx = min(len(items)-1, current_idx + page_size)
                         continue

            elif key == 'UP':
                current_idx = (current_idx - 1) % len(items)
            elif key == 'DOWN':
                current_idx = (current_idx + 1) % len(items)
            elif key == 'PAGE_UP' or key == '[':
                current_idx = max(0, current_idx - page_size)
            elif key == 'PAGE_DOWN' or key == ']':
                current_idx = min(len(items) - 1, current_idx + page_size)
            elif key == 'HOME' or key == 'g':
                current_idx = 0
            elif key == 'END' or key == 'G':
                current_idx = len(items) - 1
            elif key == 'ENTER':
                # Confirm selection
                break
            elif key in ['1', '2', '3', '4', '5', '6', '7', '8', '9']:
                num = int(key)
                if 1 <= num <= len(options):
                    current_idx = num - 1
                    continue
            elif key == '0' and allow_back:
                 current_idx = len(items) - 1 
                 
    except KeyboardInterrupt:
        # Clean exit on CTRL+C
        print(RESET + "\n" * max_view_lines + "\r") # Move past menu
        if os.name != 'nt' and sys.stdout.isatty(): 
             print("\033[?1000l\033[?1006l", end="", flush=True) # Disable Mouse
             restore_mode(fd_raw, old_raw)
        print("\033[?25h", end="") # Show cursor
        return None
    finally:
        # Restore cursor & mouse
        print(RESET + "\n" * max_view_lines + "\r") # Move past menu
        if os.name != 'nt' and sys.stdout.isatty(): 
             print("\033[?1000l\033[?1006l", end="", flush=True) # Disable Mouse
             restore_mode(fd_raw, old_raw)
        print("\033[?25h", end="") # Show cursor

    selected_label, selected_val = items[current_idx]
    return selected_val

def prompt_choice(prompt, options, allow_back=True):
    """Wrapper for prompt_menu (backward usage compatibility)."""
    return prompt_menu(prompt, options, allow_back)

def prompt_text(prompt, default=None, required=True):
    """Get text input from user."""
    try:
        import readline  # Enable arrow key support for input()
    except ImportError:
        pass  # Windows doesn't have readline
    
    default_str = f" [{default}]" if default else ""
    while True:
        try:
            value = input(f"{prompt}{default_str}: ").strip()
            if not value and default:
                return default
            if not value and required:
                print("   This field is required.")
                continue
            return value
        except KeyboardInterrupt:
            return None



def check_overwrite(filepath, always_overwrite=False, never_overwrite=False):
    """
    Check if file exists and prompt user.
    Returns: (should_write, final_path, always_overwrite, never_overwrite)
    """
    if never_overwrite:
        return False, filepath, False, True
        
    if not os.path.exists(filepath) or always_overwrite:
        return True, filepath, always_overwrite, False
        
    print(f"\n⚠️  File already exists: {filepath}")
    choice = prompt_choice("Overwrite?", [
        ("Yes", "y"),
        ("No (skip file)", "n"), 
        ("Always (overwrite all remaining)", "a"),
        ("Never (skip all remaining)", "v"),
        ("Rename (auto-increment)", "r")
    ])
    
    if choice == "y":
        return True, filepath, False, False
    elif choice == "a":
        return True, filepath, True, False
    elif choice == "v":
        print(f"⏭️  Skipping {os.path.basename(filepath)} (and all remaining)")
        return False, filepath, False, True
    elif choice == "r":
        # Auto-increment rename
        base, ext = os.path.splitext(filepath)
        counter = 1
        new_path = f"{base}_{counter}{ext}"
        while os.path.exists(new_path):
            counter += 1
            new_path = f"{base}_{counter}{ext}"
        print(f"📝 Renaming to: {new_path}")
        return True, new_path, False, False
    elif choice is None: # Back/Cancel
        print("❌ Operation cancelled.")
        return False, None, False, False # Signal abort with None path
    else:
        print(f"⏭️  Skipping {os.path.basename(filepath)}")
        return False, filepath, False, False

def browse_files(start_dir="."):
    """Interactively browse file system and return selected file path."""
    current_dir = os.path.abspath(start_dir)
    
    while True:
        try:
            items = os.listdir(current_dir)
        except PermissionError:
            print(f"❌ Permission denied: {current_dir}")
            current_dir = os.path.dirname(current_dir)
            continue
            
        # Separate dirs and files
        dirs = []
        files = []
        for item in items:
            if item.startswith('.'): continue # Skip hidden
            full_path = os.path.join(current_dir, item)
            if os.path.isdir(full_path):
                dirs.append(item)
            else:
                files.append(item)
        
        dirs.sort()
        files.sort()
        
        menu_items = []
        # Add parent directory option if not at root
        if os.path.dirname(current_dir) != current_dir:
            menu_items.append(("📂 .. (Up Directory)", ".."))
        
        for d in dirs:
            menu_items.append((f"📁 {d}/", d))
            
        for f in files:
            menu_items.append((f"📄 {f}", f))
        
        # Display Menu
        print(f"\n📂 Location: {current_dir}")
        choice = prompt_choice(None, menu_items, allow_back=True)
        
        if choice is None: # Back/Cancel
            return None
        
        if choice == "..":
            current_dir = os.path.dirname(current_dir)
        else:
            selected_path = os.path.join(current_dir, choice)
            if os.path.isdir(selected_path):
                current_dir = selected_path
            else:
                return selected_path

def prompt_file(prompt, must_exist=True):
    """Get file path input from user with browsing support."""
    while True:
        # If looking for existing file, offer browse option
        if must_exist:
            print(f"\n{prompt}")
            options = [
                ("📂 Browse Files", "browse"),
                ("⌨️  Enter Path Manually", "manual"),
                ("🔙 Cancel", None)
            ]
            method = prompt_choice(None, options, allow_back=False) # We handle None manually
            
            if method is None:
                return None
            
            if method == "browse":
                path = browse_files()
                if path:
                    return path
                continue # If cancelled browsing, return to method choice
                
            elif method == "manual":
                pass # Fall through to manual input
        
        # Manual Input Logic
        try:
            path = input(f"Enter file path: ").strip()
            if not path:
                print("   This field is required.")
                continue
            if must_exist and not os.path.exists(path):
                print(f"   File not found: {path}")
                continue
            return path
        except KeyboardInterrupt:
            return None

def run_interactive(jump_point=None):
    """Run interactive mode.
    
    Args:
        jump_point: Optional jump path (e.g., 'image/sdxl', 'audio/bark')
    """
    
    # Jump point mappings: name -> (menu_action, submenu_value)
    JUMP_POINTS = {
        # By name
        'image': ('image', None),
        'image/sdxl': ('image', 'sdxl'),
        'image/sd15': ('image', 'sd-1.5'),
        'image/flux': ('image', 'flux'),
        'image/flux-dev': ('image', 'flux-dev'),
        'video': ('video', None),
        'video/zeroscope': ('video', 'zeroscope'),
        'video/modelscope': ('video', 'ms-1.7b'),
        'video/cogvideox': ('video', 'cogvideox'),
        'video/svd': ('video', 'svd'),
        'audio': ('audio', None),
        'audio/musicgen': ('audio', 'musicgen-medium'),
        'audio/musicgen-small': ('audio', 'musicgen-small'),
        'audio/musicgen-large': ('audio', 'musicgen-large'),
        'audio/audioldm2': ('audio', 'audioldm2'),
        'audio/bark': ('audio', 'bark'),
        'caption': ('caption', None),
        'article': ('article', None),
        'article/offline': ('article', 'offline'),
        'article/online': ('article', 'online'),
        'code': ('code', None),
        'chat': ('chat', None),
        'research': ('article', 'online'),  # Alias for deep research/online article
        'transform': ('transform', None),
        'transform/edit': ('transform', 'edit'),
        'transform/rembg': ('transform', 'rembg'),
        'transform/silhouette': ('transform', 'silhouette'),
        'upscale': ('upscale', None),
        'convert': ('convert', None),
        'test': ('test', None),
        'test/unit': ('test', 'unit'),
        'test/integration': ('test', 'integration'),
        'test/codec': ('test', 'codec'),
        'sysinfo': ('sysinfo', None),
        # By number (matching menu order)
        '1': ('image', None),
        '1/1': ('image', 'sdxl'),
        '1/2': ('image', 'sd-1.5'),
        '1/3': ('image', 'flux'),
        '1/4': ('image', 'flux-dev'),
        '2': ('video', None),
        '2/1': ('video', 'zeroscope'),
        '2/2': ('video', 'ms-1.7b'),
        '2/3': ('video', 'cogvideox'),
        '2/4': ('video', 'svd'),
        '3': ('audio', None),
        '3/1': ('audio', 'musicgen-medium'),
        '3/2': ('audio', 'musicgen-small'),
        '3/3': ('audio', 'musicgen-large'),
        '3/4': ('audio', 'audioldm2'),
        '3/5': ('audio', 'bark'),
        '4': ('caption', None),
        '5': ('article', None),
        '5/1': ('article', 'offline'),
        '5/2': ('article', 'online'),
        '6': ('code', None),
        '7': ('chat', None),
        '8': ('transform', None),
        '8/1': ('transform', 'edit'),
        '8/2': ('transform', 'rembg'),
        '8/3': ('transform', 'silhouette'),
        '9': ('convert', None),
        '10': ('upscale', None),
        '11': ('test', None),
        '11/1': ('test', 'unit'),
        '11/2': ('test', 'integration'),
        '11/3': ('test', 'codec'),
        '12': ('sysinfo', None),
    }
    
    # Parse jump point
    initial_action = None
    initial_model = None
    if jump_point and jump_point != 'menu':
        jp_lower = jump_point.lower()
        if jp_lower in JUMP_POINTS:
            initial_action, initial_model = JUMP_POINTS[jp_lower]
        else:
            print(f"⚠️  Unknown jump point: {jump_point}")
            print("   Run with --help or see README for valid jump points.")
    
    def system_info_menu():
        """Display system information."""
        clear_screen()
        show_header("System Information")
        
        # Display Loading Indicator
        print("\n⏳ Loading system information...", end="", flush=True)
        
        import platform
        import psutil
        import torch
        
        import subprocess
        
        # OS Info
        os_name = platform.system()
        os_release = platform.release()
        os_ver = platform.version()
        if os_name == "Darwin":
            mac_ver = platform.mac_ver()[0]
            os_info = f"macOS {mac_ver} ({platform.machine()})"
        elif os_name == "Windows":
             # platform.platform() gives "Windows-10-...", platform.release() gives "10" or "11"
            os_info = f"{platform.system()} {platform.release()} (Build {platform.version()})"
        else:
            os_info = f"{platform.system()} {platform.release()} ({platform.machine()})"
        
        # CPU Info
        cpu_count = psutil.cpu_count(logical=True)
        cpu_percent = psutil.cpu_percent(interval=0.1)
        
        # CPU Model Name
        cpu_model = platform.processor()
        try:
            if os_name == "Windows":
                # Windows CPU Name (Registry)
                import winreg
                key = winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE, r"HARDWARE\DESCRIPTION\System\CentralProcessor\0")
                cpu_model = winreg.QueryValueEx(key, "ProcessorNameString")[0].strip()
            elif os_name == "Darwin":
                # Mac CPU Name (sysctl)
                result = subprocess.run(["sysctl", "-n", "machdep.cpu.brand_string"], capture_output=True, text=True)
                if result.returncode == 0:
                    cpu_model = result.stdout.strip()
            elif os_name == "Linux":
                # Linux CPU Name (/proc/cpuinfo)
                if os.path.exists("/proc/cpuinfo"):
                    with open("/proc/cpuinfo", "r") as f:
                        for line in f:
                            if "model name" in line:
                                cpu_model = line.split(":", 1)[1].strip()
                                break
        except:
            pass
        
        # RAM Info
        mem = psutil.virtual_memory()
        ram_total = f"{mem.total / (1024**3):.1f} GB"
        ram_used = f"{mem.used / (1024**3):.1f} GB"
        ram_avail = f"{mem.available / (1024**3):.1f} GB"
        ram_percent = f"{mem.percent}%"
        
        # GPU Info
        if torch.backends.mps.is_available():
            gpu_info = "MPS (Apple Silicon) ✅ Available"
            try:
                # [NEW] Mac MPS Stats
                mem_curr = torch.mps.current_allocated_memory() / (1024**3)
                mem_driver = torch.mps.driver_allocated_memory() / (1024**3)
                
                indent = " " * 13
                if mem_curr > 0.05 or mem_driver > 0.05:
                    gpu_info += f"\n{indent}Allocs: {mem_curr:.2f} GB Current | {mem_driver:.2f} GB Driver"
            except:
                pass
        elif torch.cuda.is_available():
            gpu_name = torch.cuda.get_device_name(0)
            vram = torch.cuda.get_device_properties(0).total_memory / (1024**3)
            
            # Fetch Nvidia Driver Info
            driver_info = ""
            try:
                # Run nvidia-smi to get system-level driver info
                # Output example: | NVIDIA-SMI 531.14  Driver Version: 531.14  CUDA Version: 12.1 |
                res = subprocess.run("nvidia-smi", capture_output=True, text=True)
                if res.returncode == 0:
                    import re
                    driver_match = re.search(r"Driver Version:\s*([\d\.]+)", res.stdout)
                    cuda_match = re.search(r"CUDA Version:\s*([\d\.]+)", res.stdout)
                    
                    parts = []
                    if driver_match: parts.append(f"NVIDIA Driver {driver_match.group(1)}")
                    if cuda_match: parts.append(f"CUDA {cuda_match.group(1)}")
                    
                    if parts:
                        # Format on next line with indentation (13 spaces for "🎮 GPU:      ")
                        indent = " " * 13
                        driver_info = f"\n{indent}{', '.join(parts)}"
            
            except:
                pass
            
            # Fetch Real-time GPU Stats (VRAM & Load)
            gpu_stats = ""
            try:
                # Output: 2569, 24576, 4
                stat_cmd = "nvidia-smi --query-gpu=memory.used,memory.total,utilization.gpu --format=csv,noheader,nounits"
                stat_res = subprocess.run(stat_cmd, capture_output=True, text=True)
                if stat_res.returncode == 0:
                    parts = stat_res.stdout.strip().split(',')
                    if len(parts) >= 3:
                        used_mb = int(parts[0])
                        total_mb = int(parts[1])
                        load_pct = int(parts[2])
                        
                        avail_gb = (total_mb - used_mb) / 1024
                        total_gb = total_mb / 1024
                        used_pct = (used_mb / total_mb) * 100
                        
                        indent = " " * 13
                        gpu_stats = f"\n{indent}Memory: {avail_gb:.1f} GB Available / {total_gb:.1f} GB Total ({used_pct:.1f}% Used) | Load: {load_pct}%"
            except:
                pass
            
            gpu_info = f"CUDA ({gpu_name}, {vram:.1f} GB VRAM) ✅ Available{driver_info}{gpu_stats}"
        else:
            gpu_info = "CPU Only (No Acceleration Detected)"
            
        # Clear Loading Indicator (Overwrite line)
        print("\r" + " " * 50 + "\r", end="", flush=True)
            
        print(f"💻 OS:       {os_info}")
        print(f"🧠 CPU:      {cpu_model} | {cpu_count} Cores (Usage: {cpu_percent}%)")
        print(f"💾 RAM:      {ram_avail} Available / {ram_total} Total ({ram_percent} Used)")
        print(f"🎮 GPU:      {gpu_info}")
        print()
        
        prompt_menu(None, [], allow_back=True)

    def main_menu():
        """Show main menu and return action."""
        clear_screen()
        show_header("AI-Media Interactive Mode")
        print("📋 What would you like to do?\n")
        
        options = [
            ("🖼️   Generate Image", "image"),
            ("🎬  Generate Video", "video"),
            ("🎵  Generate Audio", "audio"),
            ("📝  Generate Description", "caption"),
            ("📰  Generate Article", "article"),
            ("💻  Generate Code", "code"),
            ("💬  Chat", "chat"),
            ("✨  Transform/Edit Image", "transform"),
            ("🔄  Convert Media", "convert"),
            ("📄  Convert Document", "doc_convert"),
            ("📈  Upscale Media", "upscale"),
            ("🧪  Run Tests", "test"),
            ("ℹ️   System Information", "sysinfo"),
            ("❌  Exit", None)
        ]
        
        return prompt_choice(None, options, allow_back=False)
    
    def image_menu(preset_model=None):
        """Image generation submenu."""
        clear_screen()
        show_header("Image Generation")
        
        # Model selection (skip if preset)
        if preset_model:
            model = preset_model
            print(f"📦 Model: {model}\n")
        else:
            print("📦 Select Model:\n")
            # Build model options - hide Mac-specific notes on Windows
            if os.name == 'nt':  # Windows
                model_options = [
                    ("SDXL Turbo (Default, Fast) ~8GB", "sdxl"),
                    ("SD 1.5 (Lightweight) ~4GB", "sd-1.5"),
                    ("Flux Schnell (High Quality) ~12GB", "flux"),
                    ("Flux Dev (Professional) ~16GB", "flux-dev"),
                ]
            else:  # Mac/Linux
                model_options = [
                    ("SDXL Turbo (Default, Fast) ~8GB", "sdxl"),
                    ("SD 1.5 (Lightweight) ~4GB", "sd-1.5"),
                    ("Flux Schnell (High Quality, Slow on Mac) ~12GB", "flux"),
                    ("Flux Dev (Professional, Very Slow on Mac) ~16GB", "flux-dev"),
                ]
            model = prompt_choice("Model", model_options)
            if model is None:
                return
        
        # Prompt
        print()
        prompt = prompt_text("📝 Enter prompt")
        if prompt is None:
            return
        
        # Resolution
        print("\n📐 Select Resolution:\n")
        size_options = [
            ("720p (1280x720)", "720p"),
            ("1080p (1920x1080)", "1080p"),
            ("4K (3840x2160)", "4k"),
            ("64x64 (Quick Test)", "64x64"),
            ("Custom Resolution", "custom"),
        ]
        size = prompt_choice("Size", size_options)
        if size is None:
            return
            
        if size == "custom":
            print()
            size = prompt_text("Enter resolution (e.g. 512x512, 1024x768)")
            if not size:
                return
        
        # Orientation
        print("\n🔄 Select Orientation:\n")
        orient_options = [
            ("Landscape (Default)", "landscape"),
            ("Portrait", "portrait"),
            ("Square", "square"),
        ]
        orientation = prompt_choice("Orientation", orient_options)
        if orientation is None:
            return
        
        # Output
        print()
        output = prompt_text("💾 Output filename (or press Enter for auto)", required=False)
        
        # Build and run command
        cmd = f"-i -p \"{prompt}\" -s {size} -otn {orientation} --image-model {model}"
        if output:
            cmd += f" -o \"{output}\""
        
        run_self_command(cmd)
        input("\nPress Enter to continue...")
    
    def video_menu(preset_model=None):
        """Video generation submenu."""
        clear_screen()
        show_header("Video Generation")
        
        # Model selection (skip if preset)
        if preset_model:
            model = preset_model
            print(f"📦 Model: {model}\n")
        else:
            print("📦 Select Model:\n")
            model_options = [
                ("Zeroscope (Default, No Watermarks)", "zeroscope"),
                ("ModelScope (General Purpose, Has Watermarks)", "ms-1.7b"),
                ("CogVideoX (State of the Art, Slow) ~15GB", "cogvideox"),
                ("Wan 2.2 (Alibaba, 14B, High Quality) ~24GB", "wan-2.2"),
                ("LTX-Video (Lightricks, Fast DiT) ~16GB", "ltx-video"),
                ("Mochi 1 (Genmo, Motion SOTA) ~19GB", "mochi-1"),
                ("HunyuanVideo (Tencent, Cinematic) ~24GB", "hunyuan"),
                ("Stable Video Diffusion (Image-to-Video only)", "svd"),
            ]
            model = prompt_choice("Model", model_options)
            if model is None:
                return
        
        # Prompt or Input Image (for SVD)
        prompt = None
        input_image = None
        
        if model == 'svd':
             print("\n🖼️ Select Input Image for SVD:")
             input_image = prompt_file("Input Image")
             if not input_image:
                 return
        else:
            print()
            prompt = prompt_text("📝 Enter prompt")
            if prompt is None:
                return
        
        # Duration
        print("\n⏱️ Select Duration:\n")
        length_options = [
            ("2 seconds (Quick)", "2s"),
            ("5 seconds", "5s"),
            ("10 seconds", "10s"),
            ("Custom Duration", "custom"),
        ]
        
        # Build and run command (partial logic needed inside/after check)
        # We need duration first.
        
        length = prompt_choice("Duration", length_options)
        if length is None:
            return
            
        if length == "custom":
            print()
            length = prompt_text("Enter duration (e.g. 3s, 15s)")
            if not length:
                return

        # Resolution (for Zeroscope dynamic upscaling)
        size = None
        if model == 'zeroscope':
            print("\n📐 Select Output Resolution:\n")
            
            # Check if MPS (XL is skipped on Apple Silicon)
            import torch
            is_mps = torch.backends.mps.is_available() and not torch.cuda.is_available()
            
            if is_mps:
                print("   💡 Zeroscope uses Real-ESRGAN upscaling (XL skipped on Mac)")
                size_options = [
                    ("576x320 (Native, Fast)", "576x320"),
                    ("1024x576 (ESRGAN)", "1024x576"),
                    ("720p (ESRGAN)", "720p"),
                    ("1080p (ESRGAN)", "1080p"),
                    ("Custom Resolution", "custom"),
                ]
            else:
                print("   💡 Zeroscope uses smart upscaling for higher resolutions")
                size_options = [
                    ("576x320 (Native, Fast)", "576x320"),
                    ("1024x576 (XL Upscale)", "1024x576"),
                    ("720p (XL + ESRGAN)", "720p"),
                    ("1080p (XL + ESRGAN)", "1080p"),
                    ("Custom Resolution", "custom"),
                ]
            size = prompt_choice("Size", size_options)
            if size is None:
                return
            if size == "custom":
                print()
                size = prompt_text("Enter resolution (e.g. 1280x720)")
                if not size:
                    return

        # Input Image (Optional - for non-SVD models to enable Image-to-Video)
        if model != 'svd':
            print()
            add_image = input("📂 Add input image for Image-to-Video? [y/N]: ").strip().lower()
            if add_image in ['y', 'yes']:
                input_image = prompt_file("Input Image")

        # Audio Prompt (Optional - for Video-with-Audio, available for ALL models)
        print()
        audio_prompt = prompt_text("🎵 Audio prompt (for background music, Optional)", required=False)
        audio_model = None
        if audio_prompt:
             print("\n📦 Select Audio Model for Background:\n")
             audio_model_options = [
                ("MusicGen Medium (Default)", "musicgen-medium"),
                ("MusicGen Small (Fast)", "musicgen-small"),
                ("AudioLDM2 (Sound Effects)", "audioldm2"),
             ]
             audio_model = prompt_choice("Audio Model", audio_model_options)

        # Output
        print()
        output = prompt_text("💾 Output filename (or press Enter for auto)", required=False)
        
        # Build Command
        cmd = f"-v -l {length} --video-model {model}"
        if size:
            cmd += f" -s {size}"
        if prompt:
            cmd += f" -p \"{prompt}\""
        if input_image:
            cmd += f" -ii \"{input_image}\""
        if audio_prompt:
            cmd += f" -ap \"{audio_prompt}\""
            if audio_model:
                cmd += f" -am {audio_model}"
        if output:
             cmd += f" -o \"{output}\""
             
        run_self_command(cmd)
        input("\nPress Enter to continue...")

    
    def audio_menu(preset_model=None):
        """Audio generation submenu."""
        clear_screen()
        show_header("Audio Generation")
        
        # Model selection (skip if preset)
        if preset_model:
            model = preset_model
            print(f"📦 Model: {model}\n")
        else:
            print("📦 Select Model:\n")
            model_options = [
                ("MusicGen Medium (Default)", "musicgen-medium"),
                ("MusicGen Small (Fast)", "musicgen-small"),
                ("MusicGen Large (High Quality)", "musicgen-large"),
                ("AudioLDM2 (Sound Effects)", "audioldm2"),
                ("Bark (Speech/TTS)", "bark"),
            ]
            model = prompt_choice("Model", model_options)
            if model is None:
                return
        
        # Prompt
        print()
        if model == "bark":
            prompt = prompt_text("📝 Enter text to speak")
        else:
            prompt = prompt_text("📝 Enter audio description")
        if prompt is None:
            return
        
        # Duration (not for Bark)
        length = "10s"
        if model != "bark":
            print("\n⏱️ Select Duration:\n")
            length_options = [
                ("5 seconds", "5s"),
                ("10 seconds", "10s"),
                ("30 seconds", "30s"),
                ("Custom Duration", "custom"),
            ]
            length = prompt_choice("Duration", length_options)
            if length is None:
                return
            if length == "custom":
                print()
                length = prompt_text("Enter duration (e.g. 8s, 1m)")
                if not length:
                    return
        
        # Output
        print()
        output = prompt_text("💾 Output filename (or press Enter for auto)", required=False)
        
        # Build and run command
        cmd = f"-a -p \"{prompt}\" -l {length} --audio-model {model}"
        if output:
            cmd += f" -o \"{output}\""
        
        run_self_command(cmd)
        input("\nPress Enter to continue...")
    
    def transform_menu(preset_operation=None):
        """Image transformation submenu."""
        clear_screen()
        show_header("Transform/Edit Image")
        
        # Input file
        print("📂 Select input image:\n")
        input_file = prompt_file("Input Image")
        if input_file is None:
            return
        
        # Operation (skip if preset)
        if preset_operation:
            operation = preset_operation
            print(f"✨ Operation: {operation}\n")
        else:
            print("\n✨ Select Operation:\n")
            op_options = [
                ("Creative Edit (AI Instruction)", "edit"),
                ("Remove Background", "rembg"),
                ("Create Silhouette", "silhouette"),
            ]
            operation = prompt_choice("Operation", op_options)
            if operation is None:
                return
        
        # Build command based on operation
        if operation == "edit":
            print()
            instruction = prompt_text("📝 Enter edit instruction (e.g., 'Make it anime')")
            if instruction is None:
                return
            cmd = f"-ti \"{input_file}\" -tp \"{instruction}\""
        elif operation == "rembg":
            cmd = f"-ti \"{input_file}\" --remove-background"
        elif operation == "silhouette":
            cmd = f"-ti \"{input_file}\" --remove-background --silhouette"
        
        # Output
        print()
        output = prompt_text("💾 Output filename (or press Enter for auto)", required=False)
        if output:
            cmd += f" -o \"{output}\""
        
        run_self_command(cmd)
        input("\nPress Enter to continue...")
    
    def upscale_menu():
        """Upscale media submenu."""
        clear_screen()
        show_header("Upscale Media")
        
        # Media type
        print("📁 Select Media Type:\n")
        type_options = [
            ("Image", "image"),
            ("Video", "video"),
        ]
        media_type = prompt_choice("Type", type_options)
        if media_type is None:
            return
        
        # Input file
        print("\n📂 Select input file:\n")
        input_file = prompt_file("Enter file path")
        if input_file is None:
            return
        
        # Upscale factor
        print("\n📈 Select Upscale Factor:\n")
        factor_options = [
            ("2x (Default)", "2x"),
            ("4x", "4x"),
            ("Custom Factor", "custom"),
        ]
        factor = prompt_choice("Factor", factor_options)
        if factor == "custom":
            print()
            factor = prompt_text("Enter factor (e.g. 3x, 8x)")
            if not factor:
                return
        if factor is None:
            return
        
        # Method
        print("\n⚙️ Select Method:\n")
        method_options = [
            ("AI Upscale (High Quality)", "ai"),
            ("Simple Upscale (Fast, Lanczos)", "simple"),
        ]
        method = prompt_choice("Method", method_options)
        if method is None:
            return
        
        ai_model = "realesrgan"
        video_codec = "auto"
        
        if method == "ai":
            # Select Model
            print("\n🤖 Select AI Model:\n")
            model_options = [
                ("Real-ESRGAN (Fast, Faithful)", "realesrgan"),
                ("Stable Diffusion (Slow, Creative)", "sd"),
            ]
            ai_model = prompt_choice("Model", model_options)
            if ai_model is None:
                return

            if media_type == "video":
                # Select Codec for Video
                print("\n🎥 Select Video Codec:\n")
                codec_options = [
                    ("Auto (Default)", "auto"),
                    ("H.264", "h264"),
                    ("HEVC (H.265)", "hevc"),
                    ("AV1", "av1"),
                ]
                video_codec = prompt_choice("Codec", codec_options)
                if video_codec is None:
                    return

        # Build command
        if media_type == "image":
            cmd = f"-ui \"{input_file}\" -uf {factor}"
            if method == "ai":
                cmd += f" -iu {ai_model}"
            else:
                cmd += " -su"
        else:
            cmd = f"-uv \"{input_file}\" -uf {factor}"
            if method == "ai":
                cmd += f" -vu {ai_model}"
                if video_codec != "auto":
                    cmd += f" -vc {video_codec}"
            else:
                cmd += " -su"
        
        run_self_command(cmd)
        input("\nPress Enter to continue...")
    
    def convert_menu():
        """Convert media submenu."""
        clear_screen()
        show_header("Convert Media")
        
        # Media type
        print("📁 Select Media Type:\n")
        type_options = [
            ("Image", "image"),
            ("Video", "video"),
            ("Audio", "audio"),
        ]
        media_type = prompt_choice("Type", type_options)
        if media_type is None:
            return
        
        # Input file
        print("\n📂 Select input file:\n")
        input_file = prompt_file("Enter file path")
        if input_file is None:
            return
        
        # Target format
        print("\n🎯 Select Target Format:\n")
        if media_type == "image":
            format_options = [
                ("PNG", "png"),
                ("JPG", "jpg"),
                ("WebP", "webp"),
            ]
        elif media_type == "video":
            format_options = [
                ("MP4", "mp4"),
                ("WebM", "webm"),
                ("AVI", "avi"),
            ]
        else:
            format_options = [
                ("MP3", "mp3"),
                ("WAV", "wav"),
                ("FLAC", "flac"),
            ]
        target_format = prompt_choice("Format", format_options)
        if target_format is None:
            return
        
        # Build command
        if media_type == "image":
            cmd = f"-ci \"{input_file}\" -cit {target_format}"
        elif media_type == "video":
            cmd = f"-cv \"{input_file}\" -cvt {target_format}"
        else:
            cmd = f"-ca \"{input_file}\" -cat {target_format}"
        
        run_self_command(cmd)
        input("\nPress Enter to continue...")
    
    def document_convert_menu():
        """Convert document format submenu."""
        clear_screen()
        show_header("Convert Document")
        
        # Input file
        print("📂 Select input document:\n")
        input_file = prompt_file("Enter file path")
        if input_file is None:
            return
        
        # Target format
        print("\n🎯 Select Target Format:\n")
        format_options = [
            ("PDF - Portable Document", "pdf"),
            ("DOCX - Microsoft Word", "docx"),
            ("HTML - Web Page", "html"),
            ("MD - Markdown", "md"),
            ("RTF - Rich Text Format", "rtf"),
            ("TXT - Plain Text", "txt"),
            ("JSON - Structured Data", "json"),
        ]
        target_format = prompt_choice("Format", format_options)
        if target_format is None:
            return
        
        # Build command
        cmd = f"-cd \"{input_file}\" -cdt {target_format}"
        
        run_self_command(cmd)
        input("\nPress Enter to continue...")
    
    def caption_menu(preset_model=None):
        """Generate caption submenu."""
        clear_screen()
        show_header("Generate Description")
        
        # Input file
        print("📂 Select input image or video:\n")
        input_file = prompt_file("Enter file path")
        if input_file is None:
            return
        
        # Model (skip if preset)
        if preset_model:
            model = preset_model
            print(f"📦 Model: {model}\n")
        else:
            print("\n📦 Select Caption Model:\n")
            model_options = [
                ("Florence-2 (Default, SOTA)", "florence"),
                ("BLIP", "blip"),
            ]
            model = prompt_choice("Model", model_options)
            if model is None:
                return
        
        # Build command
        cmd = f"-gd \"{input_file}\" -cm {model}"
        
        run_self_command(cmd)
        input("\nPress Enter to continue...")
    
    def article_menu(preset_mode=None):
        """Generate article/research submenu."""
        while True:
            clear_screen()
            show_header("Generate Article")
            
            # Mode selection (Offline/Online) - skip if preset
            if preset_mode:
                mode = preset_mode
            else:
                print("📝 Select article mode:\n")
                mode_options = [
                    ("Offline (Model's internal knowledge)", "offline"),
                    ("Online Research (Live web search)", "online"),
                ]
                mode = prompt_choice("Mode", mode_options)
                if mode is None:
                    return
            
            online = mode == "online"
            mode_label = "Deep Research (Online)" if online else "Article (Offline)"
            
            clear_screen()
            show_header(mode_label)
            
            # Topic/Prompt
            print("✏️  Enter your topic or prompt:\n")
            topic = prompt_text("Topic")
            if not topic:
                return
            
            # Model selection
            print("\n📦 Select Model:\n")
            model_options = [
                ("Llama 3.1-8B (Default)", "llama-3.1-8b"),
                ("DeepSeek R1-Qwen-7B (~7GB)", "deepseek-r1-qwen-7b"),
                ("DeepSeek R1-Qwen-14B (~14GB)", "deepseek-r1-qwen-14b"),
                ("DeepSeek R1-Qwen-32B (⚠️ ~24GB RAM!)", "deepseek-r1-qwen-32b"),
                ("DeepSeek R1-Llama-8B (~8GB)", "deepseek-r1-llama-8b"),
                ("DeepSeek R1-Llama-70B (⚠️ ~40GB RAM!)", "deepseek-r1-llama-70b"),
                ("Qwen3-8B (Newer knowledge)", "qwen3-8b"),
                ("Qwen 2.5-14B (Larger)", "qwen-2.5-14b"),
                ("Mistral Nemo-12B", "mistral-nemo-12b"),
            ]
            model = prompt_choice("Model", model_options)
            if model is None:
                return
            
            # Output format
            print("\n📄 Select output format:\n")
            format_options = [
                ("Markdown (.md)", "md"),
                ("PDF (.pdf)", "pdf"),
                ("Word Document (.docx)", "docx"),
                ("HTML (.html)", "html"),
                ("Plain Text (.txt)", "txt"),
            ]
            output_format = prompt_choice("Format", format_options)
            if output_format is None:
                return
            
            # Research iterations (online only)
            research_iter = 3
            if online:
                print("\n🔄 Research iterations (sources to read):\n")
                iter_options = [
                    ("3 sources (Default)", "3"),
                    ("5 sources", "5"),
                    ("10 sources", "10"),
                    ("Custom", "custom"),
                ]
                iter_choice = prompt_choice("Iterations", iter_options)
                if iter_choice is None:
                    return
                if iter_choice == "custom":
                    custom_iter = prompt_text("Number of sources")
                    if custom_iter and custom_iter.isdigit():
                        research_iter = int(custom_iter)
                else:
                    research_iter = int(iter_choice)
            
            # Article length
            print("\n📏 Article Length:\n")
            length_options = [
                ("Quick (~500 words, fast) (Default)", "quick"),
                ("Standard (~1500 words)", "standard"),
                ("Detailed (~3000 words, comprehensive)", "detailed"),
            ]
            length = prompt_choice("Length", length_options)
            if length is None:
                return
            
            # Output file path (optional)
            print("\n📁 Output file path (leave empty for auto-name):\n")
            output_path = prompt_text("File path", required=False)
            
            # Build command
            flag = "-gr" if online else "-ga"
            cmd = f'{flag} -p "{topic}" -atm {model} --output-format {output_format} -al {length}'
            if online:
                cmd += f" -ri {research_iter}"
            if output_path:
                cmd += f' -o "{output_path}"'
            
            run_self_command(cmd)
            
            # Interactive result wait
            prompt_menu(None, [], allow_back=True)
    
    def code_menu(preset_model=None):
        """Generate code submenu."""
        while True:
            clear_screen()
            show_header("Generate Code")
            
            # Code description/prompt
            print("✏️  Describe what code you want to generate:")
            print("   (be more specific for better results)\n")
            print("💡 Tip: Include folder name for multi-file projects, e.g.:")
            print('   "Create React example in folder react-example"\n')
            print("(Leave empty to go back)\n")
            description = prompt_text("Description", required=False)
            if not description:
                return
            
            # Model selection
            print("\n📦 Select Code Model:\n")
            model_options = [
                ("Llama 3.1-8B (Default)", "llama-3.1-8b"),
                ("DeepSeek R1-Qwen-7B (~7GB)", "deepseek-r1-qwen-7b"),
                ("DeepSeek R1-Qwen-14B (~14GB)", "deepseek-r1-qwen-14b"),
                ("DeepSeek R1-Qwen-32B (⚠️ ~24GB RAM!)", "deepseek-r1-qwen-32b"),
                ("DeepSeek R1-Llama-8B (~8GB)", "deepseek-r1-llama-8b"),
                ("DeepSeek R1-Llama-70B (⚠️ ~40GB RAM!)", "deepseek-r1-llama-70b"),
                ("Qwen3-8B (Newer knowledge)", "qwen3-8b"),
                ("Qwen 2.5-14B (Larger)", "qwen-2.5-14b"),
            ]
            model = prompt_choice("Model", model_options)
            if model is None:
                return
            
            # Output path (optional)
            print("\n📁 Output path (optional):")
            print("   (Leave empty: uses paths/filenames from your description)")
            print("   (Existing folder: saves all generated files inside it)")
            print("   (Filename: override output name if single file)\n")
            output_path = prompt_text("Output path", required=False)
            
            # Build command
            cmd = f'-gc -p "{description}" -cdm {model}'
            if output_path:
                cmd += f' -o "{output_path}"'
            
            run_self_command(cmd)
            input("\nPress Enter to continue...")
    
    def chat_menu(preset_model=None):
        """Interactive chat submenu."""
        clear_screen()
        show_header("Chat")
        
        # Model selection (skip if preset)
        if preset_model:
            model = preset_model
            print(f"📦 Model: {model}\n")
        else:
            print("📦 Select Chat Model:\n")
            model_options = [
                ("Llama 3.1-8B (Default)", "llama-3.1-8b"),
                ("DeepSeek R1-Qwen-7B (~7GB)", "deepseek-r1-qwen-7b"),
                ("DeepSeek R1-Qwen-14B (~14GB)", "deepseek-r1-qwen-14b"),
                ("DeepSeek R1-Qwen-32B (⚠️ ~24GB RAM!)", "deepseek-r1-qwen-32b"),
                ("DeepSeek R1-Llama-8B (~8GB)", "deepseek-r1-llama-8b"),
                ("DeepSeek R1-Llama-70B (⚠️ ~40GB RAM!)", "deepseek-r1-llama-70b"),
                ("Qwen3-8B (Newer knowledge)", "qwen3-8b"),
                ("Qwen 2.5-14B (Larger)", "qwen-2.5-14b"),
                ("Mistral Nemo-12B", "mistral-nemo-12b"),
            ]
            model = prompt_choice("Model", model_options)
            if model is None:
                return
        
        # Build command and run
        cmd = f"-c --chat-model {model}"
        
        run_self_command(cmd)
        # No "Press Enter" needed - chat exits naturally
    
    def test_menu(preset_submenu=None):
        """Test selection submenu - choose between Unit Tests and Integration Tests."""
        # If preset submenu provided, go directly to that menu
        if preset_submenu == 'unit':
            unit_test_menu()
            return
        elif preset_submenu == 'integration':
            integration_test_menu()
            return
        elif preset_submenu == 'codec':
            # Run codec test directly then return
            script_dir = os.path.dirname(os.path.abspath(__file__))
            codec_test = os.path.join(script_dir, "tests", "test_codec_limits.py")
            if os.path.exists(codec_test):
                clear_screen()
                show_header("Codec Limits Test")
                print("Running codec limits stress test...")
                print("This tests H.264, HEVC, and AV1 encoder limits up to 20K.\n")
                import subprocess
                subprocess.run([sys.executable, codec_test])
                
                # Forcefully reset terminal to sane state (raw mode from child process leaks)
                try:
                    os.system('stty sane')  # Reset terminal on Unix
                except:
                    pass
                
                print("\n" + "="*60)
                input("Press Enter to return to menu...")
            else:
                clear_screen()
                show_header("Codec Limits Test")
                print("❌ tests/test_codec_limits.py not found.")
                input("\nPress Enter to continue...")
            # Fall through to test menu loop (don't return to main menu)
        
        while True:
            clear_screen()
            show_header("Run Tests")
            
            options = [
                ("🧪  Unit Tests (Python unittest)", "UNIT"),
                ("🚀  Integration Tests (tests/integration-tests.json)", "INTEGRATION"),
                ("📊  Codec Limits Test (tests/test_codec_limits.py)", "CODEC"),
            ]
            
            choice = prompt_menu("Select test type:", options)
            
            if choice is None: return
            
            if choice == "UNIT":
                unit_test_menu()
            elif choice == "INTEGRATION":
                integration_test_menu()
            elif choice == "CODEC":
                # Run codec limits test directly
                script_dir = os.path.dirname(os.path.abspath(__file__))
                codec_test = os.path.join(script_dir, "tests", "test_codec_limits.py")
                if os.path.exists(codec_test):
                    clear_screen()
                    show_header("Codec Limits Test")
                    print("Running codec limits stress test...")
                    print("This tests H.264, HEVC, and AV1 encoder limits up to 20K.\n")
                    import subprocess
                    subprocess.run([sys.executable, codec_test])
                    
                    # Forcefully reset terminal to sane state (raw mode from child process leaks)
                    try:
                        os.system('stty sane')
                    except:
                        pass
                    
                    print("\n" + "="*60)
                    input("Press Enter to return to menu...")
                else:
                    clear_screen()
                    show_header("Codec Limits Test")
                    print("❌ tests/test_codec_limits.py not found.")
                    print(f"   Expected location: {codec_test}")
                    input("\nPress Enter to continue...")
    
    def unit_test_menu():
        """Unit Tests submenu - dynamically loads test classes from tests/ai-media_test.py."""
        script_dir = os.path.dirname(os.path.abspath(__file__))
        test_module = os.path.join(script_dir, "tests", "ai-media_test.py")
        
        # Check if test file exists
        if not os.path.exists(test_module):
            clear_screen()
            show_header("Unit Tests")
            print("❌ tests/ai-media_test.py not found.")
            print(f"   Expected location: {test_module}")
            input("\nPress Enter to continue...")
            return
        
        # Extract test classes and count tests per class from tests/ai-media_test.py
        test_classes = {}  # {class_name: test_count}
        test_method_count = 0
        try:
            with open(test_module, 'r', encoding='utf-8') as f:
                lines = f.readlines()
            
            import re
            current_class = None
            class_pattern = r'^class\s+(Test\w+)\s*\(\s*(?:unittest\.)?TestCase\s*\)\s*:'
            method_pattern = r'^\s+def\s+(test_\w+)\s*\('
            
            for line in lines:
                # Check for class definition
                class_match = re.match(class_pattern, line)
                if class_match:
                    current_class = class_match.group(1)
                    test_classes[current_class] = 0
                # Check for test method
                elif current_class and re.match(method_pattern, line):
                    test_classes[current_class] += 1
                    test_method_count += 1
        except Exception as e:
            clear_screen()
            show_header("Unit Tests")
            print(f"❌ Error parsing test file: {e}")
            input("\nPress Enter to continue...")
            return
        
        if not test_classes:
            clear_screen()
            show_header("Unit Tests")
            print("❌ No test classes found in tests/ai-media_test.py")
            input("\nPress Enter to continue...")
            return
        
        # Build options
        options = []
        class_count = len(test_classes)
        options.append((f"🚀  Run All ({class_count} classes, {test_method_count} tests) [Summary]", "ALL_QUIET"))
        options.append((f"📜  Run All ({class_count} classes, {test_method_count} tests) [Verbose]", "ALL_VERBOSE"))
        
        for cls in sorted(test_classes.keys()):
            # Get method count for this class
            try:
                # We can inspect the class directly since we imported the module
                cls_obj = getattr(module, cls)
                method_count = len([m for m in dir(cls_obj) if m.startswith('test_')])
                options.append((f"{cls} ({method_count} tests)", cls))
            except:
                 options.append((f"{cls}", cls))

        while True:
            clear_screen()
            show_header("Unit Tests")
            prompt_text = (
                 f"Select a test class to run:\n\n"
                 f"ℹ️  Individual tests are always run in VERBOSE mode\n\n"
                 f"ℹ️  {class_count} test classes ({test_method_count} tests) found in tests/ai-media_test.py\n"
            )
            choice = prompt_choice(prompt_text, options, allow_back=True)

            if choice is None: return

            if choice == "ALL_QUIET":
                # Run all unit tests (Quiet)
                print("\n🧪 Running all unit tests (Summary)...\n")
                print("=" * 60)
                os.system(f'"{sys.executable}" -m unittest tests.ai-media_test')
            elif choice == "ALL_VERBOSE":
                # Run all unit tests (Verbose)
                print("\n🧪 Running all unit tests (Verbose)...\n")
                print("=" * 60)
                os.system(f'"{sys.executable}" -m unittest tests.ai-media_test -v')
            else:
                # Run specific test class (Always Verbose)
                print(f"\n🧪 Running {choice}...\n")
                print("=" * 60)
                os.system(f'"{sys.executable}" -m unittest tests.ai-media_test.{choice} -v')

            prompt_menu("Press Enter to continue...", [], allow_back=True)

    def integration_test_menu():
        """Integration Tests submenu - tests from tests/integration-tests.json."""
        # Load tests
        script_dir = os.path.dirname(os.path.abspath(__file__))
        test_file = os.path.join(script_dir, "tests", "integration-tests.json")
        try:
            with open(test_file, "r") as f:
                data = json.load(f)
            tests = data.get("tests", [])
        except Exception as e:
            clear_screen()
            show_header("App Run Tests")
            print(f"❌ Error loading tests: {e}")
            prompt_menu("Press Enter...", [], allow_back=True)
            return

        if not tests:
            clear_screen()
            show_header("App Run Tests")
            print("❌ No tests found.")
            input("Press Enter...")
            return

        # Build options
        options = []
        max_desc_len = 40  # Truncate descriptions to this length
        for t in tests:
            name = t.get("name", "Unnamed Test")
            desc = t.get("description", "")
            if desc:
                if len(desc) > max_desc_len:
                    desc = desc[:max_desc_len-3] + "..."
                display = f"{name} ({desc})"
            else:
                display = name
            options.append((display, name))
        
        # Prepend 'Run All' options
        count = len(tests)
        options.insert(0, (f"📜  Run All {count} Tests (Verbose)", "ALL_VERBOSE"))
        options.insert(0, (f"🚀  Run All {count} Tests (Summary)", "ALL_QUIET"))

        while True:
            clear_screen()
            show_header("App Run Tests")
            choice = prompt_menu("Select a test to run:\n\nℹ️  Individual tests are always run in VERBOSE mode\n", options)
            
            if choice is None: return
            
            if choice == "ALL_QUIET":
                # Run all tests (quiet/summary mode)
                run_self_command("--test")
            elif choice == "ALL_VERBOSE":
                # Run all tests (verbose mode)
                run_self_command("--test-verbose")
            else:
                # Run specific test
                # Always use verbose for single test as requested
                run_self_command(f"--test-verbose \"{choice}\"")
                
            input("\nPress Enter to continue...")


    
    # Main loop
    first_run = True
    while True:
        # Use jump point on first run if provided
        if first_run and initial_action:
            action = initial_action
            first_run = False
        else:
            action = main_menu()
        
        if action is None:
            print("\n👋 Goodbye!")
            break
        elif action == "image":
            image_menu(initial_model if first_run or initial_action == 'image' else None)
            initial_model = None  # Clear after first use
        elif action == "video":
            video_menu(initial_model if first_run or initial_action == 'video' else None)
            initial_model = None
        elif action == "audio":
            audio_menu(initial_model if first_run or initial_action == 'audio' else None)
            initial_model = None
        elif action == "transform":
            transform_menu(initial_model if first_run or initial_action == 'transform' else None)
            initial_model = None
        elif action == "upscale":
            upscale_menu()
        elif action == "convert":
            convert_menu()
        elif action == "doc_convert":
            document_convert_menu()
        elif action == "caption":
            caption_menu(initial_model if first_run or initial_action == 'caption' else None)
            initial_model = None
        elif action == "article":
            article_menu(initial_model if first_run or initial_action == 'article' else None)
            initial_model = None
        elif action == "code":
            code_menu(initial_model if first_run or initial_action == 'code' else None)
            initial_model = None
        elif action == "chat":
            chat_menu(initial_model if first_run or initial_action == 'chat' else None)
            initial_model = None
        elif action == "test":
            test_menu(initial_model if first_run or initial_action == 'test' else None)
            initial_model = None
        elif action == "sysinfo":
            system_info_menu()

# --- Test Runner ---


def run_tests(verbose=False, test_filter=None, exit_on_finish=True):
    """Run test suite from tests/integration-tests.json."""
    import shlex
    import subprocess
    
    # Use global test state for CTRL+C handling
    global _test_state
    
    script_dir = os.path.dirname(os.path.abspath(__file__))
    test_file = os.path.join(script_dir, "tests", "integration-tests.json")
    
    if not os.path.exists(test_file):
        print(f"{emoji('❌ ', 'Error: ')}Test file not found: {test_file}")
        if exit_on_finish: sys.exit(1)
        return False
    
    with open(test_file, "r") as f:
        data = json.load(f)
    
    tests = data.get("tests", [])
    if not tests:
        print(f"{emoji('❌ ', 'Error: ')}No tests found in tests/integration-tests.json")
        if exit_on_finish: sys.exit(1)
        return False
    
    # Filter tests if requested
    if test_filter:
        if isinstance(test_filter, list):
            print(f"{emoji('🔎 ', '')}Test Filter Mode: Searching for {len(test_filter)} specific tests...")
            tests = [t for t in tests if t.get("name") in test_filter]
        else:
             print(f"{emoji('🔎 ', '')}Single Test Mode: Searching for '{test_filter}'...")
             tests = [t for t in tests if t.get("name") == test_filter]
        
        if not tests:
            print(f"{emoji('❌ ', 'Error: ')}No tests found matching filter: {test_filter}")
            print("   Available tests are listed in the ID list, or:")
            print(f"\n{emoji('👉 ', '')}Redirecting to Interactive Test Menu in ", end="", flush=True)
            for i in range(3, 0, -1):
                print(f"{i}...", end="", flush=True)
                time.sleep(1)
            print()
            # Call interactive mode jumping to 'test' menu
            # We need to access run_interactive. It is defined in the same global scope.
            run_interactive(jump_point="test")
            return False
        
        print(f"{emoji('✅ ', '')}Found {len(tests)} test(s) matching filter.\n")

    # Warning prompt
    print(f"\n{'='*60}")
    print(f"{emoji('⚠️  ', '   ')}WARNING: Test Suite")
    print(f"{'='*60}")
    print(f"   • Integration tests can take a long time")
    print(f"   • Models will be downloaded if not present (2-30GB each)")
    print(f"   • High system resource consumption")
    print(f"   • Press CTRL+C at any time to interrupt")
    print(f"{'='*60}")
    
    if os.environ.get("AI_MEDIA_FORCE") != "1":
        try:
            choice = input(f"\n   Continue? [Y/n]: ").lower().strip()
            if choice in ['n', 'no']:
                print("❌ Test cancelled.")
                if exit_on_finish: sys.exit(0)
                return False
        except KeyboardInterrupt:
            print("\n❌ Test cancelled.")
            if exit_on_finish: sys.exit(0)
            return False
    else:
        print(f"\n   (Skipping confirmation due to --force)\n")
    
    print(f"\n{'='*60}")
    print(f"🧪 Running {len(tests)} test(s)")
    print(f"{'='*60}\n")
    
    passed = 0
    failed = 0
    skipped = 0
    results = []
    ran_count = 0
    
    # Resource aggregation variables
    total_ram = 0.0
    total_vram = 0.0
    total_cpu = 0.0
    total_gpu = 0.0
    resource_count = 0
    
    # Set global test state for CTRL+C handler
    _test_state['active'] = True
    _test_state['total'] = len(tests)
    _test_state['passed'] = 0
    _test_state['failed'] = 0

    suite_start_time = time.time()
    suite_start_time = time.time()
    # timestamp with milliseconds
    start_dt = datetime.fromtimestamp(suite_start_time)
    suite_timestamp = start_dt.strftime("%Y%m%d-%H%M%S-%f")[:-3]
    
    for i, test in enumerate(tests):
        test_name = test.get("name", f"Test {i+1}")
        command = test.get("command", "")
        expected_inputs = test.get("expectedInputItems", [])
        expected_outputs = test.get("expectedOutputItems", [])
        
        # Check for skip flag
        if test.get("skip") is True:
            print(f"\n{emoji('⏭️  ', '(-)')}Skipping test: {test_name} (skip: true)")
            skipped += 1
            continue
        
        # Formatting
        description = test.get("description", "")
        header = f"{emoji('📋 ', '')}Test {i+1}/{len(tests)}: {test_name}"
        desc = f"   {description}" if description else ""
        start_t_str = datetime.now().strftime("%H:%M:%S")
        time_line = f"   Start at: {start_t_str}"
        
        # Calculate dynamic width (min 50)
        lines = [header, desc, time_line]
        max_len = max(50, *[len(l) for l in lines if l])
        sep = "-" * max_len
        
        print(f"\n{sep}")
        print(header)
        if description:
            print(desc)
        print("")
        print(time_line)
        print(f"{sep}")
        
        test_passed = True
        failure_reason = None
        
        # 1. Check expected input items exist
        for input_item in expected_inputs:
            input_path = os.path.join(script_dir, input_item)
            if not os.path.exists(input_path):
                print(f"{emoji('❌ ', '[X] ')}Missing input: {input_item}")
                test_passed = False
                failure_reason = f"Missing input: {input_item}"
                break
        
        if not test_passed:
            print(f"{emoji('⏭️  ', '')}Skipping due to missing inputs")
            failed += 1
            results.append((test_name, False, failure_reason))
            continue
        
        # 2. Delete expected outputs before run (clean slate)
        # 2. Delete expected outputs before run (clean slate)
        for output_item in expected_outputs:
            output_path = os.path.join(script_dir, output_item)
            if os.path.exists(output_path):
                os.remove(output_path)
                print(f"{emoji('🗑️  ', '(-) ')}Deleted: {output_item}")
        
        # Prepare JSON Report Path
        import tempfile
        # Use suite-wide timestamp for better file grouping
        json_report_path = os.path.join(script_dir, f"{suite_timestamp}-{i+1:03d}-temp-performance.json")
        
        # Add --report-json arg if command supports it (assuming all ai-media commands do)
        full_command = [sys.executable, "-u", os.path.join(script_dir, 'ai-media.py')] + shlex.split(command) + ["--report-json", json_report_path]
        
        # For logging, show command without the hidden report arg
        cmd_display = f"python ai-media.py {command}"
        print(f"{emoji('🚀 ', '(>) ')}Running: {cmd_display}")
            
            # ... rest of subprocess creation ...
        
        is_interactive = test.get("interactive", False)
        interactive_wait = test.get("interactiveWait", 2.0)
        expected_stdout_items = test.get("expectedStdoutItems", [])
        
        start_time = time.time()
        current_process = None
        try:
            # Set UTF-8 for subprocess to handle emoji on Windows
            env = os.environ.copy()
            env["PYTHONIOENCODING"] = "utf-8"
            
            # Use Popen for better control over the subprocess
            # On Windows, create subprocess in new process group so it doesn't receive console Ctrl+C
            # This allows parent to catch KeyboardInterrupt and kill subprocess properly
            creation_flags = 0
            if os.name == 'nt':
                creation_flags = subprocess.CREATE_NEW_PROCESS_GROUP
            
            current_process = subprocess.Popen(
                full_command,
                cwd=script_dir,
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT, # Merge stderr into stdout for simple streaming
                stdin=subprocess.DEVNULL if is_interactive else None,  # Isolate interactive tests from parent tty
                text=True,
                encoding='utf-8',
                errors='replace',
                env=env,
                bufsize=1, # Line buffered
                universal_newlines=True,
                creationflags=creation_flags
            )
            
            # Track subprocess for signal handler cleanup
            _test_state['current_process'] = current_process
            
            stdout_lines = []
            
            try:
                if is_interactive:
                    if verbose: print(f"{emoji('⏳ ', 'Wait: ')}Waiting {interactive_wait}s for interactive output...")
                    time.sleep(interactive_wait)
                    
                    # Terminate interactive process (Windows-compatible)
                    if os.name == 'nt':
                        current_process.terminate()  # Windows: terminate gracefully
                    else:
                        current_process.send_signal(signal.SIGINT)  # Unix: send CTRL+C
                        
                    # For interactive tests, we don't need real-time streaming, just capture final output
                    # and ensure it doesn't hang.
                    timeout_limit = test.get("timeout", 600)
                    stdout, stderr = current_process.communicate(timeout=5) # Short timeout after term
                    stdout_lines = [stdout] if stdout else []
                    
                    # Note: stderr is merged to stdout in Popen, so stderr var will be None/Empty
                    
                else:
                    # Non-interactive (Batch) mode: Stream output real-time
                    timeout_limit = test.get("timeout", 600)
                    start_read_time = time.time()
                    
                    while True:
                        # Check for timeout
                        if time.time() - start_read_time > timeout_limit:
                            raise subprocess.TimeoutExpired(full_command, timeout_limit)
                        
                        # Read line
                        line = current_process.stdout.readline()
                        if not line and current_process.poll() is not None:
                            break
                        
                        if line:
                            # Stream to user if verbose mode is on (but not for interactive tests)
                            if verbose:
                                print(line, end='', flush=True) 
                            stdout_lines.append(line)
                            
                    current_process.wait() # Ensure clean exit

                elapsed = time.time() - start_time
                stdout = "".join(stdout_lines) 
                # stderr is merged into stdout via Popen(stderr=subprocess.STDOUT)
                
                # Show verbose output marking
                if verbose and not is_interactive:
                    print(f"\n--- END ---\n")
                
                if current_process.returncode != 0 and not is_interactive:
                    # Check if it was a Ctrl+C interrupt
                    # Unix: 130 (128 + SIGINT=2), Windows: -2 or 3221225786 (STATUS_CONTROL_C_EXIT)
                    is_ctrl_c = current_process.returncode in [130, -2, 3221225786, -1073741510]
                    
                    if is_ctrl_c:
                        print(f"\n\n⚠️  Interrupted! Cleaning up...")
                        # Re-raise to trigger the outer handler
                        raise KeyboardInterrupt()
                    else:
                        # Interactive tests expect SIGINT exit code (usually 130 or 1 or 0 handling)
                        # If caught cleanly it might be 0.
                        # We only fail non-interactive tests on non-zero exit code here unless specific check later.
                        print(f"{emoji('❌ ', 'Error: ')}Command failed with exit code {current_process.returncode}")
                        # Since stderr is merged, we can't print it separately, but it's already on screen.
                        test_passed = False
                        failure_reason = f"Exit code {current_process.returncode}"
                
                # Check STDOUT items
                if test_passed and expected_stdout_items:
                    for item in expected_stdout_items:
                        if item not in stdout:
                            print(f"{emoji('❌ ', 'Error: ')}Missing stdout item: '{item}'")
                            test_passed = False
                            failure_reason = f"Missing stdout: '{item}'"
                            break
                        else:
                            if verbose: print(f"✓ Found stdout item: '{item}'")
                            
            except subprocess.TimeoutExpired:
                current_process.kill()
                current_process.wait()
                elapsed = time.time() - start_time
                print(f"{emoji('❌ ', 'Error: ')}Command timed out after {timeout_limit} seconds")
                test_passed = False
                failure_reason = "Timeout"
                
        except KeyboardInterrupt:
            # Terminate subprocess and re-raise to be caught by outer handler
            print(f"\n\n⚠️  Interrupted! Cleaning up subprocess...")
            if current_process:
                pid = current_process.pid
                
                # On Windows, use taskkill to kill the entire process tree
                if os.name == 'nt':
                    try:
                        # /T = kill child processes, /F = force
                        subprocess.run(['taskkill', '/T', '/F', '/PID', str(pid)], 
                                       capture_output=True, timeout=5)
                        print(f"   ✅ Process tree (PID {pid}) terminated.")
                    except Exception as e:
                        print(f"   ⚠️  taskkill failed: {e}, trying fallback...")
                        current_process.kill()
                else:
                    # Unix: kill process group
                    try:
                        os.killpg(os.getpgid(pid), signal.SIGKILL)
                    except:
                        current_process.kill()
                
                try:
                    current_process.wait(timeout=2)
                except:
                    pass
            raise
            
        except Exception as e:
            elapsed = time.time() - start_time
            print(f"{emoji('❌ ', 'Error: ')}Command exception: {e}")
            test_passed = False
            failure_reason = str(e)
        
        # 4. Check expected output items exist
        if test_passed:
            # Parse Resource Usage from stdout -> DEPRECATED/REMOVED in favor of JSON IPC
            # (Old regex parsing block removed)
            pass

            for output_item in expected_outputs:
                output_path = os.path.join(script_dir, output_item)
                if not os.path.exists(output_path):
                    print(f"{emoji('❌ ', 'Error: ')}Missing output: {output_item}")
                    test_passed = False
                    failure_reason = f"Missing output: {output_item}"
                    break
                else:
                    print(f"✓ Output exists: {output_item}")
        
        if test_passed:
            # 5. Read JSON Report if available
            if json_report_path and os.path.exists(json_report_path):
                try:
                    with open(json_report_path, 'r') as f:
                        stats = json.load(f)
                        
                    # Extract Data
                    r_time = stats.get("time", 0)
                    r_ram = stats.get("ram", 0)
                    r_vram = stats.get("vram", 0)
                    r_cpu = stats.get("cpu", 0)
                    r_gpu = stats.get("gpu", 0)
                    
                    total_ram += r_ram
                    total_vram += r_vram
                    total_cpu += r_cpu
                    total_gpu += r_gpu
                    resource_count += 1
                    
                    # Store exact generation time in results if available
                    results.append((test_name, True, f"{r_time:.1f}s"))
                        
                except Exception as e:
                    print(f"{emoji('⚠️ ', 'Warning: ')}Failed to read stats JSON: {e}")
                    results.append((test_name, True, f"{elapsed:.1f}s"))
            else:
                # Fallback to elapsed time if no JSON report
                results.append((test_name, True, f"{elapsed:.1f}s"))
        
        # Always cleanup report file if it exists (even if test failed)
        if json_report_path and os.path.exists(json_report_path):
            try:
                os.remove(json_report_path)
            except:
                pass

        if test_passed:
            print(f"{emoji('✅ ', '')}PASSED ({elapsed:.1f}s)")
            passed += 1
            _test_state['passed'] = passed
        else:
            print(f"{emoji('❌ ', '')}FAILED ({elapsed:.1f}s)")
            failed += 1
            _test_state['failed'] = failed
            results.append((test_name, False, failure_reason))
    
        ran_count += 1
    
    # Mark test as no longer active
    _test_state['active'] = False
    
    total_duration = time.time() - suite_start_time

    # Print summary
    print(f"\n{'='*60}")
    print(f"{emoji('📊 ', '')}TEST SUMMARY")
    print(f"{'='*60}")
    print(f"   Total:  {len(tests)}")
    print(f"   Passed: {passed} {emoji('✅', '')}")
    

    
    if failed > 0:
        print(f"   Failed: {failed} {emoji('❌', '')}")
    else:
        print(f"   Failed: {failed}")
        
    if skipped > 0:
        print(f"   Skipped: {skipped} {emoji('⏭️', '')}")
        
    print(f"   Duration: {format_time(total_duration)}")
    
    if resource_count > 0:
        avg_ram = total_ram / resource_count
        avg_vram = total_vram / resource_count
        avg_cpu = total_cpu / resource_count
        avg_gpu = total_gpu / resource_count
        
        if len(tests) >= 2:
             print(f"\n   {emoji('⚖️  ', '')}Averages:\n")
        
        print(f"   RAM: {avg_ram:.1f} GB")
        print(f"   VRAM: {avg_vram:.1f} GB")
        print(f"   CPU: {avg_cpu:.1f} %")
        print(f"   GPU: {avg_gpu:.1f} %")
        
    print(f"{'='*60}")
    
    if failed > 0:
        print(f"\n❌ Failed Tests:")
        for name, success, reason in results:
            if not success:
                print(f"   - {name}: {reason}")
    
    print(f"\n✅ Test Run Complete")
    
    sys.exit(0 if failed == 0 else 1)

# -------------------------------------------------------------------------
# NEW: Article Generation & Deep Research
# -------------------------------------------------------------------------

class ArticleGenerator:
    def __init__(self, model_name="llama-3.1-8b", device=None):
        self.model_name = TEXT_MODELS.get(model_name.lower(), model_name)
        if model_name.lower() == "default": self.model_name = TEXT_MODELS["default"]
        
        self.device = device or get_optimal_device_and_dtype(quiet=True)[0]
        self.pipeline = None
        self.ddgs = DDGS()
        
    def _load_model(self):
        if self.pipeline: return
        
        print(f"📚 Loading Text Model: {self.model_name}...")
        try:
            dtype = torch.float16 if self.device.type in ["cuda", "mps"] else torch.float32
            
            
            # Workaround: Qwen3/Llama3 have numerical instability on MPS float16, use fp32
            # Note: This increases RAM usage significantly (e.g. 8B model -> ~32GB RAM)
            if self.device.type == "mps" and any(m in self.model_name.lower() for m in ["qwen3", "llama"]):
                print(f"   ⚠️  {self.model_name} detected on MPS - using fp32 for stability...")
                dtype = torch.float32
                import os
                os.environ["PYTORCH_ENABLE_MPS_FALLBACK"] = "1"
            
            # Load tokenizer and model separately to handle chat templates
            tokenizer = AutoTokenizer.from_pretrained(self.model_name)
            
            # Memory optimization: 4-bit/8-bit if CUDA, otherwise standard
            quantization_config = None
            if self.device.type == "cuda":
                try:
                    from transformers import BitsAndBytesConfig
                    quantization_config = BitsAndBytesConfig(load_in_4bit=True)
                except ImportError:
                    pass
            
            model_kwargs = {"torch_dtype": dtype}
            if quantization_config:
                model_kwargs["quantization_config"] = quantization_config
                # device_map="auto" is usually required for quantization
                model_kwargs["device_map"] = "auto"
            else:
                model_kwargs["device_map"] = self.device
                
            self.pipeline = pipeline(
                "text-generation",
                model=self.model_name,
                tokenizer=tokenizer,
                **model_kwargs
            )
            print("✅ Model loaded.")
        except RuntimeError as e:
            error_msg = str(e)
            if "Invalid buffer size" in error_msg or "out of memory" in error_msg.lower():
                # Extract the size if present (e.g., "59.58 GiB")
                import re
                size_match = re.search(r'(\d+\.?\d*)\s*(GiB|GB|MiB|MB)', error_msg)
                size_info = f" (model requires ~{size_match.group(0)})" if size_match else ""
                
                print(f"\n❌ Model too large for this system{size_info}")
                print(f"   The model '{self.model_name}' cannot fit in available memory.")
                print(f"   💡 Try a smaller model like:")
                print(f"      - deepseek-r1-qwen-7b (~7GB)")
                print(f"      - deepseek-r1-llama-8b (~8GB)")
                print(f"      - llama-3.1-8b (~16GB)")
                return  # Return gracefully without raising
            else:
                print(f"❌ Failed to load model: {e}")
                raise
        except OSError as e:
            error_msg = str(e)
            if "not a valid model identifier" in error_msg or "Repository Not Found" in error_msg:
                print(f"\n❌ Model not found: '{self.model_name}'")
                print(f"   This model doesn't exist on HuggingFace.")
                print(f"   💡 Available models:")
                print(f"      - deepseek-r1-qwen-7b, deepseek-r1-qwen-14b, deepseek-r1-qwen-32b")
                print(f"      - deepseek-r1-llama-8b, deepseek-r1-llama-70b")
                print(f"      - llama-3.1-8b, qwen3-8b, qwen-2.5-14b, mistral-nemo-12b")
                return  # Return gracefully without raising
            else:
                print(f"❌ Failed to load model: {e}")
                raise
        except (ValueError, Exception) as e:
            error_msg = str(e)
            # Check if this is a wrapped memory error (transformers wraps RuntimeError in ValueError)
            if "Invalid buffer size" in error_msg or "out of memory" in error_msg.lower():
                import re
                size_match = re.search(r'(\d+\.?\d*)\s*(GiB|GB|MiB|MB)', error_msg)
                size_info = f" (model requires ~{size_match.group(0)})" if size_match else ""
                
                print(f"\n❌ Model too large for this system{size_info}")
                print(f"   The model '{self.model_name}' cannot fit in available memory.")
                print(f"   💡 Try a smaller model like:")
                print(f"      - deepseek-r1-qwen-7b (~7GB)")
                print(f"      - deepseek-r1-llama-8b (~8GB)")
                print(f"      - llama-3.1-8b (~16GB)")
                return  # Return gracefully without raising
            else:
                print(f"❌ Failed to load model: {e}")
                raise

    def deep_research(self, query, iterations=3):
        """Perform recursive web search and summarization."""
        print(f"\n🔎 Deep Researching: '{query}' ({iterations} iterations)...")
        results = []
        
        # 1. Initial Broad Search
        try:
            search_results = list(self.ddgs.text(query, max_results=iterations))
            pad_width = len(str(iterations))
            for i, res in enumerate(search_results, 1):
                num_str = str(i).zfill(pad_width)
                print(f"   Reading [{num_str}]: {res['title']}...")
                # Basic scraping (just using description/snippet for now to be safe/fast)
                content = res.get('body', '') or res.get('snippet', '')
                
                # Attempt deep scraping for better context
                try:
                    import requests
                    from bs4 import BeautifulSoup
                    
                    # 5 second timeout to keep it snappy
                    page = requests.get(res['href'], timeout=5, headers={"User-Agent": "Mozilla/5.0"})
                    if page.status_code == 200:
                        soup = BeautifulSoup(page.text, 'html.parser')
                        # Extract text from paragraphs (simple generic scraper)
                        paragraphs = [p.get_text().strip() for p in soup.find_all('p')]
                        full_text = ' '.join(p for p in paragraphs if p)
                        
                        # Only use if we got substantial content (>200 chars), otherwise fallback to snippet
                        if len(full_text) > 200:
                            # Limit to ~4000 chars per source to save context window
                            content = full_text[:4000] + "..."
                except Exception:
                    # Silently fail back to snippet
                    pass
                
                results.append(f"Source: {res['title']}\nURL: {res['href']}\nContent: {content}\n")
                
                # Rate limit handling (pause)
                time.sleep(1.0) 
        except Exception as e:
            print(f"⚠️ Search error: {e}")
        
        # 2. Image Search - Find real, publicly accessible images
        image_results = []
        try:
            image_query = f"{query} photos"
            print(f"   🖼️  Searching for images...")
            image_search = list(self.ddgs.images(image_query, max_results=5))
            
            for img in image_search:
                img_url = img.get('image', '')
                img_title = img.get('title', 'Image')
                if img_url and img_url.startswith('http'):
                    image_results.append(f"![{img_title}]({img_url})")
            
            if image_results:
                print(f"   ✅ Found {len(image_results)} images")
        except Exception as e:
            print(f"   ⚠️  Image search failed: {e}")
            
        research_context = "\n\n".join(results)
        
        # Append image URLs to research context
        if image_results:
            research_context += "\n\n## Available Images (use these exact URLs)\n"
            research_context += "\n".join(image_results)
        
        return research_context

    def chat_session(self):
        """Interactive Chat Loop."""
        console = Console()
        
        self._load_model()
        history = []
        pending_context = "" # Buffer for file context
        
        # Custom Lexer for command highlighting
        class ChatLexer(Lexer):
            def lex_document(self, document):
                def get_line_tokens(line_number):
                    line = document.lines[line_number]
                    # Colors the command part of the string
                    for cmd in ['/read', '/save', '/search', '/online-search']:
                        if line.startswith(cmd):
                            # Check for iteration modifier (e.g. /search|5)
                            base_len = len(cmd)
                            if len(line) > base_len and line[base_len] == '|':
                                # Find end of command (space or end of line)
                                end_pos = line.find(' ', base_len)
                                if end_pos == -1: end_pos = len(line)
                                return [
                                    ('class:command', line[:end_pos]),
                                    ('', line[end_pos:])
                                ]
                            return [
                                ('class:command', cmd),
                                ('', line[base_len:])
                            ]
                    return [('', line)]
                return get_line_tokens

        chat_style = Style.from_dict({
            'command': '#ff00ff bold', # Magenta/Pink to match header
        })

        # Initialize PromptSession for arrow key support / history
        session = PromptSession(
            history=InMemoryHistory(),
            lexer=ChatLexer(),
            style=chat_style
        )

        # Setup Autocomplete for Slash Commands (Fuzzy + Case Insensitive)
        path_completer = FuzzyCompleter(PathCompleter(expanduser=True))
        completer = NestedCompleter.from_nested_dict({
            '/read': path_completer,
            '/save': path_completer,
            '/search': None,
            '/online-search': None,
            'exit': None,
            'quit': None,
        })
        
        console.print(f"\n💬 [bold]Chat Session Started[/bold] (Model: [bold cyan]{self.model_name}[/bold cyan])")
        console.print("   Type '[bold]exit[/bold]' or '[bold]quit[/bold]' to end.")
        console.print("   Commands: [bold]/read <path>[/bold], [bold]/save[/bold][bold]|all[/bold] [bold]<path>[/bold], [bold]/search[/bold][bold]|N[/bold] [bold]<query>[/bold]")
        console.print("   [dim]💡 Tip: /save saves last code or full response. Use |all for full history.[/dim]")
        console.print("   [dim]💡 Tip: Use /search query or /search|5 query for deeper results.[/dim]\n")
        
        while True:
            try:
                # Use prompt_toolkit for input (enables arrow keys & autocomplete)
                user_input = session.prompt(HTML('<b fg="blue">You:</b> '), completer=completer, complete_while_typing=True)
                if user_input.strip().lower() in ['exit', 'quit']:
                    break
                
                # --- Slash Commands ---
                if user_input.startswith("/read "):
                    file_path = user_input[6:].strip()
                    if os.path.exists(file_path):
                        try:
                            with open(file_path, "r", encoding="utf-8") as f:
                                content = f.read()
                            
                            # Append to pending context instead of history immediately
                            pending_context += f"\n\n[File Context: {file_path}]\n{content}\n"
                            console.print(f"📄 [bold green]Added file context:[/bold green] {file_path} [dim](Stored in memory. You can now ask questions or provide prompts about its content!)[/dim]")
                        except Exception as e:
                            console.print(f"[bold red]❌ Error reading file:[/bold red] {e}")
                    else:
                        console.print(f"[bold red]❌ File not found:[/bold red] {file_path}")

                    continue # Skip generation for this turn

                is_search = False
                for s_cmd in ["/search", "/online-search"]:
                    if user_input.startswith(s_cmd + " ") or user_input.startswith(s_cmd + "|") or user_input == s_cmd:
                        is_search = True
                        break

                if is_search:
                    # Parse command and optional iterations: /search|5 query
                    parts = user_input.split(' ', 1)
                    cmd_part = parts[0]
                    query = parts[1].strip() if len(parts) > 1 else ""
                    
                    # Extract iterations if pipe exists
                    iterations = 3
                    if '|' in cmd_part:
                        try:
                            iter_str = cmd_part.split('|', 1)[1]
                            if iter_str.isdigit():
                                iterations = int(iter_str)
                        except:
                            pass
                    
                    if query:
                        try:
                            # Use existing deep_research method
                            search_results = self.deep_research(query, iterations=iterations)
                            
                            if search_results:
                                pending_context += f"\n\n[Online Search Context: '{query}']\n{search_results}\n"
                                console.print(f"🌍 [bold green]Added search results for:[/bold green] '{query}' [dim](Context enriched. Ask your question now to analyze these findings!)[/dim]")
                            else:
                                console.print(f"[bold yellow]⚠️ No results found for:[/bold yellow] '{query}'")
                        except Exception as e:
                             console.print(f"[bold red]❌ Search error:[/bold red] {e}")
                    else:
                        console.print("[bold red]❌ Please provide a search query.[/bold red]")
                    continue

                if user_input.startswith("/save"):
                    # Handle /save|all syntax or regular /save
                    parts = user_input.split(' ', 1)
                    cmd_part = parts[0]
                    file_path = parts[1].strip() if len(parts) > 1 else ""
                    
                    save_all = "|all" in cmd_part.lower()
                    content_to_save = None
                    label = "response"
                    ext_suggestion = ".md"

                    if save_all:
                        # Format full history as markdown
                        history_content = "# Chat Conversation History\n\n"
                        for msg in history:
                            role = "User" if msg["role"] == "user" else "Assistant"
                            history_content += f"## {role}\n{msg['content']}\n\n"
                        content_to_save = history_content
                        label = "full conversation history"
                        ext_suggestion = ".md"
                    else:
                        # 1. Try to find last code block
                        for msg in reversed(history):
                            if msg["role"] == "assistant":
                                content = msg["content"]
                                import re
                                # Attempt to find code blocks and language
                                matches = re.findall(r"```(.*?)\n(.*?)```", content, re.DOTALL)
                                if matches:
                                    lang, code = matches[-1]
                                    content_to_save = code
                                    label = "code block"
                                    # Suggest extension based on language
                                    lang_map = {"python": ".py", "bash": ".sh", "javascript": ".js", "html": ".html", "css": ".css", "markdown": ".md"}
                                    ext_suggestion = lang_map.get(lang.strip().lower(), ".txt")
                                    
                                    # Heuristic: look for filename in the text before the block (only if user didn't provide one)
                                    if not file_path:
                                        fn_match = re.search(r"(\w+[\.\w]+)", content[:content.find("```")].split("\n")[-1])
                                        if fn_match and "." in fn_match.group(1):
                                            suggested_fn = fn_match.group(1)
                                            # Only use if it has a valid extension
                                            if os.path.splitext(suggested_fn)[1] in lang_map.values():
                                                file_path = suggested_fn
                                    break
                                else:
                                    # 2. Fallback to full last response
                                    content_to_save = content
                                    label = "last response"
                                    ext_suggestion = ".md"
                                    break

                    if not file_path:
                        # Try to get a descriptive name from context
                        context_str = ""
                        if save_all and history:
                            # Use first user message that isn't a command
                            for msg in history:
                                if msg["role"] == "user" and not msg["content"].strip().startswith("/"):
                                    context_str = msg["content"]
                                    break
                        elif history:
                            # Use last user message that isn't a command
                            for msg in reversed(history):
                                if msg["role"] == "user" and not msg["content"].strip().startswith("/"):
                                    context_str = msg["content"]
                                    break
                        
                        # Slugify: lowercase, alphanumeric and underscores only
                        import re
                        # Clean up: remove Markdown, special chars, etc.
                        clean_text = re.sub(r'[^a-zA-Z0-9\s]', '', context_str).strip()
                        slug = re.sub(r'\s+', '_', clean_text[:25]).lower()
                        
                        ts = int(time.time())
                        if save_all:
                            prefix = f"chat_{slug}" if slug else "chat_all"
                        else:
                            type_prefix = "code" if label == "code block" else "resp"
                            prefix = f"{type_prefix}_{slug}" if slug else type_prefix
                            
                        file_path = f"{prefix}_{ts}{ext_suggestion}"
                    elif "." not in os.path.basename(file_path):
                        # Add suggested extension if missing
                        file_path += ext_suggestion


                    if content_to_save:
                        # Check overwrite before saving
                        always_overwrite = self.args.force if hasattr(self, 'args') else False
                        should_write, final_path, _, _ = check_overwrite(file_path, always_overwrite=always_overwrite)
                        if should_write:
                            try:
                                with open(final_path, "w", encoding="utf-8") as f:
                                    f.write(content_to_save)
                                console.print(f"💾 [bold green]Exported {label} to:[/bold green] {final_path}")
                            except Exception as e:
                                console.print(f"[bold red]❌ Error saving file:[/bold red] {e}")
                        else:
                            console.print(f"\n[bold yellow]⏭️  Save cancelled (skipped).[/bold yellow]")
                    else:
                         console.print("[bold red]❌ No conversation content found to save.[/bold red]")
                    continue

                # Construct Chat Prompt (using apply_chat_template if available)
                final_content = user_input
                if pending_context:
                    final_content = pending_context + "\n" + user_input
                    pending_context = "" # Clear buffer
                
                history.append({"role": "user", "content": final_content})
                
                # Check context length (simple trimming)
                if len(history) > 20: history = history[-10:] # Keep last 10 turns
                
                prompt = self.pipeline.tokenizer.apply_chat_template(
                    history, tokenize=False, add_generation_prompt=True
                )
                
                # Rich "Thinking..." Spinner
                mps_switch_needed = False
                with console.status("[yellow]Thinking...[/yellow]", spinner="dots"):
                    try:
                        # Generate
                        outputs = self.pipeline(
                            prompt, 
                            max_new_tokens=512, 
                            do_sample=True, 
                            temperature=0.7,
                            top_p=0.9,
                        )
                    except RuntimeError as e:
                        # Handle MPS probability tensor error (large context instability)
                        if ("probability tensor" in str(e) or "out of range integral" in str(e)) and self.device.type == "mps":
                            mps_switch_needed = True
                        else:
                            raise
                
                # Handle MPS model switch OUTSIDE the spinner context
                if mps_switch_needed:
                    console.print("\n[bold yellow]⚠️  MPS precision issue with large context[/bold yellow]")
                    console.print("[dim]This model struggles with large file contexts on Apple Silicon.[/dim]\n")
                    
                    # Offer interactive model switch
                    console.print("[bold]Switch to a more stable model?[/bold]")
                    console.print("  1. [cyan]llama-3.1-8b[/cyan] - Most stable on MPS")
                    console.print("  2. [cyan]deepseek-r1-llama-8b[/cyan] - Llama-based, good stability")
                    console.print("  3. [cyan]deepseek-r1-qwen-7b[/cyan] - Reasoning-focused")
                    console.print("  0. Continue with current model (may fail again)\n")
                    
                    try:
                        # Use prompt_toolkit's prompt (not input()) since it's managing the terminal
                        choice = session.prompt("Choice [1]: ").strip() or "1"
                        
                        model_map = {
                            "1": "meta-llama/Meta-Llama-3.1-8B-Instruct",
                            "2": "deepseek-ai/DeepSeek-R1-Distill-Llama-8B",
                            "3": "deepseek-ai/DeepSeek-R1-Distill-Qwen-7B",
                        }
                        
                        if choice in model_map:
                            new_model = model_map[choice]
                            
                            # Unload current model and load new one with status spinner
                            with console.status("[yellow]🔄 Switching model...[/yellow]", spinner="dots"):
                                self.pipeline = None
                                self.model_name = new_model
                                self._load_model()
                            
                            console.print("[bold green]✅ Model switched![/bold green]")
                            
                            # Re-build prompt with new tokenizer and retry
                            prompt = self.pipeline.tokenizer.apply_chat_template(
                                history, tokenize=False, add_generation_prompt=True
                            )
                            
                            with console.status("[yellow]Retrying with new model...[/yellow]", spinner="dots"):
                                outputs = self.pipeline(
                                    prompt, 
                                    max_new_tokens=512, 
                                    do_sample=True, 
                                    temperature=0.7,
                                    top_p=0.9,
                                )
                        else:
                            console.print("[dim]Continuing with current model...[/dim]\n")
                            continue
                    except KeyboardInterrupt:
                        console.print("\n[dim]Cancelled[/dim]")
                        continue
                
                console.print("[bold green]Bot:[/bold green] ", end="")
                
                generated_text = outputs[0]['generated_text']
                # Extract only new response (remove prompt)
                response = generated_text[len(prompt):].strip()
                
                # Detect DeepSeek R1 reasoning: everything before </think> is reasoning
                # The opening <think> is in the prompt template, closing </think> is in response
                if '</think>' in response:
                    parts = response.split('</think>', 1)
                    reasoning = parts[0].strip()
                    final_answer = parts[1].strip() if len(parts) > 1 else ""
                    
                    # Display reasoning in dim italic
                    console.print("")
                    console.print("[dim italic]💭 Reasoning:[/dim italic]")
                    console.print(f"[dim italic]{reasoning}[/dim italic]")
                    console.print("")  # Spacer
                    if final_answer:
                        console.print("[bold]Answer:[/bold]")
                        console.print(Markdown(final_answer))
                else:
                    # No reasoning tags, render normally
                    console.print(Markdown(response))
                console.print("")
                
                history.append({"role": "assistant", "content": response})
                
            except KeyboardInterrupt:
                console.print("\n")
                break
            except Exception as e:
                console.print(f"[bold red]❌ Error:[/bold red] {e}")

    def generate_article(self, topic, output_file, format="md", online=False, research_iter=3, length="quick"):
        """Generate full article with optional research.
        
        Args:
            length: 'quick' (512 tokens), 'standard' (2048), 'detailed' (4096 - default)
        """
        
        # Length presets
        from rich.console import Console
        console = Console()
        
        length_config = {
            "quick": {"tokens": 512, "desc": "concise"},
            "standard": {"tokens": 2048, "desc": "balanced"},
            "detailed": {"tokens": 4096, "desc": "comprehensive"},
        }
        config = length_config.get(length, length_config["detailed"])
        max_tokens = config["tokens"]
        style = config["desc"]
        
        research_data = ""
        if online:
            with console.status(f"[bold green]Thinking... (Deep Research Iterations {research_iter})[/bold green]", spinner="dots"): 
                research_data = self.deep_research(topic, iterations=research_iter)
        
        self._load_model()
        
        print(f"✍️  Writing {style} article on '{topic}'...")
        
        # Prompt Engineering (adjusted based on length)
        if research_data:
            system_prompt = (
                f"You are an expert investigative journalist. Write a {style}, well-structured "
                "article based on the following research context. Use Markdown formatting. "
                "Cite sources where appropriate."
            )
            user_prompt = f"Topic: {topic}\n\nResearch Context:\n{research_data}\n\nArticle:"
        else:
            system_prompt = (
                f"You are a creative writer and expert knowledge base. Write a {style}, "
                "well-structured article on the following topic. Use Markdown formatting."
            )
            user_prompt = f"Topic: {topic}\n\nArticle:"

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
        
        # Apply Template
        full_prompt = self.pipeline.tokenizer.apply_chat_template(
            messages, tokenize=False, add_generation_prompt=True
        )
        
        # Generate Content
        # Generate Content
        with console.status("[bold green]Thinking... (Writing Article)[/bold green]", spinner="dots"):
            outputs = self.pipeline(
                full_prompt, 
                max_new_tokens=max_tokens, 
                do_sample=True, 
                temperature=0.7
            )
        
        # Extract markdown
        final_md = outputs[0]['generated_text'][len(full_prompt):].strip()
        
        # Extract and save <think> blocks separately (for reasoning models like DeepSeek R1, Qwen3)
        import re
        think_matches = re.findall(r'<think>(.*?)</think>', final_md, re.DOTALL)
        if think_matches:
            # Remove all think blocks from main output
            final_md = re.sub(r'<think>.*?</think>\s*', '', final_md, flags=re.DOTALL).strip()
            # Save all think blocks to separate file
            base, ext = os.path.splitext(output_file)
            think_file = f"{base}-think.md"
            try:
                with open(think_file, "w", encoding="utf-8") as f:
                    f.write("# Reasoning Process\n\n")
                    for i, block in enumerate(think_matches, 1):
                        if len(think_matches) > 1:
                            f.write(f"## Block {i}\n\n")
                        f.write(block.strip() + "\n\n")
                print(f"💭 Reasoning saved to: {think_file}")
            except Exception as e:
                print(f"⚠️  Could not save reasoning: {e}")
        
        # Save Formats (pass online flag to track image failures)
        failed_images = self._save_formatted(final_md, output_file, format, online=online)
        
        # If offline and images failed, offer to retry with research
        if not online and failed_images > 0:
            print(f"\n⚠️  {failed_images} image(s) could not be fetched (hallucinated URLs).")
            print("💡 Tip: Offline models (-ga) cannot provide real image URLs.")
            print("   Options:")
            print("   • Use Deep Research (-gr) to find real images from the web")
            print("   • Remove 'images' from your prompt for text-only articles")
            
            # Offer retry
            retry = prompt_choice("What would you like to do?", [
                ("Retry with Deep Research (online)", "y"),
                ("Keep current output (no images)", "n")
            ])
            
            if retry == "y":
                print("\n🔄 Retrying with Deep Research...")
                self.generate_article(
                    topic=topic,
                    output_file=output_file,
                    format=format,
                    online=True,
                    research_iter=research_iter,
                    length=length
                )

    def generate_code(self, prompt, output_file=None):
        """Generate Code from Prompt (supports multi-file output)."""
        self._load_model()
        
        from rich.console import Console
        console = Console()
        console.print(f"💻 Generating Code for: '{prompt}'...")
        
        system_prompt = (
            "You are an expert coding assistant. Write clean, efficient, and well-commented code "
            "based on the user's request. Return ONLY the code blocks. "
            "IMPORTANT: Before EACH code file, include a comment line with the filename, "
            "e.g., '# filename: my_script.py' or '// filename: src/utils.js'. "
            "You can use folder paths like 'src/module/file.py'. "
            "If multiple files are needed, separate them with filename comments. "
            "Do not include markdown backticks or explanations unless asked. "
            "Make sure the filename extension matches the code language."
        )
        user_prompt = f"Request: {prompt}\n\nCode:"
        
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
        
        full_prompt = self.pipeline.tokenizer.apply_chat_template(
            messages, tokenize=False, add_generation_prompt=True
        )
        
        # Show thinking spinner during generation
        # Add space to text to align with '💻' icon above (which is often double-width)
        console.print()
        with console.status("[yellow] Thinking...[/yellow]", spinner="dots"):
            outputs = self.pipeline(
                full_prompt, 
                max_new_tokens=4096, 
                do_sample=True, 
                temperature=0.2,
                top_p=0.9,
            )
        
        generated_text = outputs[0]['generated_text']
        response = generated_text[len(full_prompt):].strip()
        
        # Remove markdown code fences if present
        import re
        response = re.sub(r"```\w*\n?", "", response)
        
        # Parse multiple files from response
        # Matches: "# filename: path.ext" OR "# path.ext" OR "// filename: path.ext" OR "// path.ext"
        # The filename must contain at least one "/" or have a file extension
        file_pattern = re.compile(
            r"^(?:#|//)\s*(?:filename:\s*)?([^\s]+\.(?:py|js|ts|jsx|tsx|html|css|java|cpp|c|h|go|rs|rb|php|sh|sql|json|yaml|yml|md|txt))\s*$",
            re.IGNORECASE | re.MULTILINE
        )
        
        # Split response by filename markers
        parts = file_pattern.split(response)
        
        files_to_write = []
        
        if len(parts) > 1:
            # Multi-file output: parts = [preamble, filename1, content1, filename2, content2, ...]
            # Skip preamble (parts[0]), then pair (filename, content)
            for i in range(1, len(parts), 2):
                if i + 1 < len(parts):
                    filename = parts[i].strip()
                    content = parts[i + 1].strip()
                    if filename and content:
                        files_to_write.append((filename, content))
                elif parts[i].strip():
                    # Last filename without content (edge case)
                    pass
        else:
            # Single file output (no filename marker found)
            content = response.strip()
            if output_file:
                if "." in os.path.basename(output_file):
                    files_to_write.append((output_file, content))
                else:
                    ext = self._infer_extension(content)
                    files_to_write.append((f"{output_file}{ext}", content))
            else:
                ext = self._infer_extension(content)
                files_to_write.append((f"generated_code_{int(time.time())}{ext}", content))
        
        # Write all files
        # Check if output_file is an existing directory once
        output_is_dir = output_file and os.path.isdir(output_file)
        
        # Respect --force flag to bypass overwrite prompts
        always_overwrite = self.args.force if hasattr(self, 'args') else False
        never_overwrite = False
        
        for filepath, content in files_to_write:
            try:
                final_path = filepath
                
                if output_file:
                    if output_is_dir:
                        # User provided an EXISTING directory -> Use as base
                        final_path = os.path.join(output_file, filepath)
                    elif len(files_to_write) == 1:
                        # User provided non-directory path (or non-existent) AND single file
                        # Treat as filename
                        final_path = output_file
                    else:
                        # Multi-file AND non-directory output path
                        pass 

                # Overwrite check
                should_write, final_path, always_overwrite, never_overwrite = check_overwrite(
                    final_path, always_overwrite, never_overwrite
                )
                
                if final_path is None: # User cancelled/Back
                    print("🛑 Code generation cancelled.")
                    break
                
                if not should_write:
                    continue

                # Create directories if needed
                dir_path = os.path.dirname(final_path)
                if dir_path:
                    os.makedirs(dir_path, exist_ok=True)
                
                with open(final_path, "w", encoding="utf-8") as f:
                    f.write(content)
                print(f"✅ Code saved to: {final_path}")
            except Exception as e:
                print(f"❌ Error saving {filepath}: {e}")
                
    def _infer_extension(self, code_content):
        """Infer file extension from code content."""
        if "import " in code_content or "def " in code_content or "print(" in code_content or "class " in code_content:
            return ".py"
        elif "function " in code_content or "const " in code_content or "let " in code_content or "console.log" in code_content:
            return ".js"
        elif "#include" in code_content:
            return ".cpp"
        elif "public class" in code_content:
            return ".java"
        elif "<html" in code_content:
            return ".html"
        elif "package main" in code_content:
            return ".go"
        elif "fn main" in code_content or "use std::" in code_content:
            return ".rs"
        return ".txt"

    def _save_formatted(self, markdown_text, filename, fmt, online=False):
        """Convert and save to specific format.
        
        Returns:
            int: Number of failed image fetches (for offline mode warning)
        """
        failed_image_count = 0
        
        base, _ = os.path.splitext(filename)
        
        # Ensure we have the right extension
        if not filename.lower().endswith(f".{fmt}"):
            filename = f"{base}.{fmt}"
            
        print(f"💾 Saving as {fmt.upper()}...")
        
        if fmt == "md":
            with open(filename, "w", encoding="utf-8") as f:
                f.write(markdown_text)
                
        elif fmt == "html" or fmt == "xhtml":
            html = markdown.markdown(markdown_text, extensions=['extra', 'codehilite'])
            # Wrap in basic HTML structure
            enc = "utf-8"
            full_html = (
                f"<!DOCTYPE html><html><head><meta charset='utf-8'><title>Article</title>"
                f"<style>body{{font-family:sans-serif;max-width:800px;margin:2em auto;padding:1em;line-height:1.6}}"
                f"pre{{background:#f4f4f4;padding:1em;border-radius:5px}}</style></head>"
                f"<body>{html}</body></html>"
            )
            with open(filename, "w", encoding=enc) as f:
                f.write(full_html)
                
        elif fmt == "json":
            import json
            data = {"content": markdown_text, "html": markdown.markdown(markdown_text)}
            with open(filename, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2)

        elif fmt == "docx":
            import io
            import urllib.request
            import re as re_module
            from docx.shared import Inches
            
            doc = docx.Document()
            MIN_IMAGE_SIZE = 5 * 1024  # 5KB threshold
            
            def fetch_image_for_docx(url):
                """Fetch image and return as BytesIO for docx embedding."""
                nonlocal failed_image_count
                try:
                    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
                    with urllib.request.urlopen(req, timeout=10) as response:
                        image_data = response.read()
                    
                    if len(image_data) < MIN_IMAGE_SIZE:
                        failed_image_count += 1
                        print(f"⚠️  Image too small (likely placeholder): {url[:50]}...")
                        return None
                    
                    return io.BytesIO(image_data)
                except Exception as e:
                    failed_image_count += 1
                    print(f"⚠️  Could not fetch image: {url[:50]}... ({e})")
                    return None
            
            # Process markdown line by line
            for line in markdown_text.split('\n'):
                # Check for markdown image: ![alt](url)
                img_match = re_module.match(r'!\[([^\]]*)\]\((https?://[^\)]+)\)', line)
                if img_match:
                    alt_text = img_match.group(1)
                    img_url = img_match.group(2)
                    img_stream = fetch_image_for_docx(img_url)
                    if img_stream:
                        try:
                            doc.add_picture(img_stream, width=Inches(5))
                            # Add caption if alt text exists
                            if alt_text:
                                caption = doc.add_paragraph(alt_text)
                                caption.alignment = 1  # Center
                        except Exception as e:
                            print(f"⚠️  Could not embed image: {e}")
                            doc.add_paragraph(f"[Image: {alt_text}]")
                    else:
                        doc.add_paragraph(f"[Image: {alt_text}]")
                elif line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                else:
                    doc.add_paragraph(line)
            doc.save(filename)
        
        elif fmt == "rtf":
            # Markdown -> RTF conversion with image support
            import urllib.request
            import re as re_module
            import binascii
            
            MIN_IMAGE_SIZE = 5 * 1024  # 5KB threshold
            
            def rtf_escape(text):
                return text.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
            
            def fetch_image_for_rtf(url):
                """Fetch image and return hex-encoded data with format info."""
                nonlocal failed_image_count
                try:
                    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
                    with urllib.request.urlopen(req, timeout=10) as response:
                        image_data = response.read()
                    
                    if len(image_data) < MIN_IMAGE_SIZE:
                        failed_image_count += 1
                        print(f"⚠️  Image too small (likely placeholder): {url[:50]}...")
                        return None
                    
                    # Determine format from header bytes
                    if image_data[:3] == b'\xff\xd8\xff':
                        img_format = 'jpegblip'
                    elif image_data[:8] == b'\x89PNG\r\n\x1a\n':
                        img_format = 'pngblip'
                    else:
                        # Default to JPEG
                        img_format = 'jpegblip'
                    
                    # Hex encode the image data
                    hex_data = binascii.hexlify(image_data).decode('ascii')
                    return (img_format, hex_data)
                except Exception as e:
                    failed_image_count += 1
                    print(f"⚠️  Could not fetch image: {url[:50]}... ({e})")
                    return None
            
            rtf_lines = []
            rtf_lines.append(r'{\rtf1\ansi\deff0')
            rtf_lines.append(r'{\fonttbl{\f0 Helvetica;}{\f1 Courier;}}')
            rtf_lines.append(r'{\colortbl;\red0\green0\blue0;\red51\green51\blue51;}')
            rtf_lines.append(r'\f0\fs24')  # Default font 12pt
            
            for line in markdown_text.split('\n'):
                # Check for markdown image: ![alt](url)
                img_match = re_module.match(r'!\[([^\]]*)\]\((https?://[^\)]+)\)', line)
                if img_match:
                    alt_text = img_match.group(1)
                    img_url = img_match.group(2)
                    img_data = fetch_image_for_rtf(img_url)
                    if img_data:
                        img_format, hex_data = img_data
                        # RTF picture: width ~400 pixels (8000 twips), scale to fit
                        rtf_lines.append(r'\pard\qc\sb200\sa100')
                        rtf_lines.append(r'{\pict\\' + img_format + r'\picwgoal6000\pichgoal4000')
                        rtf_lines.append(hex_data)
                        rtf_lines.append(r'}')
                        if alt_text:
                            rtf_lines.append(r'\pard\qc\i\fs20 ' + rtf_escape(alt_text) + r'\i0\fs24\par')
                    else:
                        rtf_lines.append(r'\pard\sa100 [Image: ' + rtf_escape(alt_text) + r']\par')
                    continue
                
                line = rtf_escape(line)
                if line.startswith('# '):
                    rtf_lines.append(r'\pard\sb400\sa200\b\fs48 ' + line[2:] + r'\b0\fs24\par')
                elif line.startswith('## '):
                    rtf_lines.append(r'\pard\sb300\sa150\b\fs36 ' + line[3:] + r'\b0\fs24\par')
                elif line.startswith('### '):
                    rtf_lines.append(r'\pard\sb200\sa100\b\fs28 ' + line[4:] + r'\b0\fs24\par')
                elif line.startswith('- ') or line.startswith('* '):
                    rtf_lines.append(r'\pard\li720\fi-360\bullet  ' + line[2:] + r'\par')
                elif line.startswith('```'):
                    continue
                elif line.strip():
                    rtf_lines.append(r'\pard\sa100 ' + line + r'\par')
                else:
                    rtf_lines.append(r'\par')
            
            rtf_lines.append('}')
            
            with open(filename, "w", encoding="utf-8") as f:
                f.write('\n'.join(rtf_lines))
        elif fmt == "pdf":
            # Download remote images and convert to base64 data URIs
            import base64
            import urllib.request
            import re as re_module
            
            # Pre-process: Convert markdown links that look like images to proper image syntax
            # Model often outputs [Image: ...](url) instead of ![Image: ...](url)
            processed_md = re_module.sub(
                r'\[([Ii]mage[^\]]*)\]\((https?://[^\)]+\.(jpg|jpeg|png|gif|webp)[^\)]*)\)',
                r'![\1](\2)',
                markdown_text
            )
            
            # Remove bolding from table rows to prevent xhtml2pdf artifacts (double printing)
            # Line-by-line approach is more robust than regex
            md_lines = processed_md.split('\n')
            for idx, line in enumerate(md_lines):
                if line.strip().startswith('|'):
                    md_lines[idx] = line.replace('**', '')
            processed_md = '\n'.join(md_lines)
            
            # Convert MD -> HTML with fenced code blocks support
            # 'toc' extension generates IDs for headers (e.g. #header-name), needed for internal links
            html_content = markdown.markdown(processed_md, extensions=['extra', 'fenced_code', 'tables', 'toc'])
            
            # Replace symbols with text equivalents - xhtml2pdf doesn't support Unicode symbols
            # Map common symbols to ASCII equivalents
            symbol_replacements = {
                '✓': '[Y]', '✔': '[Y]', '☑': '[Y]',  # Checkmarks
                '✗': '[N]', '✘': '[N]', '☐': '[ ]',  # X marks and empty boxes
                '■': '[*]', '□': '[ ]', '●': '[*]',  # Filled/empty shapes
                '★': '[*]', '☆': '[ ]',              # Stars
                '→': '->', '←': '<-', '↔': '<->',    # Arrows
                '📦': '', '📚': '', '📄': '',        # Common emojis
                '🖼': '', '🎬': '', '🎵': '',
                '📝': '', '📰': '', '💻': '',
                '💬': '', '✨': '', '🔄': '',
                '📈': '', '🧪': '', 'ℹ': '',
                '❌': '[X]', '⚠': '[!]', '✅': '[Y]',
                '🪄': '', '🔒': '[LOCK]',
            }
            
            for symbol, replacement in symbol_replacements.items():
                html_content = html_content.replace(symbol, replacement)
            
            # Strip remaining emojis that weren't explicitly mapped
            def strip_emojis(text):
                """Remove emoji characters that xhtml2pdf can't render."""
                result = []
                for char in text:
                    code = ord(char)
                    # Skip emoji and symbol ranges
                    if (0x1F300 <= code <= 0x1FFFF or  # All emoji ranges
                        0x2600 <= code <= 0x27BF or    # Misc Symbols & Dingbats
                        0x2300 <= code <= 0x23FF or    # Misc Technical
                        0xFE00 <= code <= 0xFE0F):     # Variation selectors
                        continue
                    result.append(char)
                return ''.join(result)
            
            html_content = strip_emojis(html_content)
            
            html_content = strip_emojis(html_content)
            
            def fetch_and_encode_image(url):
                """Fetch remote image and return base64 data URI."""
                nonlocal failed_image_count
                MIN_IMAGE_SIZE = 5 * 1024  # 5KB - filter out placeholder/blocked images
                try:
                    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
                    with urllib.request.urlopen(req, timeout=10) as response:
                        image_data = response.read()
                    
                    # Check for likely placeholder/geo-blocked images (too small)
                    if len(image_data) < MIN_IMAGE_SIZE:
                        failed_image_count += 1
                        print(f"⚠️  Image too small (likely placeholder): {url[:50]}... ({len(image_data)//1024}KB)")
                        return None
                    
                    # Determine MIME type
                    content_type = response.headers.get('Content-Type', 'image/jpeg')
                    if 'png' in url.lower():
                        content_type = 'image/png'
                    elif 'gif' in url.lower():
                        content_type = 'image/gif'
                    elif 'webp' in url.lower():
                        content_type = 'image/webp'
                    
                    b64_data = base64.b64encode(image_data).decode('utf-8')
                    return f'data:{content_type};base64,{b64_data}'
                except Exception as e:
                    failed_image_count += 1
                    print(f"⚠️  Could not fetch image: {url[:60]}... ({e})")
                    return None
            
            def replace_src(match):
                url = match.group(1)
                data_uri = fetch_and_encode_image(url)
                if data_uri:
                    return f'src="{data_uri}"'
                return match.group(0)
            
            def replace_href_with_img(match):
                """Convert <a href="image_url">text</a> to <img src="data:...">"""
                url = match.group(1)
                alt_text = match.group(2)
                data_uri = fetch_and_encode_image(url)
                if data_uri:
                    return f'<img src="{data_uri}" alt="{alt_text}" style="max-width:100%">'
                return match.group(0)
            
            # Replace img src attributes
            html_content = re_module.sub(r'src="(https?://[^"]+)"', replace_src, html_content)
            
            # Also convert anchor links that point to images
            html_content = re_module.sub(
                r'<a href="(https?://[^"]+\.(?:jpg|jpeg|png|gif|webp)[^"]*)">([^<]+)</a>',
                replace_href_with_img,
                html_content,
                flags=re_module.IGNORECASE
            )
            
            # Add ID attributes to headings for internal anchor links
            # This makes #section-name links work in PDF
            def add_heading_ids(match):
                tag = match.group(1)
                content = match.group(2)
                # Generate ID from heading text (lowercase, replace spaces with dashes)
                heading_id = re_module.sub(r'[^\w\s-]', '', content.lower())
                heading_id = re_module.sub(r'[-\s]+', '-', heading_id).strip('-')
                return f'<{tag} id="{heading_id}">{content}</{tag}>'
            
            html_content = re_module.sub(
                r'<(h[1-6])>([^<]+)</\1>',
                add_heading_ids,
                html_content
            )
            
            # CRITICAL: xhtml2pdf doesn't handle <pre> whitespace properly
            # Convert newlines in <pre> and <code> blocks to <br/> tags
            def fix_pre_blocks(match):
                content = match.group(1)
                # Replace newlines with <br/> and preserve indentation with &nbsp;
                lines = content.split('\n')
                fixed_lines = []
                for line in lines:
                    # Count leading spaces and convert to &nbsp;
                    stripped = line.lstrip(' ')
                    indent = len(line) - len(stripped)
                    nbsp_indent = '&nbsp;' * indent
                    fixed_lines.append(nbsp_indent + stripped)
                return '<pre><code>' + '<br/>'.join(fixed_lines) + '</code></pre>'
            
            html_content = re_module.sub(r'<pre><code[^>]*>(.*?)</code></pre>', fix_pre_blocks, html_content, flags=re_module.DOTALL)
            html_content = re_module.sub(r'<pre>(.*?)</pre>', fix_pre_blocks, html_content, flags=re_module.DOTALL)
            
            # Add explicit column widths via inline styles on header cells
            # xhtml2pdf needs explicit widths for proper table layout
            def add_table_column_widths(html):
                """Add width styles to table header cells based on column count."""
                
                def add_width_to_th(th_html, width_pct):
                    """Add or merge width style into a th element."""
                    # Check if th has existing style
                    if 'style="' in th_html:
                        # Append width to existing style
                        return th_html.replace('style="', f'style="width:{width_pct}%; ')
                    else:
                        # Add new style attribute
                        return th_html.replace('<th', f'<th style="width:{width_pct}%"', 1)
                
                def process_table(table_match):
                    table_html = table_match.group(0)
                    
                    # Find all th elements in thead
                    thead_match = re_module.search(r'<thead>(.*?)</thead>', table_html, flags=re_module.DOTALL)
                    if not thead_match:
                        return table_html
                    
                    thead_content = thead_match.group(1)
                    th_elements = re_module.findall(r'<th[^>]*>.*?</th>', thead_content, flags=re_module.DOTALL)
                    col_count = len(th_elements)
                    
                    if col_count == 0:
                        return table_html
                    
                    headers_text = [re_module.sub(r'<[^>]+>', '', th).strip().lower() for th in th_elements]
                    
                    # Content-aware width maps
                    if 'menu #' in headers_text or 'jump point' in headers_text:
                         # Fast Jump Points: | Menu # | Task | Jump Point | Description |
                         widths = [10, 20, 30, 40]
                    elif 'model' in headers_text and 'vram' in headers_text:
                         # Text Models: | Model | VRAM | RAM | Notes |
                         widths = [35, 12, 12, 41]
                    elif 'argument' in headers_text:
                         # Upscaling: | Argument | Description | Default |
                         widths = [35, 45, 20]
                    elif 'goal' in headers_text and col_count == 2:
                         # Transformation: | Goal | Command Pattern |
                         widths = [25, 75]
                    else:
                        # Fallback defaults
                        width_maps = {
                            2: [25, 75],
                            3: [33, 34, 33],
                            4: [25, 25, 25, 25],
                            5: [10, 15, 25, 25, 25],
                            6: [10, 15, 15, 15, 15, 30],
                            7: [10, 12, 12, 12, 12, 12, 30],
                            8: [8, 12, 10, 10, 10, 10, 10, 30],
                        }
                        widths = width_maps.get(col_count, [100 // col_count] * col_count)
                    
                    # Replace each th with width-styled version
                    new_thead_content = thead_content
                    for i, th in enumerate(th_elements):
                        if i < len(widths):
                            styled_th = add_width_to_th(th, widths[i])
                            new_thead_content = new_thead_content.replace(th, styled_th, 1)
                    
                    return table_html.replace(thead_match.group(1), new_thead_content)
                
                return re_module.sub(r'<table>.*?</table>', process_table, html, flags=re_module.DOTALL)
            
            html_content = add_table_column_widths(html_content)
            
            # Fix internal anchor links - ensure href anchors match heading IDs exactly
            # xhtml2pdf needs explicit name attributes on targets
            def fix_anchor_targets(match):
                tag = match.group(1)
                heading_id = match.group(2)
                content = match.group(3)
                # Add both id and name attributes for maximum compatibility
                return f'<{tag} id="{heading_id}"><a name="{heading_id}"></a>{content}</{tag}>'
            
            html_content = re_module.sub(
                r'<(h[1-6]) id="([^"]+)">([^<]+)</\1>',
                fix_anchor_targets,
                html_content
            )
            
            # Wrap in full HTML document for xhtml2pdf
            full_html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
@page {{
    size: a4 portrait;
    margin: 1.5cm;
}}
body {{ 
    font-family: Helvetica, sans-serif; 
    font-size: 9pt; 
    line-height: 1.4; 
}}
h1 {{ font-size: 18pt; color: #333; margin-top: 0.8em; margin-bottom: 0.4em; page-break-after: avoid; }}
h2 {{ font-size: 14pt; color: #444; margin-top: 0.6em; margin-bottom: 0.3em; page-break-after: avoid; }}
h3 {{ font-size: 12pt; color: #555; margin-top: 0.5em; margin-bottom: 0.2em; page-break-after: avoid; }}
h4 {{ font-size: 10pt; color: #666; margin-top: 0.4em; margin-bottom: 0.2em; }}
p {{ margin: 0.3em 0; }}
pre {{ 
    background: #f4f4f4; 
    padding: 6px; 
    border: 1px solid #ddd;
    font-family: Courier, monospace; 
    font-size: 6pt;
    line-height: 1.2;
    word-wrap: break-word;
    word-break: break-all;
    overflow-wrap: break-word;
}}
code {{ 
    background: #f0f0f0; 
    padding: 1px 2px;
    font-family: Courier, monospace;
    font-size: 6pt;
    word-wrap: break-word;
    word-break: break-all;
}}
/* Lists - explicit styling for xhtml2pdf */
ul {{ 
    margin: 0.4em 0 0.4em 1.5em; 
    padding: 0;
}}
ol {{ 
    margin: 0.4em 0 0.4em 1.5em; 
    padding: 0;
}}
li {{ 
    margin: 0.15em 0; 
    padding-left: 0.3em;
}}
ul li {{ list-style-type: disc; }}
ul ul li {{ list-style-type: circle; }}
ol li {{ list-style-type: decimal; }}
img {{ max-width: 100%; height: auto; }}
/* Tables - auto layout for content-aware column sizing */
table {{ 
    border-collapse: collapse; 
    width: 100%; 
    font-size: 6pt;
    margin: 0.4em 0;
}}
th, td {{ 
    border: 1px solid #999; 
    padding: 2px 4px; 
    text-align: left;
    word-wrap: break-word;
}}
th {{ background: #e8e8e8; font-weight: bold; }}
a {{ color: #0066cc; text-decoration: underline; }}
blockquote {{
    border-left: 2px solid #999;
    margin: 0.4em 0;
    padding-left: 0.8em;
    color: #555;
    font-size: 8pt;
}}
/* GitHub-style alerts */
.alert {{
    border-left: 3px solid #999;
    padding: 0.5em 0.8em;
    margin: 0.5em 0;
    background: #f8f8f8;
    font-size: 8pt;
}}
</style>
</head>
<body>
{html_content}
</body>
</html>"""
            
            with open(filename, "wb") as f:
                pisa_status = pisa.CreatePDF(full_html, dest=f)
            
            if pisa_status.err:
                print("❌ PDF conversion failed")
            else:
                 print(f"✅ PDF saved: {filename}")
                 
        elif fmt == "txt":
            # Strip markdown chars via BS4
            html = markdown.markdown(markdown_text)
            text = BeautifulSoup(html, "html.parser").get_text()
            with open(filename, "w", encoding="utf-8") as f:
                f.write(text)
                
        else:
            print(f"⚠️ Unknown format '{fmt}', saving as MD.")
            with open(f"{base}.md", "w", encoding="utf-8") as f:
                f.write(markdown_text)

        print(f"✅ Saved to {filename}")
        return failed_image_count


# --- Main Logic ---

def run_unit_tests(module_name=None, verbose=False):
    """Run unit tests using python's unittest module."""
    import subprocess
    target = module_name if module_name else "tests.ai-media_test"
    
    print(f"{emoji('🧪 ', '')}Running Unit Tests: {target}\n")
    print("=" * 60)
    
    cmd = [sys.executable, "-m", "unittest", target]
    if verbose:
        cmd.append("-v")
    
    env = os.environ.copy()
    env["PYTHONIOENCODING"] = "utf-8"
    
    # Run synchronously
    try:
        subprocess.run(cmd, env=env, check=False) # check=False to custom handle exit
    except Exception as e:
        print(f"Error running tests: {e}")
        sys.exit(1)



class CleanHelpFormatter(argparse.RawTextHelpFormatter):
    """Custom formatter that hides metavar and uses wider columns."""
    def __init__(self, prog):
        super().__init__(prog, max_help_position=40, width=120)
    
    def _format_action_invocation(self, action):
        if not action.option_strings:
            return super()._format_action_invocation(action)
        # For options with nargs or that take values, hide the metavar
        if action.nargs != 0 and action.option_strings:
            # Return just the option strings without the metavar
            return ', '.join(action.option_strings)
        return super()._format_action_invocation(action)

def main():
    # Cancel loading message timer (imports complete)
    if _loading_timer:
        _loading_timer.cancel()
    
    parser = argparse.ArgumentParser(
        description="Generate, describe, upscale, and convert media using AI and FFmpeg.",
        formatter_class=CleanHelpFormatter,
        epilog="""
Examples:
  -- Image Generation --
  python ai-media.py -i -p "Cyberpunk city" -o city.png -s 720p
  python ai-media.py -i -p "Forest" -o forest.jpg -s 4k
  
  -- Video Generation --
  python ai-media.py -v -p "Robot dancing" -o robot.mp4 -l 5s
  python ai-media.py -v -p "Camera pans left" -ii ./start.png -o output.mp4 (Image-to-Video)
  python ai-media.py -v -p "Dancer" -ap "Techno beat" -o party.mp4 (Video+Audio Mux)
  python ai-media.py -v -p "Future city" -o city.mp4 --video-model wan2.2 (SOTA Quality)
  python ai-media.py -v -p "Driving car" -o car.mp4 --video-model ltx-video (Fast/High-Res)
  python ai-media.py -v -p "Liquid simulation" --video-model mochi-1 (High Motion)
  python ai-media.py -v -p "Cinematic panda" --video-model hunyuan (13B Model)
  
  -- Audio Generation --
  python ai-media.py -a -p "Jazz saxophone" -o jazz.mp3 -l 30s
  python ai-media.py -a -p "Rainforest" -o rain.wav --audio-model audioldm2
  python ai-media.py -a -p "Spooky" -ii ./haunted.jpg -o spooky.mp3 (Image-to-Audio)
  python ai-media.py -a -ii ./image.jpg -cm blip (Image-to-Audio w/ BLIP)
  python ai-media.py -a -ii ./image.jpg (Auto-caption + Audio)
  python ai-media.py -a -ii ./video.mp4 (Auto-caption Video + Audio)
  python ai-media.py -a -p "♪ In the jungle ♪ [laughter]" --audio-model bark (Bark Creative)

  -- Text Generation (Articles, Research, Chat, Code) --
  python ai-media.py -ga -p "Future of AI" -o article.md (Offline)
  python ai-media.py -gr -p "Latest Quantum Computing News" -o research.pdf (Online Deep Research)
  python ai-media.py -c --chat-model llama-3.1-8b (Interactive Chat)
  python ai-media.py -gc 'Write a Snake game in Python' (Code Gen)
  python ai-media.py -ga -p "Story about a cat" -o story.docx -atm mistral-nemo-12b

  -- Generate Description --
  python ai-media.py -gd -ii video.mp4
  python ai-media.py -gd -ii image.jpg -cm blip (Use simpler model)

  -- Creative Image Transformation --
  python ai-media.py -ti "photo.jpg" -p "Make it look like an anime drawing"
  python ai-media.py -ti "photo.jpg" -p "Make it anime" -o "edits/anime_version.png"
  python ai-media.py -ti "photo.jpg" --remove-background
  python ai-media.py -ti "photo.jpg" --remove-background -o "no_bg/photo_clean.png"

  -- Media Conversion --
  python ai-media.py -ci photo.gif -cit png
  python ai-media.py -cv clip.mov -cvt mp4
  python ai-media.py -ca song.wav -cat mp3
  
  -- AI Upscaling --
  python ai-media.py -ui input.jpg -uf 2x
  python ai-media.py -ui input.jpg -uf 4x
  python ai-media.py -ui input.jpg -uf 4x -su (Simple Upscale)
  python ai-media.py -uv input.mp4 -vu realesrgan (Fast AI - Recommended)
  python ai-media.py -uv input.mp4 -uf 2x -vu sd (High Detail AI)


Supported Models:
  Images:
    - sdxl (default)           : ~8GB  | Fast, high quality.
    - sd-1.5                   : ~4GB  | Lightweight, lower VRAM.
    - flux                     : ~24GB | High quality (🔒 Gated - Free Login Required)
    - flux-dev                 : ~24GB | Professional creative work (🔒 Gated - Free Login Required)
  
  Video:
    - zeroscope (default)      : ~4GB  | Fast, no watermarks. Auto-upscales with XL.
    - ms-1.7b                  : ~10GB | General purpose (has watermark issues).
    - cogvideox                : ~15GB | High fidelity.
    - svd                      : ~4GB  | I2V Only.
    - wan2.2                   : ~30GB | SOTA (2025). Excellent quality.
    - ltx-video                : ~12GB | Balanced speed/quality. Good motion.
    - mochi-1                  : ~19GB | High motion fidelity.
    - hunyuan                  : ~25GB | Massive scale.
    
  Audio:
    - musicgen-small           : ~2GB  | Fast, lightweight. Good for quick sketches.
    - musicgen-medium (default): ~6GB  | Balanced quality/speed.
    - musicgen-large           : ~10GB | High fidelity. Slower.
    - audioldm2                : ~4GB  | Specialized in Sound Effects (SFX), foley, environmental.
    - stable-audio             : ~10GB | Variable-length, high-quality music/SFX (🔒 Gated - Free Login Required)
    - bark                     : ~4GB  | Realistic speech, music, and sound effects.
    
  Text (Articles, Research, Chat, Code):
    - deepseek-r1-qwen-7b      : ~7GB  | R1 distilled to Qwen-7B. Step-by-step reasoning. (Ungated)
    - deepseek-r1-qwen-14b     : ~14GB | R1 distilled to Qwen-14B. Better reasoning. (Ungated)
    - deepseek-r1-qwen-32b     : ~24GB | ⚠️ HIGH RAM! R1 distilled to Qwen-32B. (Ungated)
    - deepseek-r1-llama-8b     : ~8GB  | R1 distilled to Llama-8B. Reasoning-focused. (Ungated)
    - deepseek-r1-llama-70b    : ~40GB | ⚠️ HIGH RAM! R1 distilled to Llama-70B. (Ungated)
    - llama-3.1-8b (default)   : ~16GB | Writing, chat, and reasoning (🔒 Gated - Free Login Required)
    - mistral-nemo-12b         : ~24GB | Powerful 12B model. Large context and reasoning.
    - qwen3-8b                 : ~16GB | Latest Qwen model. Strong instruction-following.
    - qwen-2.5-14b             : ~28GB | Larger Qwen model. Great at detailed formatting.

  Description Generation:
    - florence (default)       : ~1.5GB | SOTA details, rich descriptions, "seeing" the scene.
    - blip                     : ~1GB   | Simple, concise captions. Faster but less detailed.

  Creative Image Transformation:
    - instruct-pix2pix         : ~4GB  | Instructional image editing (e.g., "Make it anime").
    - instruct-pix2pix-sdxl    : ~8GB  | High quality, slow.
    - remove-bg                : ~1GB  | Background removal and silhouette creation.

  Upscaling:
    - x2 (≤2x factor)          : ~4GB  | Fast, preserves original style.
    - x4 (>2x factor)          : ~8GB  | High detail, sharpens textures.
    - Real-ESRGAN x4plus       : ~0.3GB| Fast, faithful upscaling, better temporal consistency.
        """
    )
    
    # Generation Mode (First - what do you want to create?)
    mode_group = parser.add_argument_group("Generation Mode")
    mode_group.add_argument("-i", "--generate-image", action="store_true", help="Generate Image")
    mode_group.add_argument("-v", "--generate-video", action="store_true", help="Generate Video")
    mode_group.add_argument("-a", "--generate-audio", action="store_true", help="Generate Audio")
    
    # NEW: Article Modes
    mode_group.add_argument("-ga", "--generate-article", action="store_true", help="Generate Article (Offline)")
    mode_group.add_argument("-gr", "--generate-research", action="store_true", help="Generate Article + Research (Online)")
    mode_group.add_argument("-c", "--chat", action="store_true", help="Interactive Chat Mode")
    mode_group.add_argument("-gc", "--generate-code", nargs="?", const=True, default=False, 
                            help="Generate code. Supports multi-file projects & auto-naming. E.g.: -gc 'Write a REST API'")
    
    mode_group.add_argument("-gd", "--generate-description", nargs="?", const="USE_INPUT_IMAGE", help="Generate Description (Caption) for Image or Video.")
    mode_group.add_argument("-ti", "--transform-image", nargs="?", const="USE_GENERATED", metavar="FILE", help="Transform an image. Omit FILE to auto-use generated output from -i.")
    
    # Common Parameters (applies to most modes)
    common_group = parser.add_argument_group("Common Parameters")
    common_group.add_argument("-p", "--prompt", required=False, help="Text prompt description (Required for generation modes)")
    common_group.add_argument("-o", "--output", help="Output file path. Auto-generated from prompt if omitted.")
    common_group.add_argument("--force", action="store_true", help="Skip all confirmation prompts (overwrites files, ignores resource warnings).")
    common_group.add_argument("-f", "--format", help="File format. Image: jpg/png (default: jpg). Video: mp4. Audio: mp3/wav (default: mp3). Article: md/pdf/doc/html.")
    common_group.add_argument("-s", "--size", help="Resolution for Image/Video: '720p', '1080p', '4k', '1280x720'. For zeroscope: triggers dynamic upscaling (XL + Real-ESRGAN) for targets > 576x320. Default: 720p")
    common_group.add_argument("-npt", "--no-performance-tracking", action="store_true", help="Disable performance tracking (performance.json).")
    
    
    # Specific options
    image_group = parser.add_argument_group("Image Options")
    image_models_help = [k + " (Gated)" if k in ["flux", "flux-dev"] else k for k in IMAGE_MODELS.keys()]
    image_group.add_argument("--image-model", default="default", help=f"Model: {', '.join(image_models_help)}")
    image_group.add_argument("-otn", "--orientation", choices=["landscape", "portrait", "square"], default="landscape",
                              help="Orientation for SDXL/Flux generation. 'portrait' swaps width/height.")
    image_group.add_argument("--unsafe", action="store_true", help="Disable NSFW safety checker (Use with caution).")
    
    video_group = parser.add_argument_group("Video Options")
    video_group.add_argument("--video-model", default="default", help=f"Model: {', '.join(VIDEO_MODELS.keys())} (default: zeroscope)")
    video_group.add_argument("-l", "--length", default="2s", help="Duration (e.g. '2s', '5s', '1m', '{m:1, s:30}'). Default: 2s")
    video_group.add_argument("-ii", "--input-image", help="Input image for Image-to-Video generation.")
    video_group.add_argument("-ap", "--audio-prompt", help="Audio prompt for 'Video with Audio' generation (merged via FFmpeg).")
    video_group.add_argument("-vc", "--video-codec", choices=['auto', 'h264', 'hevc', 'av1'], default='auto',
                             help="Video Codec: h264, hevc, av1, or auto. AV1 uses hardware encoder (av1_nvenc) when available. Default: auto")
    
    audio_group = parser.add_argument_group("Audio Options")
    audio_models_help = [k + " (Gated)" if k in ["stable-audio"] else k for k in AUDIO_MODELS.keys()]
    audio_group.add_argument("-am", "--audio-model", default="default", help=f"Model: {', '.join(audio_models_help)}")
    audio_group.add_argument("--voice-preset", default="v2/en_speaker_6", help="Bark Voice Preset (e.g. 'v2/en_speaker_6'). Default: v2/en_speaker_6")
    audio_group.add_argument("-m", "--sampling-rate", type=str, default="32000", help="Sampling rate (e.g. 32000, 44.1k). Default: 32000.")
    audio_group.add_argument("-b", "--bit-depth", type=int, choices=[16, 24, 32], default=16, help="Bit depth for audio conversion.")
    audio_group.add_argument("-r", "--bit-rate", help="Bit rate (e.g. 192k) for audio conversion.")
    
    # NEW: Text Options
    text_group = parser.add_argument_group("Text/Article Options (Articles, Research, Chat, Code)")
    text_models_help = [k + " (Gated)" if k in ["llama-3.1-8b"] else k for k in TEXT_MODELS.keys()]
    text_group.add_argument("-atm", "--article-model", default="default", 
                            help=f"Model for articles/research (-ga/-gr). Options: {', '.join(text_models_help)}")
    text_group.add_argument("-chm", "--chat-model", default="default", 
                            help=f"Model for chat (-c). Options: {', '.join(text_models_help)}")
    text_group.add_argument("-cdm", "--code-model", default="default", 
                            help=f"Model for code gen (-gc). Options: {', '.join(text_models_help)}")
    text_group.add_argument("--output-format", choices=["md", "pdf", "docx", "rtf", "html", "xhtml", "json", "txt"], default="md", 
                            help="Article/research output format. Default: md")
    text_group.add_argument("-ri", "--research-iter", type=int, default=3, 
                            help="Deep research: number of sources to read. Default: 3")
    text_group.add_argument("-al", "--article-length", choices=["quick", "standard", "detailed"], default="quick",
                            help="Article length: quick (fast, ~500 words, default), standard (~1500 words), detailed (comprehensive, ~3000 words).")

    # Description Generation Options
    caption_group = parser.add_argument_group("Description Generation Options")
    caption_group.add_argument("-cm", "--caption-model", default="florence", choices=["florence", "blip"], help="Model for description generation: 'florence' (default, SOTA) or 'blip'.")
    
    # Creative Image Transformation
    transform_group = parser.add_argument_group("Creative Image Transformation Options")
    transform_group.add_argument("-tp", "--transform-prompt", help="Edit instruction for InstructPix2Pix (e.g., 'Make it anime'). Used with -ti.")
    transform_group.add_argument("-rb", "--remove-background", action="store_true", help="Remove background (Transparent PNG).")
    transform_group.add_argument("--silhouette", action="store_true", help="Create a black silhouette (requires -rb).")
    transform_group.add_argument("--image-guidance", type=float, default=1.5, help="Image guidance scale (default: 1.5). Higher = closer to original.")

    # Media Conversion (Standalone - No AI)
    convert_group = parser.add_argument_group("Media Conversion Options")
    convert_group.add_argument("-ci", "--convert-image", metavar="FILE", help="Convert image format (e.g., gif→png)")
    convert_group.add_argument("-cit", "--convert-image-to", metavar="FMT", help="Output format (png, .webp, out.jpg)")
    convert_group.add_argument("-cv", "--convert-video", metavar="FILE", help="Convert video (mov→mp4)")
    convert_group.add_argument("-cvt", "--convert-video-to", metavar="FMT", help="Output format (mp4, .webm, out.avi)")
    convert_group.add_argument("-ca", "--convert-audio", metavar="FILE", help="Convert audio (wav→mp3)")
    convert_group.add_argument("-cat", "--convert-audio-to", metavar="FMT", help="Output format (mp3, .flac, out.ogg)")
    convert_group.add_argument("--convert-image-engine", choices=["pil", "ffmpeg"], default="pil", help="pil (default) or ffmpeg")
    
    # Document Conversion
    doc_conv_group = parser.add_argument_group("Document Conversion Options")
    doc_conv_group.add_argument("-cd", "--convert-document", metavar="FILE", help="Convert document format (e.g., report.docx→pdf)")
    doc_conv_group.add_argument("-cdt", "--convert-document-to", metavar="FMT", help="Output format: md, html, pdf, docx, rtf, txt, json")
    
    # AI Upscaling (Standalone Mode)
    upscale_mode_group = parser.add_argument_group("AI Upscaling Options")
    upscale_mode_group.add_argument("-ui", "--upscale-image", metavar="FILE", help="Upscale an existing image")
    upscale_mode_group.add_argument("-uv", "--upscale-video", metavar="FILE", help="Upscale an existing video")
    upscale_mode_group.add_argument("-iu", "--image-upscaler", choices=["sd", "realesrgan"], default="realesrgan", help="Model for image upscaling: 'realesrgan' (Fast, faithful, default) or 'sd' (Stable Diffusion, slower, creative)")
    upscale_mode_group.add_argument("-vu", "--video-upscaler", choices=["sd", "realesrgan"], default="realesrgan", help="Model for video upscaling: 'realesrgan' (Fast, faithful, default) or 'sd' (Stable Diffusion, slow, detailed)")

    # Upscaling Options (applies to both standalone and chained upscaling)
    upscale_group = parser.add_argument_group("Upscaling Options")
    upscale_group.add_argument("-uf", "--upscale-factor", help="Upscale factor (e.g. '2x', '4'). Default: 2x")
    upscale_group.add_argument("--upscale", action="store_true", help="Enable AI Upscaling after generation (chained mode).")
    upscale_group.add_argument("-uof", "--upscaled-output-file", help="Custom filename for the upscaled output (e.g. 'highres.png').")
    upscale_group.add_argument("-us", "--upscale-strength", type=float, default=0.0, help="Upscale creativity/strength (0.0-1.0). Default: 0.0")
    upscale_group.add_argument("-su", "--simple-upscale", action="store_true", help="Use simple non-AI upscaling (PIL Lanczos for images, FFmpeg for videos). Very fast.")
    
    # Testing
    test_group = parser.add_argument_group("Testing")
    test_group.add_argument("--test", nargs="*", help="Run integration tests from tests/integration-tests.json (quiet mode). Optional: Space-separated list of Test Names.")
    test_group.add_argument("--test-verbose", nargs="*", help="Run integration tests with full output. Optional: Space-separated list of Test Names.")
    test_group.add_argument("--unittests", nargs="?", const="tests.ai-media_test", metavar="MODULE",
                            help="Run Python unit tests (Quiet/Summary mode). Default: all tests. Examples: tests.ai-media_test.TestParseSize")
    test_group.add_argument("--unittests-verbose", nargs="?", const="tests.ai-media_test", metavar="MODULE",
                            help="Run Python unit tests (Verbose mode). Default: all tests.")
    
    # Interactive Mode
    parser.add_argument("-I", "--interactive", nargs="?", const="menu", metavar="JUMP",
                        help="Run in interactive mode. Optional: Jump point (e.g., 'image/sdxl', 'audio/bark').")
    
    parser.add_argument("--report-json", help="Path to write a JSON report of the generation stats")
    
    global args
    args = parser.parse_args()
    
    # Interactive Mode Trigger
    if args.interactive or len(sys.argv) == 1:
        run_interactive(jump_point=args.interactive if args.interactive != "menu" else None)
        sys.exit(0)

    # Test Triggers (Priority over generation)
    if args.test is not None:
        run_tests(verbose=False, test_filter=args.test if len(args.test) > 0 else None)
    if args.test_verbose is not None:
        run_tests(verbose=True, test_filter=args.test_verbose if len(args.test_verbose) > 0 else None)
    if args.unittests is not None:
        run_unit_tests(module_name=args.unittests, verbose=False) 
        sys.exit(0)
    if args.unittests_verbose is not None:
        run_unit_tests(module_name=args.unittests_verbose, verbose=True)
        sys.exit(0)

    # Prompt Check (Required for non-chat modes)
    if not args.chat and not args.prompt and not any([args.generate_description, args.transform_image, args.convert_image, args.convert_video, args.convert_audio, args.convert_document, args.upscale_image, args.upscale_video, args.generate_code]):
        print("❌ Error: Prompt is required (use -p 'Your prompt')")
        sys.exit(1)
        
    # --- Dispatch ---
    
    # NEW: Article Generation Dispatch
    if args.chat:
        gen = ArticleGenerator(model_name=args.chat_model)
        gen.chat_session()
        sys.exit(0)
        
        sys.exit(0)

    elif args.generate_code:
        gen = ArticleGenerator(model_name=args.code_model)
        
        # Determine prompt: use -gc value if string, else use main prompt
        prompt = args.generate_code if isinstance(args.generate_code, str) else args.prompt
        
        if not prompt:
            print("❌ Error: Code generation code requires a prompt (e.g. -gc 'Write a script...' or -gc -p '...')")
            sys.exit(1)
            
        gen.generate_code(prompt, output_file=args.output)
        sys.exit(0)

    elif args.generate_article or args.generate_research:
        online = args.generate_research
        gen = ArticleGenerator(model_name=args.article_model)
        
        # Determine output format - prefer extension from -o, then --output-format
        output_format = args.output_format  # Default from --output-format flag
        
        if args.output:
            outfile = args.output
            # Infer format from filename extension if present
            _, ext = os.path.splitext(args.output)
            if ext:
                ext_lower = ext.lower().lstrip('.')
                valid_formats = ["md", "pdf", "docx", "rtf", "html", "xhtml", "json", "txt"]
                if ext_lower in valid_formats:
                    output_format = ext_lower
        else:
            # Slugify the prompt: lowercase, replace spaces with underscores, remove special chars
            import re
            slug = re.sub(r'[^\w\s-]', '', args.prompt.lower())
            slug = re.sub(r'[-\s]+', '_', slug).strip('_')[:50]  # Limit to 50 chars
            outfile = f"{slug}.{output_format}" if slug else f"article_{int(time.time())}.{output_format}"
        
        # Ensure extension matches format
        base, _ = os.path.splitext(outfile)
        if not outfile.lower().endswith(f".{output_format}"):
            outfile = f"{base}.{output_format}"
        
        # Pre-generation overwrite check (unless --force)
        if not args.force:
            should_write, outfile, _, _ = check_overwrite(outfile)
            if not should_write:
                print("⏭️  Skipped.")
                sys.exit(0)
        
        gen.generate_article(
            topic=args.prompt, 
            output_file=outfile, 
            format=output_format,
            online=online,
            research_iter=args.research_iter,
            length=args.article_length
        )
        sys.exit(0)

    # -------------------------------------------------------------------
    # Auto-generate output filename from prompt if not provided
    # This centralized logic handles all generation modes
    # -------------------------------------------------------------------
    if any([args.generate_image, args.generate_video, args.generate_audio, args.transform_image]):
        import re
        if not args.output:
            # Sanitize prompt to create safe filename (first 2 words, alphanumeric only)
            if args.prompt:
                words = re.findall(r'[a-zA-Z0-9]+', args.prompt.lower())[:2]
                if words:
                    args.output = "-".join(words)
                    print(f"ℹ️  No output specified. Using: {args.output}")
                else:
                    print("⚠️  Cannot auto-generate filename: prompt contains no valid words.")
                    args.output = f"output_{int(time.time())}"
            elif args.input_image:
                # Use input filename as base
                base = os.path.splitext(os.path.basename(args.input_image))[0]
                args.output = f"audio_{base}" if args.generate_audio else f"video_{base}"
                print(f"ℹ️  No output specified. Using input basename: {args.output}")
            elif args.transform_image:
                # Use transform input filename as base
                base = os.path.splitext(os.path.basename(args.transform_image))[0]
                args.output = f"{base}_transformed"
                print(f"ℹ️  No output specified. Using transform basename: {args.output}")
            else:
                args.output = f"output_{int(time.time())}"

        # Smart Extension Handling
        if args.format:
            ext = f".{args.format.lower().lstrip('.')}"
            if not args.output.lower().endswith(ext):
                args.output += ext
                print(f"ℹ️  Appended extension '{ext}' to output path.")
        else:
            # No format specified - check if output has an extension
            _, existing_ext = os.path.splitext(args.output)
            if not existing_ext:
                # Auto-append default extension based on mode
                if args.generate_image:
                    args.output += ".jpg"
                elif args.generate_video:
                    args.output += ".mp4"
                elif args.generate_audio:
                    args.output += ".mp3"
                elif args.transform_image:
                    args.output += ".png"

        ensure_paths(args.output)
        
        # Check for existing file
        if os.path.exists(args.output) and not args.force:
            print(f"⚠️  File '{args.output}' already exists.")
            try:
                choice = input(f"   Overwrite? [y/N]: ").lower().strip()
                if choice not in ['y', 'yes']:
                    print("❌ Operation cancelled.")
                    sys.exit(0)
                print("") # Spacer
            except KeyboardInterrupt:
                print("\n❌ Operation cancelled.")
                sys.exit(0)

    if args.generate_image:
        w, h = parse_size(args.size)
        if args.orientation == "portrait": w, h = h, w
        
        outfile = args.output  # Already set by centralized auto-filename logic
        
        # Check resources
        if not check_resources_and_warn(args.image_model, w, h, force=args.force):
           sys.exit(0)
           
        generate_image(args.prompt, outfile, w, h, args.image_model, unsafe=args.unsafe)
        
        if args.upscale or args.simple_upscale:
            base, ext = os.path.splitext(outfile)
            suffix = "simple_upscaled" if args.simple_upscale else "upscaled"
            upscale_out = args.upscaled_output_file or f"{base}_{suffix}_{args.upscale_factor}{ext}"
            
            if args.simple_upscale:
                simple_upscale_image(outfile, upscale_out, factor=args.upscale_factor)
            else:
                model = args.image_upscaler if hasattr(args, 'image_upscaler') else "realesrgan"
                upscale_image_fast(outfile, upscale_out, factor=args.upscale_factor) 

    elif args.generate_audio:
        dur = parse_duration(args.length)
        outfile = args.output  # Already set by centralized auto-filename logic
             
        if args.input_image:
             if not os.path.exists(args.input_image):
                 print(f"❌ Error: Input file not found: {args.input_image}")
                 sys.exit(1)
             
             print(f"👁️  Analyzing image: {args.input_image}...")
             caption = generate_description(args.input_image, model_name=args.caption_model)
             if not caption:
                 print("   Failed to generate description.")
                 sys.exit(1)
             print(f"   📝 Caption: '{caption}'")
             
             full_prompt = caption
             if args.prompt:
                 full_prompt = f"{args.prompt}. {caption}"
             
             print(f"   🎶 Generating audio for: '{full_prompt}'")
             generate_audio(full_prompt, outfile, dur, args.audio_model)
             
        else:
             generate_audio(args.prompt, outfile, dur, args.audio_model)

    elif args.generate_video:
        w, h = parse_size(args.size) if args.size else (576, 320)
        dur = parse_duration(args.length)
        outfile = args.output  # Already set by centralized auto-filename logic
        
        generate_video(
            prompt=args.prompt,
            output_file=outfile,
            duration=dur,
            width=w, 
            height=h,
            model_name=args.video_model,
            upscale=args.upscale,
            simple_upscale=args.simple_upscale,
            upscale_factor=args.upscale_factor
        )

    # --- Standalone Modes ---
    elif args.generate_description:
        if args.generate_description != "USE_INPUT_IMAGE":
            input_file = args.generate_description
        else:
            input_file = args.input_image
             
        if not input_file:
             print("❌ Error: Input image/video required (-gd FILE or -gd -ii FILE)")
             sys.exit(1)
             
        desc = generate_description(input_file, model_name=args.caption_model)
        if desc:
            print(f"📝 Description: {desc}")
            
    elif args.upscale_image:
        outfile = args.output
        if not outfile:
             base, ext = os.path.splitext(args.upscale_image)
             outfile = f"{base}_upscaled{ext}"
             
        if args.simple_upscale:
             simple_upscale_image(args.upscale_image, outfile, factor=args.upscale_factor)
        elif args.image_upscaler == "sd":
             upscale_image_file(args.upscale_image, outfile, args.upscale_strength, factor=args.upscale_factor)
        else:
             upscale_image_fast(args.upscale_image, outfile, factor=args.upscale_factor)

    elif args.upscale_video:
         outfile = args.output
         if not outfile:
             base, ext = os.path.splitext(args.upscale_video)
             outfile = f"{base}_upscaled{ext}"
             
         if args.simple_upscale:
             simple_upscale_video(args.upscale_video, outfile, factor=args.upscale_factor)
         else:
             upscale_video_fast(args.upscale_video, outfile, factor=args.upscale_factor)

    elif args.transform_image:
         if args.transform_image != "USE_GENERATED":
             current_input = args.transform_image
         else:
             print("❌ Error: Transformation requires an input file.")
             sys.exit(1)
             
         if args.transform_prompt:
             # Edit
             generate_edit(current_input, args.transform_prompt, args.output or "edit.png", model_name="instruct-pix2pix", image_guidance_scale=args.image_guidance)
         elif args.remove_background:
             remove_background(current_input, args.output or "nobg.png", silhouette=args.silhouette)

    elif args.convert_image:
         outfile = args.convert_image_to or args.output or "out.png"
         convert_media(args.convert_image, outfile)
    elif args.convert_video:
         outfile = args.convert_video_to or args.output or "out.mp4"
         convert_media(args.convert_video, outfile)
    elif args.convert_audio:
         outfile = args.convert_audio_to or args.output or "out.mp3"
         convert_media(args.convert_audio, outfile)
    elif args.convert_document:
         outfile = args.convert_document_to or args.output or "out.pdf"
         convert_document_file(args.convert_document, outfile)
if __name__ == "__main__":
    main()
